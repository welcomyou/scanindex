"""
PDF to DOCX Position-Preserving Converter v4

Key Features:
1. Creates a NEW blank DOCX from the final OCR PDF
2. Extracts text with positions from PDF using PyMuPDF
3. Recreates text in DOCX with correct positions (using line breaks/paragraphs)
4. All text is Times New Roman 14pt
5. Tables detected by Camelot are preserved as tables

Input: _ocr.pdf (final OCR text with positions)
Output: _final.docx (clean document with preserved layout)
"""

import os
# FORCE CPU ONLY
os.environ["CUDA_VISIBLE_DEVICES"] = ""

import hashlib
import json
import re
import fitz  # PyMuPDF
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Set
from docx import Document
from docx.shared import Pt, Inches, Twips, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime



# Img2Table is retained only for historical benchmark scripts. Production table
# extraction is DocLayout bbox + GMFT-ONNX + Docling TableFormer v1 step-cache
# ONNX, so the runtime does not import img2table or use it as a fallback.
PDF = None
img2table_Image = None
IMG2TABLE_AVAILABLE = False

# GMFT-ONNX is part of the production DocLayout-anchored table pipeline.
try:
    import importlib
    _gmft_onnx_engine = importlib.import_module("scanindex.core.tables.gmft_onnx_table_engine")
    detect_tables_gmft_onnx = _gmft_onnx_engine.detect_tables_gmft_onnx
    detect_tables_gmft_onnx_on_layout_regions = _gmft_onnx_engine.detect_tables_gmft_onnx_on_layout_regions
    GMFT_ONNX_AVAILABLE = _gmft_onnx_engine.is_gmft_onnx_available()
except ImportError:
    detect_tables_gmft_onnx = None
    detect_tables_gmft_onnx_on_layout_regions = None
    GMFT_ONNX_AVAILABLE = False

# Docling TableFormer v1 accurate ONNX is used as a PyTorch-free structure recognizer on
# DocLayout table boxes. It is not used as the primary detector in production.
try:
    from scanindex.core.tables.docling_tableformer_v1_onnx_engine import (
        detect_tables_docling_tableformer_v1_onnx,
        is_docling_tableformer_v1_onnx_available,
    )
    DOCLING_TABLEFORMER_AVAILABLE = is_docling_tableformer_v1_onnx_available()
except ImportError:
    detect_tables_docling_tableformer_v1_onnx = None
    DOCLING_TABLEFORMER_AVAILABLE = False

# RapidTable/Wired variants are benchmark-only now. They are intentionally not
# imported by production runtime to keep the PDF-to-DOCX path on the selected
# GMFT + Docling v1 step-cache ONNX pipeline.
detect_tables_rapidtable_slanet = None
RAPIDTABLE_AVAILABLE = False

# Legacy PyTorch GMFT is opt-in for dev comparison only.
GMFT_AVAILABLE = False
if os.environ.get("OCRTOOL_ALLOW_PYTORCH_GMFT") == "1":
    try:
        import sys
        from scanindex.infra.paths import get_base_dir
        legacy_gmft_dir = os.path.join(
            get_base_dir(),
            "temp",
            "legacy_model_train_20260504",
            "root_cleanup",
            "antigravity-gmft",
        )
        sys.path.insert(0, legacy_gmft_dir)
        from gmft_table_engine import detect_tables_gmft, is_gmft_available
        GMFT_AVAILABLE = is_gmft_available()
    except ImportError:
        GMFT_AVAILABLE = False
else:
    detect_tables_gmft = None
    GMFT_AVAILABLE = False


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class TextSpan:
    """A span of text with formatting info."""
    text: str
    font_size: float
    y: float  # y position
    is_superscript: bool = False
    x: float = 0.0
    width: float = 0.0
    fg_gray: int = 128
    has_space_after: bool = True


@dataclass
class TextLine:
    """A line of text from PDF with position."""
    text: str
    x: float
    y: float
    width: float
    height: float
    page: int
    font_size: float = 12.0  # Average font size of spans
    is_footnote: bool = False
    spans: List[TextSpan] = None  # Individual spans for superscript detection
    block_id: int = 0          # Screen AI layout block group
    paragraph_id: int = 0      # Paragraph within block
    content_type: int = 0      # 0=printed, 1=handwritten, 4=separator, 8=signature
    fg_gray: int = 128         # Foreground grayscale (lower = darker/bolder)
    confidence: float = 0.0    # OCR confidence [0,1]
    semantic_type: str = ""    # From DocLayout-YOLO: title, text, figure, header, footer...
    order: int = 0             # OCR reading order within the page
    source_line_id: str = ""    # Canonical OCR line id, when available
    kie_labels: Set[str] = field(default_factory=set)

    @property
    def y_center(self) -> float:
        return self.y + self.height / 2

    def __post_init__(self):
        if self.spans is None:
            self.spans = []


@dataclass
class TableRegion:
    """A table detected in PDF with precise coordinates."""
    page: int
    y_top: float      # Visual Y from top
    y_bottom: float   # Visual Y from top
    cells: List[List[str]]  # Retain for structure/fallback
    row_count: int
    col_count: int
    cell_bboxes: List[List[Tuple[float, float, float, float]]] = field(default_factory=list) # [x0, y0, x1, y1] Visual coords


# ============================================================================
# LOGGING
# ============================================================================

class Logger:
    def __init__(self, log_path: Optional[str]):
        self.log_path = log_path
        self.lines = []
        
    def log(self, msg: str):
        # timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S") 
        # User requested to remove internal timestamp as GUI adds its own
        line = msg 
        self.lines.append(line)
        try:
            print(line)
        except:
            pass
        
    def save(self):
        if self.log_path:
            with open(self.log_path, 'w', encoding='utf-8') as f:
                f.write("\n".join(self.lines))
            
    def get_log_text(self) -> str:
        return "\n".join(self.lines)


# ============================================================================
# PDF TEXT EXTRACTION WITH POSITIONS
# ============================================================================

def _clean_extracted_text(text: str, strip: bool = True) -> str:
    """Normalize PDF/Word extraction artifacts without changing OCR wording."""
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    text = (
        text.replace("\u00a0", " ")
            .replace("\u202f", " ")
            .replace("\ufeff", "")
            .replace("\u00ad", "-")
            .replace("\u2010", "-")
            .replace("\u2011", "-")
            .replace("\u2012", "-")
            .replace("\u2013", "-")
            .replace("\u2014", "-")
    )
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip() if strip else text


def _json_line_xywh(line: dict) -> Tuple[float, float, float, float]:
    if all(k in line and line.get(k) is not None for k in ("x", "y", "w", "h")):
        try:
            return (
                float(line.get("x", 0.0)),
                float(line.get("y", 0.0)),
                float(line.get("w", 0.0)),
                float(line.get("h", 0.0)),
            )
        except (TypeError, ValueError):
            pass

    bbox = line.get("bbox") or [0.0, 0.0, 0.0, 0.0]
    if len(bbox) < 4:
        return (0.0, 0.0, 0.0, 0.0)
    try:
        x0, y0, x1, y1 = (float(v) for v in bbox[:4])
    except (TypeError, ValueError):
        return (0.0, 0.0, 0.0, 0.0)
    return (x0, y0, max(0.0, x1 - x0), max(0.0, y1 - y0))


def _load_canonical_companion_data(path: str, logger: Logger, purpose: str) -> Tuple[Optional[dict], Optional[Path]]:
    """Load the canonical `.json.zst` sidecar for a PDF or sidecar-like path."""
    if not path:
        return None, None
    try:
        from scanindex.core.canonical_io import load_canonical, resolve_companion

        resolved = resolve_companion(path)
        if resolved is None:
            return None, None
        return load_canonical(resolved), resolved
    except Exception as e:
        logger.log(f"Could not load OCR JSON for {purpose}: {e}")
        return None, None


def load_lines_from_companion_json(json_path: str, logger: Logger) -> Optional[List[TextLine]]:
    """Load OCR text/positions directly from the canonical JSON when present."""
    ocr_data, _resolved = _load_canonical_companion_data(json_path, logger, "text source")
    if not ocr_data:
        return None

    all_lines: List[TextLine] = []
    kie_labels_by_line = _line_labels_from_annotations(ocr_data)
    for page_idx, page in enumerate(ocr_data.get("pages", []), 1):
        words_by_line: Dict[str, List[dict]] = {}
        for word in page.get("words", []) or []:
            line_id = str(word.get("line_id") or "")
            if not line_id:
                continue
            wx, wy, ww, wh = _json_line_xywh(word)
            if ww <= 0 or wh <= 0:
                continue
            words_by_line.setdefault(line_id, []).append({
                "text": _clean_extracted_text(word.get("text") or word.get("ocr_text") or ""),
                "x": wx,
                "y": wy,
                "w": ww,
                "h": wh,
                "order": int(word.get("order", 0) or 0),
                "has_space_after": bool(word.get("has_space_after", True)),
                "fg_gray": int(word.get("fg_gray", 128) or 128),
                "confidence": float(word.get("confidence", 0.0) or 0.0),
                "source_layer": str(word.get("source_layer") or ""),
            })

        for line in page.get("lines", []):
            text = _clean_extracted_text(line.get("text") or line.get("ocr_text") or "")
            if not text:
                continue
            source_line_id = str(line.get("id") or "")
            x, y, w, h = _json_line_xywh(line)
            if w <= 0 or h <= 0:
                continue
            font_size = float(line.get("font_size") or max(h * 0.78, 4.0))
            text_line = TextLine(
                text=text,
                x=x,
                y=y,
                width=w,
                height=h,
                page=page_idx,
                font_size=font_size,
                spans=[TextSpan(text=text, font_size=font_size, y=y)],
                block_id=int(line.get("block_id", 0) or 0),
                paragraph_id=int(line.get("paragraph_id", 0) or 0),
                content_type=int(line.get("content_type", 0) or 0),
                fg_gray=int(line.get("fg_gray", 128) or 128),
                confidence=float(line.get("confidence", 0.0) or 0.0),
                order=int(line.get("order", len(all_lines)) or 0),
                source_line_id=source_line_id,
                kie_labels=set(kie_labels_by_line.get(source_line_id, set())),
            )
            setattr(text_line, "source_layer", str(line.get("source_layer") or page.get("text_source") or ""))
            line_words = words_by_line.get(str(line.get("id") or ""), [])
            if line_words:
                line_words.sort(key=lambda item: item["order"])
                setattr(text_line, "word_items", line_words)
                text_line.spans = [
                    TextSpan(
                        text=item["text"],
                        font_size=font_size,
                        y=item["y"],
                        x=item["x"],
                        width=item["w"],
                        fg_gray=int(item.get("fg_gray", 128) or 128),
                        has_space_after=bool(item.get("has_space_after", True)),
                    )
                    for item in line_words
                    if item.get("text")
                ]
            all_lines.append(text_line)

    if not all_lines:
        return None

    all_lines.sort(key=lambda l: (l.page, l.order, l.y, l.x))
    logger.log(f"Using companion OCR JSON text source: {len(all_lines)} lines")
    return all_lines


def _line_bbox(line: TextLine) -> Tuple[float, float, float, float]:
    return (line.x, line.y, line.x + line.width, line.y + line.height)


def _bbox_overlap_ratio(a: Tuple[float, float, float, float],
                        b: Tuple[float, float, float, float]) -> float:
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0 = max(ax0, bx0)
    iy0 = max(ay0, by0)
    ix1 = min(ax1, bx1)
    iy1 = min(ay1, by1)
    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0
    area_a = max((ax1 - ax0) * (ay1 - ay0), 1e-6)
    return ((ix1 - ix0) * (iy1 - iy0)) / area_a


def _unaccent_upper(text: str) -> str:
    text = (text or "").upper().replace("Đ", "D")
    text = unicodedata.normalize("NFD", text)
    return "".join(ch for ch in text if unicodedata.category(ch) != "Mn")


def _line_labels_from_annotations(ocr_data: dict) -> Dict[str, Set[str]]:
    """Map canonical OCR line ids to KIE labels from annotations."""
    labels_by_line: Dict[str, Set[str]] = {}
    ann = ocr_data.get("annotations") or {}
    fields = ann.get("field_instances") or []
    if not fields:
        return labels_by_line

    word_to_line: Dict[str, str] = {}
    for page in ocr_data.get("pages", []) or []:
        for word in page.get("words", []) or []:
            word_id = str(word.get("id") or "")
            line_id = str(word.get("line_id") or "")
            if word_id and line_id:
                word_to_line[word_id] = line_id

    for inst in fields:
        label = str(inst.get("label") or "").strip()
        if not label:
            continue
        line_ids = {str(line_id) for line_id in (inst.get("line_ids") or []) if str(line_id)}
        for word_id in inst.get("word_ids") or []:
            line_id = word_to_line.get(str(word_id))
            if line_id:
                line_ids.add(line_id)
        for line_id in line_ids:
            labels_by_line.setdefault(line_id, set()).add(label)
    return labels_by_line


def _mostly_uppercase_text(text: str, threshold: float = 0.65) -> bool:
    letters = [ch for ch in text if ch.isalpha()]
    if len(letters) < 3:
        return False
    upper = sum(1 for ch in letters if ch.upper() == ch and ch.lower() != ch)
    lower = sum(1 for ch in letters if ch.lower() == ch and ch.upper() != ch)
    return upper > 0 and upper / max(upper + lower, 1) >= threshold


def _looks_like_numbered_heading(text: str) -> bool:
    """Detect left-aligned section headings such as "II. ...", "4. ..." or "5.1. ..."."""
    stripped = _clean_extracted_text(text)
    roman_match = re.match(r"^\s*([IVXLCDM]{1,8}\.)\s+(.+)$", stripped)
    match = roman_match or re.match(r"^\s*(\d{1,3}(?:\.\d{1,3}){0,4}\.?)\s+(.+)$", stripped)
    if not match:
        return False
    marker = match.group(1).strip()
    body = match.group(2).strip()
    if not body or len(body) > 160:
        return False
    if body.endswith(":"):
        return True
    if _mostly_uppercase_text(body):
        return True
    words = [w for w in re.split(r"\s+", body) if w]
    if body.endswith((".", ";", ",")):
        return False
    if roman_match:
        return len(words) <= 14
    if "." in marker.rstrip("."):
        return len(words) <= 14 and len(body) <= 120
    return len(words) <= 14 and len(body) <= 90


def _looks_like_list_item(text: str) -> bool:
    stripped = _clean_extracted_text(text)
    return is_numbered_bullet(stripped)


def _starts_with_dash_marker(text: str) -> bool:
    return bool(re.match(r"^\s*[-–]\s+", text or ""))


def _dash_marker_can_continue_previous_sentence(prev_text: str, next_text: str) -> bool:
    prev_text = (prev_text or "").rstrip()
    next_text = (next_text or "").lstrip()
    if not prev_text or not _starts_with_dash_marker(next_text):
        return False
    if prev_text[-1] in '.?!:;)]}"':
        return False
    after_dash = re.sub(r"^\s*[-–]\s+", "", next_text, count=1).strip()
    return bool(after_dash and after_dash[0].isalpha())


def _looks_like_visual_bold_heading(
    text: str,
    semantic_type: str,
    numbered_heading: bool,
    is_doc_subject: bool,
    is_centered: bool,
) -> bool:
    """Conservative paragraph-level bold detection for OCR-exported DOCX."""
    stripped = _clean_extracted_text(text)
    if not stripped:
        return False
    if numbered_heading or is_doc_subject:
        return True

    sem = (semantic_type or "").strip().lower()
    if sem == "title":
        if len(stripped) > 160:
            return False
        if _looks_like_list_item(stripped) and len(stripped) > 120:
            return False
        return is_centered or _mostly_uppercase_text(stripped)

    return is_centered and len(stripped) <= 140 and _mostly_uppercase_text(stripped)


def _looks_like_header_emphasis_line(text: str) -> bool:
    """Detect short formal header lines by typography, not by wording."""
    stripped = _clean_extracted_text(text)
    if not stripped:
        return False
    letters = [ch for ch in stripped if ch.isalpha()]
    if len(letters) < 4:
        return False
    if not _mostly_uppercase_text(stripped, threshold=0.55):
        return False
    words = [w for w in re.split(r"\s+", stripped) if re.search(r"[A-Za-zÀ-Ỵà-ỵĐđ]", w)]
    if len(words) > 8:
        return False
    digit_count = sum(1 for ch in stripped if ch.isdigit())
    if digit_count and digit_count >= max(2, len(letters) * 0.15):
        return False
    if stripped.endswith((".", ",", ";", ":")):
        return False
    return True


def _looks_like_person_name(text: str) -> bool:
    stripped = _clean_extracted_text(text)
    if not stripped or re.search(r"\d", stripped):
        return False
    if any(ch in stripped for ch in ":;,/\\"):
        return False
    letters = [ch for ch in stripped if ch.isalpha()]
    if letters and not any(ch.islower() for ch in letters):
        return False
    words = [w for w in re.split(r"\s+", stripped) if re.search(r"[A-Za-zÀ-Ỵà-ỵĐđ]", w)]
    if not (2 <= len(words) <= 5):
        return False
    titlecase_words = 0
    for word in words:
        clean = word.strip("().,")
        is_initial = len(clean) == 1 and clean.isalpha() and clean.upper() == clean
        if len(clean) < 2 and not is_initial:
            return False
        first_alpha = next((ch for ch in clean if ch.isalpha()), "")
        if first_alpha and first_alpha.upper() == first_alpha:
            titlecase_words += 1
    return titlecase_words >= max(2, len(words) - 1)


def _is_signature_footer_line(text: str) -> bool:
    return _looks_like_header_emphasis_line(text) or _looks_like_person_name(text)


def _looks_like_qr_access_artifact(text: str) -> bool:
    normalized = _unaccent_upper(_clean_extracted_text(text))
    if "QR" not in normalized:
        return False
    return any(token in normalized for token in ("MA QR", "TRUY CAP", "DUONG LINK", "LINK", "THONG KE"))


def _line_overlaps_layout_type(line: TextLine,
                               layout_regions_by_page: Dict[int, List[dict]],
                               region_type: str,
                               min_overlap: float = 0.35) -> bool:
    for region in layout_regions_by_page.get(line.page, []):
        if region.get("type") != region_type:
            continue
        bbox = region.get("bbox_pdf")
        if not bbox or len(bbox) < 4:
            continue
        if _bbox_overlap_ratio(_line_bbox(line), tuple(float(v) for v in bbox[:4])) >= min_overlap:
            return True
    return False


def filter_figure_ocr_noise(lines: List[TextLine],
                            layout_regions_by_page: Dict[int, List[dict]],
                            logger: Logger) -> List[TextLine]:
    if not lines or not layout_regions_by_page:
        return lines

    lines_by_page: Dict[int, List[TextLine]] = {}
    for line in lines:
        lines_by_page.setdefault(line.page, []).append(line)

    figure_bboxes_by_page: Dict[int, List[Tuple[float, float, float, float]]] = {}
    for page, regions in (layout_regions_by_page or {}).items():
        page_lines = lines_by_page.get(page, [])
        max_right = max((line.x + line.width for line in page_lines), default=0.0)
        noi_nhan_anchors = [
            line for line in page_lines
            if _unaccent_upper(line.text).startswith("NOI NHAN")
        ]
        for region in regions or []:
            if _layout_region_kind(region) != "figure":
                continue
            bbox = _layout_region_bbox_pdf(region)
            if not bbox:
                continue
            fig_center_x = (bbox[0] + bbox[2]) / 2.0
            looks_like_signature = False
            if noi_nhan_anchors:
                anchor_y = max(line.y for line in noi_nhan_anchors)
                looks_like_signature = fig_center_x >= max_right * 0.45 and bbox[1] >= anchor_y - 30.0
            if looks_like_signature:
                continue
            figure_bboxes_by_page.setdefault(page, []).append(bbox)

    if not figure_bboxes_by_page:
        return lines

    kept: List[TextLine] = []
    removed = 0
    for line in lines:
        sem = (line.semantic_type or "").strip().lower()
        text = _clean_extracted_text(line.text)
        is_caption = (
            sem in {"figure_caption", "figure caption", "table_caption", "table caption"}
            or bool(re.match(r"^(figure|fig\.?|table)\s*\d+\s*[:.\-]", text, re.IGNORECASE))
        )
        if is_caption:
            kept.append(line)
            continue
        line_bbox = _line_bbox(line)
        inside_figure = False
        for fig_bbox in figure_bboxes_by_page.get(line.page, []):
            if _bbox_overlap_ratio(line_bbox, fig_bbox) >= 0.35 or _layout_center_inside(line_bbox, fig_bbox, pad=3.0):
                inside_figure = True
                break
        if inside_figure:
            removed += 1
            continue
        kept.append(line)

    if removed:
        logger.log(f"Removed {removed} OCR line(s) inside rendered figure regions")
    return kept


def enrich_lines_from_json(lines: List[TextLine], json_path: str, logger: Logger):
    """
    Enrich TextLine objects with Screen AI metadata from companion JSON.
    Matches lines by bbox overlap/proximity and assigns:
    block_id, paragraph_id, content_type, confidence, fg_gray.
    """
    ocr_data, _resolved = _load_canonical_companion_data(json_path, logger, "enrichment")
    if not ocr_data:
        return

    pages = ocr_data.get("pages", [])
    kie_labels_by_line = _line_labels_from_annotations(ocr_data)
    enriched = 0

    for tl in lines:
        page_idx = tl.page - 1  # TextLine.page is 1-based
        if page_idx < 0 or page_idx >= len(pages):
            continue

        json_lines = pages[page_idx].get("lines", [])
        if not json_lines:
            continue

        best_match = None
        best_score = None
        tl_bbox = _line_bbox(tl)
        tl_cx = tl.x + tl.width / 2.0
        tl_cy = tl.y + tl.height / 2.0
        for jl in json_lines:
            jx, jy, jw, jh = _json_line_xywh(jl)
            if jw <= 0 or jh <= 0:
                continue
            jb = (jx, jy, jx + jw, jy + jh)
            overlap = _bbox_overlap_ratio(tl_bbox, jb)
            jcx = jx + jw / 2.0
            jcy = jy + jh / 2.0
            x_dist = abs(tl_cx - jcx) / max(tl.width, jw, 1.0)
            y_dist = abs(tl_cy - jcy) / max(tl.height, jh, 1.0)

            if overlap > 0:
                score = 10.0 * overlap - x_dist - y_dist
            elif x_dist <= 0.35 and y_dist <= 0.75:
                score = 1.0 - x_dist - y_dist
            else:
                continue

            if best_score is None or score > best_score:
                best_score = score
                best_match = jl

        if best_match and best_score is not None and best_score > 0:
            json_text = _clean_extracted_text(best_match.get("text") or "")
            if json_text:
                tl.text = json_text
            source_line_id = str(best_match.get("id") or "")
            tl.source_line_id = source_line_id
            tl.kie_labels = set(kie_labels_by_line.get(source_line_id, set()))
            tl.block_id = best_match.get("block_id", 0)
            tl.paragraph_id = best_match.get("paragraph_id", 0)
            tl.content_type = best_match.get("content_type", 0)
            tl.confidence = best_match.get("confidence", 0.0)
            tl.fg_gray = best_match.get("fg_gray", 128)
            tl.order = int(best_match.get("order", tl.order) or tl.order)
            # Override geometry with OCR bbox (more accurate than rendered text).
            # direct_ocr_engine renders line as single span with approximate font size,
            # so PyMuPDF's span bbox may not match the real visual extent.
            ocr_x, ocr_y, ocr_w, ocr_h = _json_line_xywh(best_match)
            if ocr_w > 0 and ocr_h > 0:
                tl.x = ocr_x
                tl.y = ocr_y
                tl.width = ocr_w
                tl.height = ocr_h
            enriched += 1

    if enriched > 0:
        logger.log(f"Enriched {enriched}/{len(lines)} lines with Screen AI metadata")


def extract_pdf_lines(pdf_path: str, logger: Logger) -> Tuple[List[TextLine], dict]:
    """
    Extract all text lines from PDF with their positions.
    Returns: (lines, page_info) where page_info contains dimensions.
    """
    doc = fitz.open(pdf_path)
    all_lines = []
    page_info = {}
    
    for page_num, page in enumerate(doc, 1):
        page_info[page_num] = {
            "width": page.rect.width,
            "height": page.rect.height
        }

    json_lines = load_lines_from_companion_json(pdf_path, logger)
    if json_lines is not None:
        doc.close()
        logger.log(f"PDF: {len(page_info)} pages, {len(json_lines)} OCR JSON text lines loaded")
        return json_lines, page_info

    for page_num, page in enumerate(doc, 1):
        
        # Extract text as dict for detailed position info
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if "lines" not in block:
                continue
            
            for line in block["lines"]:
                line_text = ""
                x0, y0, x1, y1 = float('inf'), float('inf'), 0, 0
                font_sizes = []
                text_spans = []
                
                for span in line["spans"]:
                    text = _clean_extracted_text(span["text"], strip=False)
                    line_text += text
                    if text.strip():
                        bbox = span["bbox"]
                        x0 = min(x0, bbox[0])
                        y0 = min(y0, bbox[1])
                        x1 = max(x1, bbox[2])
                        y1 = max(y1, bbox[3])
                        span_size = span.get("size", 12)
                        span_y = bbox[1]
                        font_sizes.append(span_size)
                        text_spans.append(TextSpan(
                            text=text,
                            font_size=span_size,
                            y=span_y,
                            x=float(bbox[0]),
                            width=max(0.0, float(bbox[2]) - float(bbox[0])),
                            fg_gray=128,
                            has_space_after=text.endswith(" "),
                        ))
                
                avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
                
                # Detect superscripts: significantly smaller font + only digits
                # Compare each span's font to the LINE's average (not doc average)
                for tspan in text_spans:
                    is_small = tspan.font_size < avg_font_size * 0.85  # 90% threshold
                    is_digits = tspan.text.strip().isdigit()
                    if is_small and is_digits:
                        tspan.is_superscript = True
                
                # Collect all lines
                if line_text.strip():
                    all_lines.append(TextLine(
                        text=line_text.strip(),
                        x=x0,
                        y=y0,
                        width=x1 - x0,
                        height=y1 - y0,
                        page=page_num,
                        font_size=avg_font_size,
                        spans=text_spans
                    ))
    
    doc.close()

    # Sort by page, then by reading order (XY-Cut handles multi-column)
    sorted_lines = []
    pages = sorted(set(l.page for l in all_lines))
    for pg in pages:
        pg_lines = [l for l in all_lines if l.page == pg]
        sorted_lines.extend(xy_cut_sort(pg_lines))
    all_lines = sorted_lines

    logger.log(f"PDF: {len(page_info)} pages, {len(all_lines)} text lines extracted")

    return all_lines, page_info


def xy_cut_sort(lines, depth=0):
    """
    Recursive XY-Cut algorithm for correct reading order.
    Handles multi-column layouts by detecting vertical gaps (columns)
    and horizontal gaps (sections).
    """
    if len(lines) <= 1:
        return lines
    if depth > 20:  # prevent infinite recursion
        return sorted(lines, key=lambda l: (l.y, l.x))

    # Compute bounding region
    min_x = min(l.x for l in lines)
    max_x = max(l.x + l.width for l in lines)
    region_width = max_x - min_x

    # Find largest horizontal gap (split top/bottom)
    h_gap, h_pos = _find_largest_gap([l.y for l in lines], [l.height for l in lines])

    # Find largest vertical gap (split left/right = multi-column)
    # Exclude narrow outliers (< 10% of region width) to avoid page numbers skewing
    main_lines = [l for l in lines if l.width > region_width * 0.1] if region_width > 0 else lines
    if main_lines:
        v_gap, v_pos = _find_largest_gap([l.x for l in main_lines], [l.width for l in main_lines])
    else:
        v_gap, v_pos = 0, 0

    min_gap = 5.0  # minimum gap to consider a split (in PDF points)

    if max(h_gap, v_gap) < min_gap:
        return sorted(lines, key=lambda l: (l.y, l.x))

    if h_gap >= v_gap:
        # Horizontal cut: split into top and bottom
        top = [l for l in lines if l.y + l.height / 2 < h_pos]
        bottom = [l for l in lines if l.y + l.height / 2 >= h_pos]
        if not top or not bottom:
            return sorted(lines, key=lambda l: (l.y, l.x))
        return xy_cut_sort(top, depth + 1) + xy_cut_sort(bottom, depth + 1)
    else:
        # Vertical cut: split into left and right (multi-column!)
        left = [l for l in lines if l.x + l.width / 2 < v_pos]
        right = [l for l in lines if l.x + l.width / 2 >= v_pos]
        if not left or not right:
            return sorted(lines, key=lambda l: (l.y, l.x))
        return xy_cut_sort(left, depth + 1) + xy_cut_sort(right, depth + 1)


def _find_largest_gap(positions, sizes):
    """
    Find the largest gap between elements along one axis.
    Returns (gap_size, gap_midpoint).
    """
    if not positions:
        return 0, 0
    # Create sorted list of (start, end) intervals
    intervals = sorted(zip(positions, sizes), key=lambda x: x[0])
    best_gap = 0
    best_pos = 0
    for i in range(1, len(intervals)):
        prev_end = intervals[i - 1][0] + intervals[i - 1][1]
        curr_start = intervals[i][0]
        gap = curr_start - prev_end
        if gap > best_gap:
            best_gap = gap
            best_pos = (prev_end + curr_start) / 2
    return best_gap, best_pos


# ============================================================================
# LINE MERGING INTO PARAGRAPHS
# ============================================================================

def is_numbered_bullet(text: str) -> bool:
    """
    Check if text starts with a numbered bullet:
    - 1., 2., 3. (but NOT 18.000 which is a number)
    - I., II., IV. (roman numerals)
    - a), b), c) or 1), 2), 3)
    """
    stripped = _clean_extracted_text(text or "")
    if re.match(r'^\d{4}\.$', stripped):
        return False

    # Pattern for "1.", "2.", "1.1.", "1.2.3." etc - must be followed by space or end
    # Distinguish from "18.000" by checking what follows the dot
    if re.match(r'^\d+(?:\.\d+)*\.\s', stripped) or re.match(r'^\d+(?:\.\d+)*\.$', stripped):
        return True
    
    # Roman numerals: I., II., III., IV., V., VI., etc.
    if re.match(r'^[IVX]+\.\s', stripped) or re.match(r'^[IVX]+\.$', stripped):
        return True
    
    # a), b), c), đ) style. [^\W\d_] is a single Unicode letter.
    if re.match(r'^[^\W\d_]\)\s', stripped) or re.match(r'^[^\W\d_]\)$', stripped):
        return True
    
    # 1), 2), 3) style
    if re.match(r'^\d+\)\s', stripped) or re.match(r'^\d+\)$', stripped):
        return True

    # (1), (2), (a) style used in Vietnamese administrative lists
    if re.match(r'^\(\d+\)\s', stripped) or re.match(r'^\([^\W\d_]\)\s', stripped):
        return True
    
    return False


@dataclass
class ParagraphEdgeDecision:
    split: bool
    split_score: float
    merge_score: float
    reasons: List[str] = field(default_factory=list)
    features: Dict[str, object] = field(default_factory=dict)


def _terminal_punctuation(text: str) -> bool:
    return bool((text or "").rstrip().endswith(tuple('.?!:;)]}"')))


def _hard_terminal_punctuation(text: str) -> bool:
    return bool((text or "").rstrip().endswith(tuple('?!:')))


def _content_width(base_x: float, right_margin: float) -> float:
    return max(right_margin - base_x, 1.0)


def _median(values: List[float], default: float = 0.0) -> float:
    vals = sorted(float(v) for v in values if v is not None)
    if not vals:
        return default
    mid = len(vals) // 2
    if len(vals) % 2:
        return vals[mid]
    return (vals[mid - 1] + vals[mid]) / 2.0


def _percentile(values: List[float], q: float, default: float = 0.0) -> float:
    vals = sorted(float(v) for v in values if v is not None)
    if not vals:
        return default
    q = max(0.0, min(1.0, float(q)))
    idx = q * (len(vals) - 1)
    lo = int(idx)
    hi = min(lo + 1, len(vals) - 1)
    frac = idx - lo
    return vals[lo] * (1.0 - frac) + vals[hi] * frac


def _clamp(value: float, low: float, high: float) -> float:
    return max(low, min(high, value))


def _line_centered_in_content_band(line: TextLine, left_margin: float, right_margin: float) -> bool:
    content_width = _content_width(left_margin, right_margin)
    line_center = line.x + line.width / 2.0
    page_center = left_margin + content_width / 2.0
    left_gap = max(0.0, line.x - left_margin)
    right_gap = max(0.0, right_margin - (line.x + line.width))
    return (
        line.width < content_width * 0.82
        and abs(line_center - page_center) < content_width * 0.10
        and abs(left_gap - right_gap) < content_width * 0.18
    )


def _looks_like_standalone_centered_title_line(
    text: str,
    line: TextLine,
    left_margin: float,
    right_margin: float,
) -> bool:
    stripped = _clean_extracted_text(text)
    if not stripped or len(stripped) > 80:
        return False
    content_width = _content_width(left_margin, right_margin)
    words = [w for w in re.split(r"\s+", stripped) if w]
    return (
        len(words) <= 6
        and line.width < content_width * 0.42
        and _mostly_uppercase_text(stripped)
        and _line_centered_in_content_band(line, left_margin, right_margin)
    )


def _line_style_compatible_for_heading_wrap(line1: TextLine, line2: TextLine) -> bool:
    h1 = max(float(getattr(line1, "height", 0.0) or 0.0), 1.0)
    h2 = max(float(getattr(line2, "height", 0.0) or 0.0), 1.0)
    height_ratio = min(h1, h2) / max(h1, h2)
    try:
        gray1 = int(getattr(line1, "fg_gray", 128))
        gray2 = int(getattr(line2, "fg_gray", 128))
        gray_compatible = (
            not (_has_known_gray(gray1) and _has_known_gray(gray2))
            or abs(gray1 - gray2) <= 48
        )
    except Exception:
        gray_compatible = True
    return height_ratio >= 0.72 and gray_compatible


def _looks_like_short_heading_tail(line: TextLine, next_width: float) -> bool:
    text = _clean_extracted_text(line.text or "")
    if not text:
        return False
    words = [word for word in re.split(r"\s+", text) if word]
    return len(words) <= 3 and line.width <= next_width * 0.28


def _looks_like_structural_heading_continuation(
    line1: TextLine,
    line2: TextLine,
    prev_base: float,
    prev_right: float,
    next_base: float,
    next_right: float,
    numbered_heading_prev: bool,
    numbered_next: bool,
    numbered_heading_next: bool,
    dot_bullet_next: bool,
    dash_next: bool,
) -> bool:
    """Detect wrapped heading lines using OCR block/style/layout, not fixed wording."""
    if not numbered_heading_prev:
        return False
    if getattr(line1, "block_id", 0) <= 0 or line1.block_id != getattr(line2, "block_id", 0):
        return False
    if line1.page != line2.page:
        return False
    text1 = (line1.text or "").rstrip()
    text2 = (line2.text or "").strip()
    if not text1 or not text2:
        return False
    if _hard_terminal_punctuation(text1):
        return False
    if numbered_next or numbered_heading_next or dot_bullet_next or dash_next:
        return False

    gap = line2.y - (line1.y + line1.height)
    compact_gap = gap <= max(line1.height, line2.height) * 0.75
    if not compact_gap:
        return False

    prev_width = _content_width(prev_base, prev_right)
    next_width = _content_width(next_base, next_right)
    prev_reaches_right = (line1.x + line1.width) >= prev_right - max(30.0, prev_width * 0.075)
    next_has_heading_width = line2.width >= next_width * 0.45
    next_in_content_band = (
        abs(line2.x - next_base) <= max(42.0, next_width * 0.10)
        or abs(line2.x - line1.x) <= max(42.0, next_width * 0.10)
    )
    if not (prev_reaches_right or next_has_heading_width):
        return False
    if not next_in_content_band:
        return False

    if _looks_like_short_heading_tail(line2, next_width):
        return True

    return _line_style_compatible_for_heading_wrap(line1, line2)


def score_paragraph_edge(
    line1: TextLine,
    line2: TextLine,
    base_x: float,
    right_margin: float,
    logger: Logger = None,
    page_info: dict = None,
    margin_map: Dict[int, Tuple[float, float]] = None,
) -> ParagraphEdgeDecision:
    """
    Score whether the edge between two OCR text lines is a paragraph break.

    This keeps the old public behavior deterministic, but makes the decision
    inspectable: geometry, lexical markers, and page-boundary evidence vote
    toward split or merge instead of relying on a single brittle condition.
    """
    text1 = (line1.text or "").rstrip()
    text2 = (line2.text or "").strip()
    reasons: List[str] = []
    features: Dict[str, object] = {}
    split_score = 0.0
    merge_score = 0.0

    if not text1 or not text2:
        return ParagraphEdgeDecision(True, 100.0, 0.0, ["empty_text"], features)

    sem2 = (getattr(line2, "semantic_type", "") or "").strip().lower()
    layout_type2 = (getattr(line2, "_layout_region_type", "") or "").strip().lower()
    if (
        sem2 in {"figure_caption", "figure caption", "table_caption", "table caption"}
        or layout_type2 in {"figure_caption", "figure caption", "table_caption", "table caption"}
        or bool(re.match(r"^(figure|fig\.?|table)\s*\d+\s*[:.\-]", text2, re.IGNORECASE))
    ):
        return ParagraphEdgeDecision(
            True,
            100.0,
            0.0,
            ["caption_boundary"],
            {"semantic_type2": sem2, "layout_type2": layout_type2},
        )

    multi_col1 = getattr(line1, "_multi_column_role", None) or getattr(line1, "_digital_column", None)
    multi_col2 = getattr(line2, "_multi_column_role", None) or getattr(line2, "_digital_column", None)
    if (
        line1.page == line2.page
        and (getattr(line1, "_multi_column_page", False) or getattr(line1, "_digital_two_column_page", False))
        and (getattr(line2, "_multi_column_page", False) or getattr(line2, "_digital_two_column_page", False))
        and multi_col1
        and multi_col2
        and multi_col1 != multi_col2
    ):
        return ParagraphEdgeDecision(
            True,
            100.0,
            0.0,
            ["multi_column_boundary"],
            {"multi_col1": multi_col1, "multi_col2": multi_col2},
        )

    layout_region1 = getattr(line1, "_layout_region_order", None)
    layout_region2 = getattr(line2, "_layout_region_order", None)
    if (
        line1.page == line2.page
        and layout_region1 is not None
        and layout_region2 is not None
        and layout_region1 != layout_region2
    ):
        return ParagraphEdgeDecision(
            True,
            100.0,
            0.0,
            ["layout_region_boundary"],
            {
                "layout_region1": layout_region1,
                "layout_region2": layout_region2,
                "layout_type1": getattr(line1, "_layout_region_type", ""),
                "layout_type2": getattr(line2, "_layout_region_type", ""),
            },
        )

    prev_base, prev_right = (margin_map or {}).get(line1.page, (base_x, right_margin))
    next_base, next_right = (margin_map or {}).get(line2.page, (base_x, right_margin))
    prev_width = _content_width(prev_base, prev_right)
    next_width = _content_width(next_base, next_right)
    indent_threshold = max(24.0, next_width * 0.055)
    right_threshold = max(30.0, prev_width * 0.075)

    line1_right_edge = line1.x + line1.width
    same_page = line1.page == line2.page
    adjacent_page = line2.page == line1.page + 1
    prev_reaches_right = line1_right_edge >= prev_right - right_threshold
    prev_short = line1.width / prev_width < 0.25
    near_same_left = abs(line2.x - line1.x) <= indent_threshold
    next_body_left = (
        abs(line2.x - next_base) <= max(36.0, next_width * 0.08)
        or abs(line2.x - line1.x) <= max(42.0, next_width * 0.08)
    )
    has_start_indent = line2.x > next_base + indent_threshold
    terminal = _terminal_punctuation(text1)
    hard_terminal = _hard_terminal_punctuation(text1)
    numbered_next = is_numbered_bullet(text2)
    numbered_heading_next = _looks_like_numbered_heading(text2)
    numbered_heading_prev = _looks_like_numbered_heading(text1)
    dash_next = _starts_with_dash_marker(text2)
    dash_continuation = (
        _dash_marker_can_continue_previous_sentence(text1, text2)
        and not numbered_heading_prev
    )
    dot_bullet_next = text2.startswith("\u2022")
    starts_lower = text2[0].islower()
    starts_alpha = text2[0].isalpha()
    starts_digit = text2[0].isdigit()
    prev_centered_title = _looks_like_standalone_centered_title_line(text1, line1, prev_base, prev_right)
    next_centered = _line_centered_in_content_band(line2, next_base, next_right)
    structural_heading_wrap = _looks_like_structural_heading_continuation(
        line1,
        line2,
        prev_base,
        prev_right,
        next_base,
        next_right,
        numbered_heading_prev,
        numbered_next,
        numbered_heading_next,
        dot_bullet_next,
        dash_next,
    )

    features.update(
        {
            "same_page": same_page,
            "adjacent_page": adjacent_page,
            "prev_reaches_right": prev_reaches_right,
            "prev_short": prev_short,
            "near_same_left": near_same_left,
            "next_body_left": next_body_left,
            "has_start_indent": has_start_indent,
            "terminal": terminal,
            "hard_terminal": hard_terminal,
            "numbered_next": numbered_next,
            "numbered_heading_next": numbered_heading_next,
            "numbered_heading_prev": numbered_heading_prev,
            "dash_next": dash_next,
            "dash_continuation": dash_continuation,
            "dot_bullet_next": dot_bullet_next,
            "prev_centered_title": prev_centered_title,
            "next_centered": next_centered,
            "structural_heading_wrap": structural_heading_wrap,
        }
    )

    if numbered_next or numbered_heading_next:
        split_score += 8.0
        reasons.append("next_line_is_list_or_numbered_heading")
    if dot_bullet_next:
        split_score += 8.0
        reasons.append("next_line_is_dot_bullet")
    if dash_next:
        if numbered_heading_prev:
            split_score += 8.0
            reasons.append("dash_list_after_numbered_heading")
        elif dash_continuation:
            merge_score += 6.0
            reasons.append("dash_continues_unfinished_sentence")
        else:
            split_score += 6.0
            reasons.append("dash_starts_list_item")

    if prev_centered_title and (next_centered or line2.width >= next_width * 0.35):
        split_score += 6.0
        reasons.append("standalone_centered_title")

    if numbered_heading_prev:
        same_heading_wrap = (
            structural_heading_wrap
            or (
                near_same_left
                and _mostly_uppercase_text(text1)
                and _mostly_uppercase_text(text2)
                and not _terminal_punctuation(text2)
            )
        )
        if same_heading_wrap:
            merge_score += 4.0
            reasons.append("numbered_heading_wrap")
        else:
            split_score += 5.0
            reasons.append("previous_line_is_complete_numbered_heading")

    if not same_page:
        if not adjacent_page or not page_info:
            split_score += 6.0
            reasons.append("non_adjacent_or_unknown_page_boundary")
        else:
            prev_page_height = page_info.get(line1.page, {}).get("height", 842)
            next_page_height = page_info.get(line2.page, {}).get("height", 842)
            prev_footnote_top = page_info.get(line1.page, {}).get("footnote_top_y")
            prev_before_footnote_band = False
            if prev_footnote_top:
                footnote_gap = max(line1.height * 2.0, prev_page_height * 0.035)
                prev_before_footnote_band = (
                    line1.y < prev_footnote_top
                    and (line1.y + line1.height) >= prev_footnote_top - footnote_gap
                )
            prev_near_bottom = (
                (line1.y + line1.height) >= prev_page_height * 0.68
                or prev_before_footnote_band
            )
            next_near_top = line2.y <= next_page_height * 0.35
            features["prev_near_bottom"] = prev_near_bottom
            features["next_near_top"] = next_near_top
            features["prev_before_footnote_band"] = prev_before_footnote_band
            if prev_near_bottom and next_near_top:
                merge_score += 2.0
                reasons.append(
                    "adjacent_page_boundary_before_footnote"
                    if prev_before_footnote_band
                    else "adjacent_page_boundary_geometry"
                )
            else:
                split_score += 4.0
                reasons.append("not_page_boundary_continuation_band")

            if terminal:
                split_score += 5.0
                reasons.append("previous_text_has_terminal_punctuation")
            else:
                merge_score += 2.0
                reasons.append("previous_text_unfinished")

            if prev_reaches_right:
                merge_score += 3.0
                reasons.append("previous_line_reaches_body_right")
            else:
                split_score += 1.5
                reasons.append("previous_line_ends_short")

            if next_body_left:
                merge_score += 2.0
                reasons.append("next_line_starts_in_body_band")
            else:
                split_score += 2.5
                reasons.append("next_line_outside_body_band")

            if starts_lower:
                merge_score += 2.5
                reasons.append("next_line_starts_lowercase")
            elif starts_alpha and not terminal:
                merge_score += 1.0
                reasons.append("next_line_alpha_after_unfinished_text")
            elif starts_digit:
                split_score += 1.0
                reasons.append("next_line_starts_digit")
    else:
        gap = line2.y - (line1.y + line1.height)
        gap_ratio = gap / max(line1.height, 1.0)
        features["vertical_gap"] = gap
        features["vertical_gap_ratio"] = gap_ratio

        if gap_ratio > 1.7:
            split_score += 4.0
            reasons.append("large_vertical_gap")
        elif gap_ratio <= 0.9:
            merge_score += 0.75
            reasons.append("normal_line_gap")

        if not has_start_indent and prev_reaches_right:
            merge_score += 3.0
            reasons.append("block_wrap_geometry")
        if near_same_left and prev_reaches_right:
            merge_score += 1.5
            reasons.append("same_left_and_previous_full")

        if has_start_indent:
            if starts_lower or (starts_digit and not numbered_next and not terminal):
                merge_score += 2.0
                reasons.append("indented_continuation_after_unfinished_text")
            else:
                split_score += 2.5
                reasons.append("new_start_indent")

        if not prev_reaches_right:
            if hard_terminal:
                split_score += 3.0
                reasons.append("short_previous_line_with_hard_terminal")
            elif terminal:
                split_score += 2.0
                reasons.append("short_previous_line_with_terminal")
            elif prev_short:
                split_score += 2.5
                reasons.append("very_short_previous_line")
            else:
                merge_score += 2.0
                reasons.append("unfinished_short_previous_line")

        if hard_terminal:
            split_score += 1.5
            reasons.append("hard_terminal_punctuation")
        elif not terminal:
            merge_score += 1.0
            reasons.append("no_terminal_punctuation")
        if starts_lower:
            merge_score += 2.0
            reasons.append("next_line_starts_lowercase")

    split = split_score > merge_score
    decision = ParagraphEdgeDecision(split, split_score, merge_score, reasons, features)
    if logger and os.environ.get("OCRTOOL_DEBUG_PARAGRAPH_EDGE") == "1":
        logger.log(
            "Paragraph edge: "
            f"{'SPLIT' if decision.split else 'MERGE'} "
            f"split={decision.split_score:.2f} merge={decision.merge_score:.2f} "
            f"reasons={','.join(decision.reasons)}"
        )
    return decision


def should_split_paragraph(line1: TextLine, line2: TextLine, base_x: float, right_margin: float, logger: Logger = None) -> bool:
    """Compatibility wrapper for callers that only need the split/merge bool."""
    return score_paragraph_edge(line1, line2, base_x, right_margin, logger).split


def _set_merged_last_line(first_line: TextLine, last_line: TextLine) -> None:
    setattr(first_line, "_merged_last_line", last_line)


def _get_merged_last_line(first_line: TextLine) -> TextLine:
    return getattr(first_line, "_merged_last_line", first_line)


def _set_merged_lines(first_line: TextLine, lines: List[TextLine]) -> None:
    setattr(first_line, "_merged_lines", list(lines or [first_line]))


def _get_merged_lines(first_line: TextLine) -> List[TextLine]:
    lines = getattr(first_line, "_merged_lines", None)
    if lines:
        return list(lines)
    return [first_line]


def _should_heal_paragraph_fragment(
    prev_text: str,
    prev_first_line: TextLine,
    next_text: str,
    next_first_line: TextLine,
    margin_map: Dict[int, Tuple[float, float]],
    page_info: dict,
) -> bool:
    """Merge OCR/block fragments that are clearly one running sentence."""
    prev_text = (prev_text or "").rstrip()
    next_text = (next_text or "").lstrip()
    if not prev_text or not next_text:
        return False
    if getattr(prev_first_line, "is_footnote", False) or getattr(next_first_line, "is_footnote", False):
        return False
    if is_numbered_bullet(next_text) or next_text[0] in '•':
        return False
    if next_text[0] in '-–' and not _dash_marker_can_continue_previous_sentence(prev_text, next_text):
        return False
    if _looks_like_numbered_heading(next_text):
        return False

    prev_last_line = _get_merged_last_line(prev_first_line)
    base_x, right_margin = margin_map.get(next_first_line.page, (0, 500))
    decision = score_paragraph_edge(
        prev_last_line,
        next_first_line,
        base_x,
        right_margin,
        page_info=page_info,
        margin_map=margin_map,
    )
    return not decision.split


def _heal_fragmented_paragraphs(
    paragraphs: List[Tuple[str, TextLine, bool]],
    margin_map: Dict[int, Tuple[float, float]],
    page_info: dict,
    logger: Logger = None,
) -> List[Tuple[str, TextLine, bool]]:
    if not paragraphs:
        return paragraphs

    healed: List[Tuple[str, TextLine, bool]] = []
    merged_count = 0
    for para in paragraphs:
        if not healed:
            healed.append(para)
            continue

        prev_text, prev_first_line, prev_is_footnote = healed[-1]
        next_text, next_first_line, next_is_footnote = para
        if (
            not prev_is_footnote
            and not next_is_footnote
            and _should_heal_paragraph_fragment(
                prev_text,
                prev_first_line,
                next_text,
                next_first_line,
                margin_map,
                page_info,
            )
        ):
            healed[-1] = (
                prev_text.rstrip() + " " + next_text.lstrip(),
                prev_first_line,
                prev_is_footnote,
            )
            _set_merged_lines(
                prev_first_line,
                _get_merged_lines(prev_first_line) + _get_merged_lines(next_first_line),
            )
            _set_merged_last_line(prev_first_line, _get_merged_last_line(next_first_line))
            merged_count += 1
        else:
            healed.append(para)

    if logger and merged_count:
        logger.log(f"Healed {merged_count} OCR paragraph fragments across block/page boundaries")
    return healed


def is_page_number_text(text: str) -> bool:
    """Check if text looks like a page number."""
    text = text.strip()
    if not text:
        return False
    if len(text) >= 5:
        return False
    # Only digits, or digits + special chars
    if re.match(r'^[\d!@#$%^&*():;,.\-]+$', text):
        # Must have at least one digit
        return bool(re.search(r'\d', text))
    return False


def is_page_number(line: TextLine, page_lines: List[TextLine]) -> bool:
    """
    Check if line is a page number:
    - Only digits, or digits + special chars (!@#$%^&*():...)
    - Less than 5 characters
    - First or last line of the page
    """
    text = line.text.strip()
    
    if not is_page_number_text(text):
        return False
    
    # Must have at least one digit
    if not re.search(r'\d', text):
        return False
    
    # Relaxed rule: if it's very short (< 3 chars) and looks like a number/symbol, allow it to be filtered
    # regardless of position (catches noise like "7:", "1:")
    if len(text) < 3:
        return True

    # Must be first or last line of the page
    same_page_lines = [l for l in page_lines if l.page == line.page]
    if not same_page_lines:
        return False
    
    same_page_lines.sort(key=lambda l: l.y)
    is_first = line == same_page_lines[0]
    is_last = line == same_page_lines[-1]
    
    return is_first or is_last


def filter_page_numbers(lines: List[TextLine], logger: Logger) -> List[TextLine]:
    """Remove page number lines."""
    filtered = [line for line in lines if not is_page_number(line, lines)]
    removed = len(lines) - len(filtered)
    if removed > 0:
        logger.log(f"Removed {removed} page number lines")
    return filtered


def _has_spanning_header_title_region(page: int,
                                      page_width: float,
                                      page_height: float,
                                      layout_regions_by_page: Optional[Dict[int, List[dict]]]) -> bool:
    for region in (layout_regions_by_page or {}).get(page, []):
        if region.get("type") != "title":
            continue
        bbox = region.get("bbox_pdf")
        if not bbox or len(bbox) < 4:
            continue
        x0, y0, x1, y1 = (float(v) for v in bbox[:4])
        if (
            y0 <= page_height * 0.14
            and x0 <= page_width * 0.18
            and x1 >= page_width * 0.82
            and (y1 - y0) <= page_height * 0.08
        ):
            return True
    return False


def detect_dual_column_headers(lines: List[TextLine],
                               page_info: dict,
                               logger: Logger,
                               layout_regions_by_page: Optional[Dict[int, List[dict]]] = None) -> Dict[int, Tuple[List[TextLine], List[TextLine]]]:
    """
    Detect dual-column headers at top of pages (common in VN government docs).
    Heuristic: two groups of lines in top 20% of page, one left of center, one right,
    with overlapping Y ranges.
    Returns: {page_num: (left_lines, right_lines)}
    """
    def looks_like_dual_header(left_lines: List[TextLine], right_lines: List[TextLine]) -> bool:
        left_emphasis = [l for l in left_lines if _looks_like_header_emphasis_line(l.text)]
        right_emphasis = [l for l in right_lines if _looks_like_header_emphasis_line(l.text)]
        if not left_emphasis or not right_emphasis:
            return False

        left_y0 = min(l.y for l in left_emphasis)
        left_y1 = max(l.y + l.height for l in left_emphasis)
        right_y0 = min(l.y for l in right_emphasis)
        right_y1 = max(l.y + l.height for l in right_emphasis)
        overlaps_vertically = min(left_y1, right_y1) >= max(left_y0, right_y0) - 24
        starts_near_same_band = abs(left_y0 - right_y0) <= 60
        return overlaps_vertically and starts_near_same_band

    result = {}
    pages = set(l.page for l in lines)
    for pg in pages:
        pg_height = page_info.get(pg, {}).get("height", 842)
        pg_width = page_info.get(pg, {}).get("width", 595)
        page_center = pg_width / 2
        half_width = page_center
        if _has_spanning_header_title_region(pg, pg_width, pg_height, layout_regions_by_page):
            continue

        # Step 1: Find all lines in the top area
        y_threshold = pg_height * 0.20
        top_lines = [l for l in lines if l.page == pg and l.y < y_threshold]
        if len(top_lines) < 3:
            continue

        # Step 2: Split into true left/right header zones. Centered titles
        # can cross the page center, so require right-column lines to start
        # near the right half instead of using center-point alone.
        all_left = sorted([
            l for l in top_lines
            if l.x + l.width / 2 < page_center - 8 and l.width < half_width * 0.95
        ], key=lambda l: l.y)
        all_right = sorted([
            l for l in top_lines
            if l.x >= page_center - 15 and l.width < half_width * 0.95
        ], key=lambda l: l.y)

        if not all_left or not all_right:
            continue
        if not looks_like_dual_header(all_left, all_right):
            continue

        # Step 3: Only keep left lines that overlap Y range with right group
        # with a small allowance for left-only "Số..." lines under the header.
        ry_max = max(l.y + l.height for l in all_right)
        y_tolerance = 30
        left = [l for l in all_left if l.y < ry_max + y_tolerance]
        right = all_right

        if not left or not right:
            continue

        # Step 4: Verify this is a real dual-column HEADER (not body text)
        # Both groups must have SHORT lines (< 60% of their half-width)
        # Long body text lines would span most of their half
        left_avg_ratio = sum(l.width / half_width for l in left) / len(left) if left else 1
        right_avg_ratio = sum(l.width / half_width for l in right) / len(right) if right else 1
        if left_avg_ratio > 0.85 or right_avg_ratio > 0.85:
            continue  # Lines too wide — this is body text, not header

        ry_min = min(l.y for l in right)
        ly_max = max(l.y + l.height for l in left)
        overlap = min(ly_max, ry_max) - max(min(l.y for l in left), ry_min)
        if overlap > 0:
            result[pg] = (left, right)
            logger.log(f"Detected dual-column header on page {pg}: {len(left)}L + {len(right)}R lines")

    return result


def detect_dual_column_footers(lines: List[TextLine], page_info: dict, logger: Logger) -> Dict[int, Tuple[List[TextLine], List[TextLine]]]:
    """
    Detect bottom recipient/signature blocks rendered as two columns.
    """
    def norm_text(text: str) -> str:
        text = unicodedata.normalize("NFD", text.upper())
        return "".join(ch for ch in text if unicodedata.category(ch) != "Mn")

    result = {}
    pages = set(l.page for l in lines)
    for pg in pages:
        pg_height = page_info.get(pg, {}).get("height", 842)
        pg_width = page_info.get(pg, {}).get("width", 595)
        page_center = pg_width / 2
        page_lines = [l for l in lines if l.page == pg]
        anchors = [l for l in page_lines if norm_text(l.text).startswith("NOI NHAN")]
        if not anchors:
            continue

        anchor = sorted(anchors, key=lambda l: l.y)[-1]
        top_y = max(0, anchor.y - 12)
        bottom_y = min(pg_height * 0.95, anchor.y + 180)
        candidates = [l for l in page_lines if top_y <= l.y <= bottom_y]

        left = sorted(
            [l for l in candidates if l.x + l.width / 2 < page_center],
            key=lambda l: (l.y, l.x)
        )
        right = sorted(
            [l for l in candidates if l.x + l.width / 2 >= page_center],
            key=lambda l: (l.y, l.x)
        )
        right = [l for l in right if _is_signature_footer_line(l.text)]
        if not left or not right:
            continue

        has_signature_title = any(_looks_like_header_emphasis_line(l.text) for l in right)
        has_signature_name = any(_looks_like_person_name(l.text) for l in right)
        if not has_signature_title or not has_signature_name:
            continue

        result[pg] = (left, right)
        logger.log(f"Detected dual-column footer on page {pg}: {len(left)}L + {len(right)}R lines")

    return result


def _borderless_table_xml(table) -> None:
    """Remove table/cell borders that Word may inherit from default styles."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in list(tblPr.findall(qn("w:tblBorders"))):
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for existing in list(tcPr.findall(qn("w:tcBorders"))):
                tcPr.remove(existing)
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'none')
                el.set(qn('w:sz'), '0')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), 'auto')
                tc_borders.append(el)
            tcPr.append(tc_borders)


def add_dual_header_table(doc: Document, left_lines: List[TextLine], right_lines: List[TextLine],
                          page_width_pt: float, logger: Logger):
    """Render dual-column header as a borderless 1x2 table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _borderless_table_xml(table)

    # Column widths are based on text extents, scaled to the DOCX content width.
    # Using raw PDF x positions here can make the DOCX table wider than the
    # writable area and causes Word/layout renderers to reflow it unpredictably.
    content_width = max(float(page_width_pt or 0.0), 120.0)
    left_w = max((float(l.width) for l in left_lines), default=content_width / 2.0)
    right_w = max((float(l.width) for l in right_lines), default=content_width / 2.0)
    total = left_w + right_w
    if total > 0:
        left_pt = content_width * left_w / total
        right_pt = content_width - left_pt
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(int(content_width * 20)))
        for row in table.rows:
            row.cells[0].width = Pt(left_pt)
            row.cells[1].width = Pt(right_pt)

    right_emphasis_indices = [
        idx for idx, line in enumerate(right_lines)
        if _looks_like_header_emphasis_line(line.text)
    ]
    right_underline_idx = right_emphasis_indices[0] if right_emphasis_indices else None

    # Left cell
    cell_l = table.cell(0, 0)
    cell_l.text = ""
    for i, line in enumerate(left_lines):
        para = cell_l.paragraphs[0] if i == 0 else cell_l.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(line.text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        if _looks_like_header_emphasis_line(line.text):
            run.font.bold = True

    # Right cell
    cell_r = table.cell(0, 1)
    cell_r.text = ""
    for i, line in enumerate(right_lines):
        para = cell_r.paragraphs[0] if i == 0 else cell_r.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(line.text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        if _looks_like_header_emphasis_line(line.text):
            run.font.bold = True
        if i == right_underline_idx:
            run.font.underline = True


def _coalesce_admin_header_fragments(lines: List[TextLine]) -> List[Tuple[str, TextLine]]:
    rows: List[Tuple[str, TextLine]] = []
    for line in sorted(lines, key=lambda item: (item.y, item.x)):
        text = _clean_extracted_text(line.text)
        if not text:
            continue
        if rows and re.fullmatch(r"\d{1,2}", text):
            prev_text, prev_line = rows[-1]
            if re.search(r"\bngày\s+tháng\b", prev_text, re.IGNORECASE):
                rows[-1] = (
                    re.sub(r"\bngày\s+tháng\b", f"ngày {text} tháng", prev_text, count=1, flags=re.IGNORECASE),
                    prev_line,
                )
                continue
            if re.match(r"^Số\s*-", prev_text, re.IGNORECASE):
                rows[-1] = (
                    re.sub(r"^Số\s*", f"Số {text} ", prev_text, count=1, flags=re.IGNORECASE),
                    prev_line,
                )
                continue
        rows.append((text, line))
    return rows


def add_dual_header_paragraphs(doc: Document, left_lines: List[TextLine], right_lines: List[TextLine],
                               content_width_pt: float, logger: Logger):
    """Render an administrative two-sided header as editable tabbed paragraphs."""
    left_sorted = _coalesce_admin_header_fragments(left_lines)
    right_sorted = _coalesce_admin_header_fragments(right_lines)
    row_count = max(len(left_sorted), len(right_sorted))
    right_emphasis_indices = {
        idx for idx, (_text, line) in enumerate(right_sorted)
        if _looks_like_header_emphasis_line(line.text)
    }
    for idx in range(row_count):
        left_line = left_sorted[idx] if idx < len(left_sorted) else None
        right_line = right_sorted[idx] if idx < len(right_sorted) else None
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0
        fmt.tab_stops.add_tab_stop(Pt(max(float(content_width_pt), 120.0)), WD_TAB_ALIGNMENT.RIGHT)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        if left_line is not None:
            left_text, left_source = left_line
            run = para.add_run(left_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12.5)
            if _looks_like_header_emphasis_line(left_source.text):
                run.font.bold = True
        if right_line is not None:
            right_text, right_source = right_line
            para.add_run("\t")
            run = para.add_run(right_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12.5)
            if _looks_like_header_emphasis_line(right_source.text):
                run.font.bold = True
            if idx in right_emphasis_indices:
                run.font.underline = True


def detect_footnotes(lines: List[TextLine], page_info: dict, logger: Logger) -> List[TextLine]:
    """
    Detect footnotes conservatively.

    OCR often reports wrapped tail fragments with smaller bboxes/font sizes,
    such as the last word of a normal paragraph. Those fragments must stay
    body text, not footnotes.
    """
    if not lines:
        return lines
    
    font_sizes = [line.font_size for line in lines if line.font_size > 0]
    avg_font = sum(font_sizes) / len(font_sizes) if font_sizes else 12
    
    footnote_count = 0

    def has_footnote_marker(text: str) -> bool:
        stripped = text.strip()
        return bool(
            re.match(r'^\d{1,2}\s+\S', stripped)
            or re.match(r'^[*]\s+\S', stripped)
        )

    def is_wrapped_body_fragment(prev_line: Optional[TextLine], line: TextLine) -> bool:
        if prev_line is None or prev_line.page != line.page:
            return False
        stripped = line.text.strip()
        if not stripped or has_footnote_marker(stripped) or is_numbered_bullet(stripped):
            return False
        if stripped.startswith(("-", "+")):
            return False

        if (
            line.block_id > 0
            and prev_line.block_id == line.block_id
            and prev_line.paragraph_id == line.paragraph_id
        ):
            return True

        prev_text = prev_line.text.rstrip()
        if not prev_text or prev_text.endswith((".", ":", ";", "?", "!", ")")):
            return False
        gap = line.y - (prev_line.y + prev_line.height)
        max_h = max(prev_line.height, line.height, 1.0)
        return -max_h * 0.25 <= gap <= max_h * 1.5

    prev_by_page: Dict[int, Optional[TextLine]] = {}
    
    for line in lines:
        page_height = page_info.get(line.page, {}).get("height", 800)
        page_width = page_info.get(line.page, {}).get("width", 595)
        prev_line = prev_by_page.get(line.page)
        prev_by_page[line.page] = line
        sem = (line.semantic_type or "").strip().lower()
        text_stripped = line.text.strip()
        
        # Primary criteria
        is_very_small_font = line.font_size < avg_font * 0.75
        is_small_font = line.font_size < avg_font * 0.88
        
        in_lower_part = line.y > page_height * 0.70
        in_bottom_band = line.y > page_height * 0.82
        
        # Bonus criteria
        starts_with_num = has_footnote_marker(line.text)
        if sem == "footnote":
            if line.y > page_height * 0.55:
                line.is_footnote = True
                footnote_count += 1
                continue
            line.is_footnote = False
            continue
        if prev_line is not None and getattr(prev_line, "is_footnote", False):
            gap = line.y - (prev_line.y + prev_line.height)
            max_h = max(prev_line.height, line.height, 1.0)
            x_close = abs(line.x - prev_line.x) <= page_width * 0.12
            x_continuation_indent = prev_line.x <= line.x <= prev_line.x + page_width * 0.16
            if (
                is_small_font
                and in_lower_part
                and -max_h * 0.25 <= gap <= max_h * 1.8
                and line.x <= page_width * 0.55
                and (x_close or x_continuation_indent)
            ):
                line.is_footnote = True
                footnote_count += 1
                continue
        text_lower = text_stripped.lower()
        is_recipient_footer = (
            text_lower.startswith("nơi ")
            or text_lower.startswith("noi ")
            or text_stripped.startswith("-")
            or text_lower == "trân trọng."
        )

        if (
            is_recipient_footer
            or is_numbered_bullet(text_stripped)
            or (
                sem in {
                    "title",
                    "plain text",
                    "table",
                    "table_caption",
                    "figure_caption",
                    "page-header",
                    "page-footer",
                }
                and not starts_with_num
            )
            or is_wrapped_body_fragment(prev_line, line)
        ):
            line.is_footnote = False
            continue

        if line.x > page_width * 0.50 and not starts_with_num:
            line.is_footnote = False
            continue

        if sem == "table_footnote" and is_small_font and in_lower_part:
            line.is_footnote = True
            footnote_count += 1
            continue
        
        # Score-based detection
        score = 0
        if is_very_small_font:
            score += 60
        elif is_small_font:
            score += 40
            
        if in_bottom_band:
            score += 35
        elif in_lower_part:
            score += 30
            
        if starts_with_num:
            score += 35
        elif not in_bottom_band:
            score -= 35
            
        # Logging for debug
        if "Quản lý văn bản điều hành" in line.text:
             logger.log(f"DEBUG FN CHECK: '{line.text[:30]}...'")
             logger.log(f"  Global Avg Font: {avg_font:.2f}")
             logger.log(f"  Line Font: {line.font_size:.2f} ({line.font_size/avg_font*100:.1f}%)")
             logger.log(f"  Is Small (<88%): {is_small_font}")
             logger.log(f"  Score: {score}")
        
        # Require either an explicit marker or true bottom-band placement; small
        # font alone is too noisy for OCR-generated PDFs.
        if score >= 90 and (starts_with_num or in_bottom_band):
            line.is_footnote = True
            footnote_count += 1
    
    if footnote_count > 0:
        logger.log(f"Detected {footnote_count} footnote lines (avg font: {avg_font:.1f})")
    
    return lines




def _normalize_text_for_compare(text: str) -> str:
    return re.sub(r"\s+", " ", _clean_extracted_text(text or "")).strip().lower()


def _has_known_gray(value: object) -> bool:
    try:
        gray = int(value)
    except Exception:
        return False
    return 0 <= gray <= 255 and gray != 128


def _gray_z(gray: float, median: float, iqr: float) -> float:
    return (float(median) - float(gray)) / max(float(iqr), 1.0)


def _line_is_visually_bold(line: TextLine) -> bool:
    """Infer whether the whole OCR line is visually bold."""
    try:
        line_gray = int(line.fg_gray)
    except Exception:
        return False
    if not _has_known_gray(line_gray):
        return False

    stats = getattr(line, "_page_style_stats", {}) or {}
    if stats.get("doc_known_gray_count", 0) < 80:
        return False
    cutoff = stats.get("doc_bold_gray_cutoff")
    if cutoff is None:
        return False

    median = float(stats.get("doc_gray_median", 128.0))
    iqr = max(float(stats.get("doc_gray_iqr", 0.0)), 1.0)
    z_cutoff = float(stats.get("doc_bold_z_cutoff", _gray_z(float(cutoff), median, iqr)))
    return line_gray <= float(cutoff) and _gray_z(line_gray, median, iqr) >= z_cutoff


def _word_is_visually_bold(
    word: dict,
    line: TextLine,
    *,
    allow_line_bold: bool = True,
    allow_word_bold: bool = True,
) -> bool:
    """Infer local bold from robustly-normalized foreground gray."""
    if not word:
        return False
    if allow_line_bold and _line_is_visually_bold(line):
        return True
    if not allow_word_bold:
        return False
    try:
        gray = int(word.get("fg_gray", 128) or 128)
    except Exception:
        return False
    if not _has_known_gray(gray):
        return False

    stats = getattr(line, "_page_style_stats", {}) or {}
    if stats.get("doc_known_gray_count", 0) < 80:
        return False

    doc_median = float(stats.get("doc_gray_median", 128.0))
    doc_iqr = max(float(stats.get("doc_gray_iqr", 0.0)), 1.0)
    doc_cutoff = stats.get("doc_bold_gray_cutoff")
    if doc_cutoff is None:
        return False

    # Groundtruth calibration showed document-level robust normalization is the
    # strongest signal. The cutoff itself remains data-driven: it is the dark
    # tail of the current document, expressed as a z score for comparability.
    doc_z = _gray_z(gray, doc_median, doc_iqr)
    doc_z_cutoff = float(stats.get("doc_bold_z_cutoff", _gray_z(float(doc_cutoff), doc_median, doc_iqr)))
    if gray > float(doc_cutoff) or doc_z < doc_z_cutoff:
        return False

    page_median = float(stats.get("gray_median", doc_median))
    page_iqr = max(float(stats.get("gray_iqr", doc_iqr)), 1.0)
    page_z = _gray_z(gray, page_median, page_iqr)
    page_cutoff = stats.get("page_bold_gray_cutoff")
    page_supports_dark = page_cutoff is None or gray <= float(page_cutoff) or page_z >= 0.75

    return page_supports_dark


def _line_visual_bold_ratio(line: TextLine) -> float:
    words = _word_items_for_line(line)
    if not words:
        return 0.0
    known = [word for word in words if _has_known_gray(word.get("fg_gray", 128))]
    if not known:
        return 0.0
    bold_count = sum(1 for word in known if _word_is_visually_bold(word, line))
    return bold_count / max(len(known), 1)


def _word_items_for_line(line: TextLine) -> List[dict]:
    items = getattr(line, "word_items", None)
    if items:
        return [item for item in items if _clean_extracted_text(str(item.get("text") or ""))]
    return []


def _merged_word_tokens(
    first_line: TextLine,
    *,
    allow_line_bold: bool = True,
    allow_word_bold: bool = True,
) -> List[Tuple[str, bool, bool]]:
    """Return (text, has_space_before, is_bold) tokens for a merged paragraph."""
    tokens: List[Tuple[str, bool, bool]] = []
    previous_has_space_after = False
    for line_idx, line in enumerate(_get_merged_lines(first_line)):
        words = _word_items_for_line(line)
        if not words:
            previous_has_space_after = True
            continue
        for word_idx, word in enumerate(words):
            text = _clean_extracted_text(str(word.get("text") or ""))
            if not text:
                continue
            has_space_before = bool(tokens) and (previous_has_space_after or word_idx == 0 or line_idx > 0)
            tokens.append((
                text,
                has_space_before,
                _word_is_visually_bold(
                    word,
                    line,
                    allow_line_bold=allow_line_bold,
                    allow_word_bold=allow_word_bold,
                ),
            ))
            previous_has_space_after = bool(word.get("has_space_after", True))
        previous_has_space_after = True
    return tokens


def _token_letter_count(text: str) -> int:
    return sum(1 for ch in text if ch.isalpha())


def _is_content_word_token(text: str) -> bool:
    return _token_letter_count(text) >= 2


def _token_ends_sentence(text: str) -> bool:
    return bool(re.search(r"[.!?][)\"']*$", (text or "").strip()))


def _token_ends_clause(text: str) -> bool:
    return bool(re.search(r"[.,:;!?][)\"']*$", (text or "").strip()))


def _looks_like_lead_list_token(text: str) -> bool:
    stripped = (text or "").strip()
    return stripped == "-" or bool(re.match(r"^\(?[0-9A-Za-z]+[\).]$", stripped))


def _next_bold_run_length(tokens: List[Tuple[str, bool, bool]], flags: List[bool], start: int) -> int:
    length = 0
    for idx in range(start, len(tokens)):
        if not flags[idx]:
            break
        length += 1
    return length


def _smooth_bold_tokens(tokens: List[Tuple[str, bool, bool]]) -> List[Tuple[str, bool, bool]]:
    """Stabilize OCR bold decisions using local run and sentence evidence."""
    if not tokens:
        return tokens

    flags = [bool(token[2]) for token in tokens]

    # Very short function-like words are noisy when they are the only dark
    # token in their immediate neighborhood. Clear them before using bold runs
    # as phrase evidence.
    for idx, (token_text, _, _) in enumerate(tokens):
        letters = _token_letter_count(token_text)
        if not flags[idx] or letters == 0 or letters > 3:
            continue
        original_neighbor_bold = (
            (idx > 0 and bool(tokens[idx - 1][2]))
            or (idx + 1 < len(tokens) and bool(tokens[idx + 1][2]))
        )
        if not original_neighbor_bold:
            flags[idx] = False

    # Administrative bullet/list items often bold the lead sentence, but OCR
    # foreground gray can miss a few words inside that sentence. Promote the
    # first sentence only when its own words already show a strong bold majority.
    first_idx = next((idx for idx, (text, _, _) in enumerate(tokens) if (text or "").strip()), 0)
    if first_idx < len(tokens) and _looks_like_lead_list_token(tokens[first_idx][0]):
        sentence_end = None
        for idx in range(first_idx + 1, len(tokens)):
            if _token_ends_sentence(tokens[idx][0]):
                sentence_end = idx
                break
        if sentence_end is not None and sentence_end > first_idx:
            content_indices = [
                idx for idx in range(first_idx + 1, sentence_end + 1)
                if _is_content_word_token(tokens[idx][0])
            ]
            if content_indices:
                bold_count = sum(1 for idx in content_indices if flags[idx])
                if bold_count >= 2 and bold_count / max(len(content_indices), 1) >= 0.45:
                    for idx in range(first_idx, sentence_end + 1):
                        flags[idx] = True

    # Fill short OCR misses inside an otherwise bold phrase.
    for idx, (token_text, _, _) in enumerate(tokens):
        if flags[idx] or _token_letter_count(token_text) < 4:
            continue
        prev_bold = idx > 0 and flags[idx - 1]
        next_bold = idx + 1 < len(tokens) and flags[idx + 1]
        prev2_bold = idx > 1 and flags[idx - 2]
        if prev_bold and next_bold:
            flags[idx] = True
        elif prev_bold and prev2_bold and _token_ends_clause(token_text):
            flags[idx] = True

    # Backfill the first word of a bold phrase when OCR missed it. Keep this
    # conservative: the candidate must be a real content word, and either the
    # following bold run is substantial or the pair forms a short clause.
    for idx, (token_text, _, _) in enumerate(tokens):
        if flags[idx] or _token_letter_count(token_text) < 4:
            continue
        next_run = _next_bold_run_length(tokens, flags, idx + 1)
        boundary_before = (
            idx == 0
            or flags[idx - 1]
            or _token_ends_clause(tokens[idx - 1][0])
            or not _is_content_word_token(tokens[idx - 1][0])
        )
        if next_run >= 3 or (next_run >= 2 and boundary_before):
            flags[idx] = True
        elif (
            next_run == 1
            and idx + 1 < len(tokens)
            and _token_ends_clause(tokens[idx + 1][0])
            and (idx == 0 or _token_ends_clause(tokens[idx - 1][0]) or not _is_content_word_token(tokens[idx - 1][0]))
        ):
            flags[idx] = True

    return [(text, has_space_before, flags[idx]) for idx, (text, has_space_before, _) in enumerate(tokens)]


def _try_add_word_formatted_text(
    para,
    text: str,
    first_line: TextLine,
    apply_run_format,
    *,
    bold: bool,
) -> bool:
    tokens = _merged_word_tokens(first_line, allow_line_bold=bold, allow_word_bold=bold)
    if not tokens:
        return False

    reconstructed_parts = []
    for token_text, has_space_before, _ in tokens:
        if has_space_before and reconstructed_parts:
            reconstructed_parts.append(" ")
        reconstructed_parts.append(token_text)
    reconstructed = "".join(reconstructed_parts)

    expected_norm = _normalize_text_for_compare(text)
    reconstructed_norm = _normalize_text_for_compare(reconstructed)
    if not expected_norm or not reconstructed_norm:
        return False

    if reconstructed_norm != expected_norm:
        return False

    smoothed_tokens: List[Tuple[str, bool, bool]] = []
    run_smoothed_tokens = _smooth_bold_tokens(tokens)
    for idx, (token_text, has_space_before, token_bold) in enumerate(run_smoothed_tokens):
        if token_bold:
            has_digit = any(ch.isdigit() for ch in token_text)
            letters_only = "".join(ch for ch in token_text if ch.isalpha())
            is_short_plain_word = (
                not has_digit
                and bool(letters_only)
                and len(letters_only) <= 12
                and not _mostly_uppercase_text(letters_only, threshold=0.85)
            )
            neighbor_bold = (
                (idx > 0 and run_smoothed_tokens[idx - 1][2])
                or (idx + 1 < len(run_smoothed_tokens) and run_smoothed_tokens[idx + 1][2])
            )
            if is_short_plain_word and not neighbor_bold:
                token_bold = False
        smoothed_tokens.append((token_text, has_space_before, token_bold))

    for token_text, has_space_before, token_bold in smoothed_tokens:
        run_text = (" " if has_space_before else "") + token_text
        run = para.add_run(run_text)
        apply_run_format(run, is_bold=True if bold or token_bold else False)
    return True


def add_text_with_superscripts(para, text: str, first_line: TextLine, is_footnote: bool,
                               bold: bool = False, font_size_pt: int = 14, italic: bool = False):
    """Add text to paragraph, preserving superscripts and OCR-inferred emphasis."""

    def _apply_run_format(run, size=None, is_bold=None, is_italic=None):
        """Apply common formatting to a run."""
        run.font.name = "Times New Roman"
        run.font.size = Pt(size if size else font_size_pt)
        if is_bold is not None:
            run.font.bold = is_bold
        elif bold:
            run.font.bold = True
        if is_italic is not None:
            run.font.italic = is_italic
        elif italic:
            run.font.italic = True

    # Collect superscript positions from spans
    superscript_texts = set()
    if first_line.spans:
        for span in first_line.spans:
            if span.is_superscript and span.text.strip():
                superscript_texts.add(span.text.strip())

    # For footnotes: italic, size 11, leading number superscript
    if is_footnote:
        match = re.match(r'^(\d+)\s*', text)
        if match:
            run = para.add_run(match.group(1))
            _apply_run_format(run, size=11, is_bold=False, is_italic=False)
            run.font.superscript = True
            rest = text[len(match.group(0)):]
            if rest:
                run = para.add_run(" " + rest)
                _apply_run_format(run, size=11, is_bold=False, is_italic=True)
        else:
            run = para.add_run(text)
            _apply_run_format(run, size=11, is_bold=False, is_italic=True)
        return

    # Normal paragraph - check for inline superscripts
    if not superscript_texts:
        if _try_add_word_formatted_text(para, text, first_line, _apply_run_format, bold=bold):
            return
        run = para.add_run(text)
        _apply_run_format(run)
        return

    # Parse text and format superscripts
    i = 0
    while i < len(text):
        found = False
        for sup in superscript_texts:
            if text[i:i+len(sup)] == sup:
                before_ok = (i == 0 or not text[i-1].isdigit())
                after_ok = (i + len(sup) >= len(text) or not text[i + len(sup)].isdigit())
                if before_ok and after_ok:
                    run = para.add_run(sup)
                    _apply_run_format(run, size=9)
                    run.font.superscript = True
                    i += len(sup)
                    found = True
                    break
        if not found:
            next_pos = len(text)
            for sup in superscript_texts:
                pos = text.find(sup, i + 1)
                if pos != -1 and pos < next_pos:
                    next_pos = pos
            normal = text[i:next_pos]
            if normal:
                run = para.add_run(normal)
                _apply_run_format(run)
            i = next_pos if next_pos > i else i + 1


def _repair_drop_cap_join(text: str, first_line: TextLine) -> str:
    first_text = _clean_extracted_text(getattr(first_line, "text", ""))
    if len(first_text) != 1 or not first_text.isalpha() or not first_text.isupper():
        return text
    match = re.match(r"^([A-Z])\s+([a-z]{1,4})(\b.*)$", text or "", re.DOTALL)
    if not match:
        return text
    return match.group(1) + match.group(2) + match.group(3)


def merge_raw_paragraphs(
    lines: List[TextLine],
    margin_map: Dict[int, Tuple[float, float]],
    logger: Logger = None,
    page_info: dict = None,
) -> List[Tuple[str, TextLine, bool]]:
    """
    Core merging logic: split or merge lines based on indentation and gap.
    Returns: List of (text, first_line, is_footnote).
    """
    if not lines:
        return []
        
    raw_paragraphs = []
    current_text = lines[0].text
    current_first_line = lines[0]
    current_lines = [lines[0]]
    
    for i in range(1, len(lines)):
        # Get margins for current line's page
        # Prioritize line 2 (current line) page for margin context
        page = lines[i].page
        base_x, max_right = margin_map.get(page, (0, 500))
        
        decision = score_paragraph_edge(
            lines[i - 1],
            lines[i],
            base_x,
            max_right,
            logger,
            page_info=page_info,
            margin_map=margin_map,
        )
        if decision.split:
            _set_merged_last_line(current_first_line, lines[i - 1])
            _set_merged_lines(current_first_line, current_lines)
            raw_paragraphs.append((current_text, current_first_line, getattr(current_first_line, 'is_footnote', False)))
            current_text = lines[i].text
            current_first_line = lines[i]
            current_lines = [lines[i]]
        else:
            current_text += " " + lines[i].text
            current_lines.append(lines[i])
    
    _set_merged_last_line(current_first_line, lines[-1])
    _set_merged_lines(current_first_line, current_lines)
    raw_paragraphs.append((current_text, current_first_line, getattr(current_first_line, 'is_footnote', False)))
    return raw_paragraphs


def is_numeric_cell(text: str) -> bool:
    """
    Check if cell content should be treated as numeric (Center Aligned).
    Rule: Contains NO alphabet characters (A-Z, a-z, Vietnamese).
    Allowed: Digits, punctuation, symbols.
    """
    if not text:
        return False # Empty -> Default Left? Or irrelevant.
        
    for char in text:
        if char.isalpha():
            return False
    return True


def merge_lines_to_paragraphs(lines: List[TextLine], page_info: dict, logger: Logger) -> List[Tuple[str, TextLine, bool]]:
    """
    Merge lines into paragraphs, handling footnote reordering.
    Returns: List of (merged_text, first_line, is_footnote)
    """
    if not lines:
        return []
    
    # First filter out page numbers
    lines = filter_page_numbers(lines, logger)
    
    # Then detect footnotes
    lines = detect_footnotes(lines, page_info, logger)
    
    margin_map = {}
    
    # Calculate margins per page
    # Group lines by page first
    page_lines_map = {}
    for line in lines:
        if line.page not in page_lines_map:
            page_lines_map[line.page] = []
        page_lines_map[line.page].append(line)
        
    for page, p_lines in page_lines_map.items():
        x_positions = [l.x for l in p_lines]
        if not x_positions:
            margin_map[page] = (0, 500)
            continue
            
        # Left Margin (Base X): Use MODE (Most Frequent) instead of min/percentile
        # This avoids headers/footers/sidebars skewing the "Body Text" margin.
        from collections import Counter
        
        all_lefts = sorted(x_positions)
        
        # Round to nearest 2.0 to group similar indentations
        x_rounded = [round(x / 2.0) * 2.0 for x in all_lefts]
        common = Counter(x_rounded).most_common(1)
        
        if common:
            mode_base_x = common[0][0]
            # Use mode as base_x. 
            # Note: We want the "Main Body" left. 
            # If there are valid paragraphs to the left of the mode (e.g. outdented headers), 
            # using mode might make them look "negative indented"? code usually handles x > base.
            # But for SPLIT logic, `has_start_indent` checks `line.x > base + threshold`.
            # If line.x == mode, and base == mode, then indent is 0. -> Merge. Correct.
            base_x = mode_base_x
        else:
             sorted_lefts = sorted(all_lefts)
             base_x = sorted_lefts[int(len(sorted_lefts) * 0.05)]

        # Right Margin: Use a combination of mode and robust max
        all_rights = sorted([l.x + l.width for l in p_lines])
        robust_max = all_rights[int(len(all_rights) * 0.98)] if all_rights else 500
        
        # Calculate mode of right edges (rounded to 5.0)
        from collections import Counter
        rights_rounded = [round(r / 5.0) * 5.0 for r in all_rights]
        common_rights = Counter(rights_rounded).most_common(3)
        
        # If the most common right edge is significant and slightly less than robust_max,
        # it's likely the "Main Body" justified margin, not the absolute max (which could be a header).
        max_right = robust_max
        if common_rights:
            mode_right, count = common_rights[0]
            # If mode has ≥ 3 lines and is within 100 units of max, and more frequent than absolute max
            if count >= 3 and mode_right > robust_max - 50:
                max_right = mode_right
            elif robust_max > 0:
                max_right = robust_max
        
        margin_map[page] = (base_x, max_right)
        logger.log(f"Page {page} Margins: Left={base_x:.1f}, Right={max_right:.1f} (RobustMax={robust_max:.1f})")

    footnote_lines = [l for l in lines if l.is_footnote]
    merge_page_info = {
        pg: dict(info or {})
        for pg, info in (page_info or {}).items()
    }
    if footnote_lines:
        footnote_tops: Dict[int, float] = {}
        for line in footnote_lines:
            footnote_tops[line.page] = min(footnote_tops.get(line.page, line.y), line.y)
        for pg, top_y in footnote_tops.items():
            merge_page_info.setdefault(pg, dict((page_info or {}).get(pg, {})))
            merge_page_info[pg]["footnote_top_y"] = top_y

    def merge_subset(subset: List[TextLine]) -> List[Tuple[str, TextLine, bool]]:
        if not subset:
            return []
        # If block_id metadata available, pre-group by (page, block_id, paragraph_id).
        # Lines in different blocks should NOT merge (Screen AI already grouped them).
        if any(l.block_id > 0 for l in subset):
            from itertools import groupby
            merged = []
            for key, group in groupby(subset, key=lambda l: (l.page, l.block_id, l.paragraph_id)):
                group_lines = list(group)
                merged.extend(merge_raw_paragraphs(group_lines, margin_map, logger, merge_page_info))
            return _heal_fragmented_paragraphs(merged, margin_map, merge_page_info, logger)
        return _heal_fragmented_paragraphs(
            merge_raw_paragraphs(subset, margin_map, logger, merge_page_info),
            margin_map,
            merge_page_info,
            logger,
        )

    if footnote_lines:
        # Footnotes must not sit between two body fragments. Merge body text
        # without footnotes first, then insert each footnote after the nearest
        # preceding paragraph on the same page.
        body_paragraphs = merge_subset([l for l in lines if not l.is_footnote])
        footnote_paragraphs = merge_subset(footnote_lines)

        insertions: Dict[int, List[Tuple[str, TextLine, bool]]] = {}
        for fn_para in footnote_paragraphs:
            _, fn_line, _ = fn_para
            target_idx = -1
            for idx, (_, body_line, _) in enumerate(body_paragraphs):
                body_last = _get_merged_last_line(body_line)
                starts_before_footnote = (
                    body_line.page < fn_line.page
                    or (body_line.page == fn_line.page and body_line.y <= fn_line.y)
                )
                spans_footnote_page = body_line.page <= fn_line.page <= body_last.page
                ends_before_footnote = (
                    body_last.page < fn_line.page
                    or (body_last.page == fn_line.page and body_last.y <= fn_line.y)
                )
                if starts_before_footnote and (spans_footnote_page or ends_before_footnote):
                    target_idx = idx
            if target_idx < 0:
                for idx, (_, body_line, _) in enumerate(body_paragraphs):
                    body_last = _get_merged_last_line(body_line)
                    if body_last.page <= fn_line.page:
                        target_idx = idx
            if target_idx < 0:
                target_idx = 0
            insertions.setdefault(target_idx, []).append(fn_para)

        final_paragraphs = []
        for idx, para in enumerate(body_paragraphs):
            final_paragraphs.append(para)
            final_paragraphs.extend(insertions.get(idx, []))
    else:
        final_paragraphs = merge_subset(lines)

    logger.log(f"Merged {len(lines)} lines into {len(final_paragraphs)} paragraphs")

    return final_paragraphs


def _is_list_item_line(text: str) -> bool:
    text = (text or "").lstrip()
    return bool(re.match(r'^(\d+\.|[a-zA-Z]\)|[-+\u2022\u2023\u2043\u25e6\u2219])', text))


def _is_standalone_score_line(text: str) -> bool:
    normalized = _unaccent_upper(_clean_extracted_text(text)).strip(" .")
    normalized = re.sub(r"\s+", " ", normalized)
    return bool(re.fullmatch(r"\d+(?:[,.]\d+)?\s*(?:DIEM|D)", normalized))


def _looks_like_rubric_heading_after_score(text: str) -> bool:
    text = _clean_extracted_text(text)
    if not text or _is_list_item_line(text) or _is_standalone_score_line(text):
        return False
    if len(text.split()) < 2 or len(text) > 120 or ":" not in text:
        return False
    first_alpha = next((ch for ch in text if ch.isalpha()), "")
    return bool(first_alpha and first_alpha.isupper())


def _split_embedded_score_heading_line(line: str) -> List[str]:
    line = _clean_extracted_text(line)
    words = line.split()
    if len(words) < 3:
        return [line]
    max_score_tokens = min(4, len(words) - 1)
    for cut in range(1, max_score_tokens + 1):
        score = " ".join(words[:cut])
        tail = " ".join(words[cut:])
        if _is_standalone_score_line(score) and _looks_like_rubric_heading_after_score(tail):
            return [score, "- " + tail]
    return [line]


def _restore_missing_rubric_bullets(lines: List[str]) -> List[str]:
    has_list_context = any(_is_list_item_line(line) for line in lines)
    if not has_list_context:
        return lines

    expanded: List[str] = []
    for line in lines:
        if _is_list_item_line(line):
            expanded.append(line)
        else:
            expanded.extend(_split_embedded_score_heading_line(line))

    restored: List[str] = []
    for line in expanded:
        if (
            restored
            and _is_standalone_score_line(restored[-1])
            and _looks_like_rubric_heading_after_score(line)
        ):
            line = "- " + line
        restored.append(line)
    return restored


def _restore_inline_list_breaks(text: str) -> str:
    fixed_lines = []
    for line in (text or "").split("\n"):
        stripped = line.lstrip()
        if stripped.startswith("-") and " - " in stripped:
            indent = line[:len(line) - len(stripped)]
            line = indent + re.sub(r"\s+-\s+(?=\S)", "\n- ", stripped)
        fixed_lines.append(line)
    return "\n".join(fixed_lines)


def clean_ocr_cell_text(text: str) -> str:
    """
    Clean OCR text:
    1. Remove leading pipe characters.
    2. Smart merge lines (auto-heal wrapped text).
    """
    text = text.strip()
    # Remove leading pipes from start of text (legacy check)
    while text.startswith('|'):
        text = text[1:].strip()
        
    # Split by newline to handle internal lines
    if '\n' in text:
        lines = [l.strip() for l in text.split('\n')]
        # Filter empty lines? No, keep structure but clean pipes
        cleaned_lines = []
        for l in lines:
            l = l.strip()
            while l.startswith('|'):
                l = l[1:].strip()
            if l:
                cleaned_lines.append(l)
        lines = cleaned_lines
    else:
        # Single line case
        lines = [text]
        
    if not lines:
        return ""

    lines = _restore_missing_rubric_bullets(lines)

    # Smart merge
    merged = [lines[0]]
    
    for line in lines[1:]:
        prev = merged[-1]
        should_merge = False
        
        # Standalone score lines are complete rubric atoms; they should not
        # swallow the next criterion when OCR misses a bullet marker.
        if _is_standalone_score_line(prev):
            should_merge = False

        # Rule 1: Starts with lowercase -> Merge
        elif line and line[0].islower():
            should_merge = True
        
        # Rule 2: Prev line indicates continuation (no punctuation)
        # AND line doesn't look like a list item
        elif prev and prev[-1] not in '.!?;:':
            is_list_item = _is_list_item_line(line)
            if not is_list_item:
                should_merge = True
        
        if should_merge:
            merged[-1] = prev + " " + line
        else:
            merged.append(line)
    
    text = "\n".join(merged)
    text = _restore_inline_list_breaks(text)
        
    return text


# ============================================================================
# CAMELOT TABLE DETECTION
# ============================================================================




def sort_lines_reading_order(lines: List[TextLine]) -> List[TextLine]:
    """
    Sort lines by reading order: Top-to-Bottom, Left-to-Right.
    Handles lines that are roughly on the same vertical level (row).
    """
    if not lines:
        return []

    # 1. Sort by Y top first to allow linear processing
    lines_sorted = sorted(lines, key=lambda l: l.y)
    
    rows = []
    current_row = []
    current_row_bottom = -1.0
    
    for line in lines_sorted:
        if not current_row:
            current_row.append(line)
            current_row_bottom = line.y + line.height
            continue
            
        # Check if line belongs to current row
        # Criteria: The line's top is significantly above the current row's bottom.
        # Implies vertical overlap.
        # Use a generous overlap check logic.
        
        # If line starts above the bottom of the previous cluster (with tolerance)
        # Tolerance: 20% of line height?
        tolerance = line.height * 0.2
        if line.y < current_row_bottom - tolerance:
            current_row.append(line)
            # Update bottom to the max of the cluster
            current_row_bottom = max(current_row_bottom, line.y + line.height)
        else:
            # New row
            rows.append(current_row)
            current_row = [line]
            current_row_bottom = line.y + line.height
            
    if current_row:
        rows.append(current_row)
        
    # 2. Sort each row by X and flatten
    final_lines = []
    for row in rows:
        row.sort(key=lambda l: l.x)
        final_lines.extend(row)
        
    return final_lines


def _quantile(values: List[float], q: float, default: float = 0.0) -> float:
    if not values:
        return default
    ordered = sorted(float(v) for v in values)
    if len(ordered) == 1:
        return ordered[0]
    q = max(0.0, min(1.0, q))
    pos = (len(ordered) - 1) * q
    lo = int(pos)
    hi = min(lo + 1, len(ordered) - 1)
    frac = pos - lo
    return ordered[lo] * (1.0 - frac) + ordered[hi] * frac


def _cluster_objects_by_x(objects: List[object], k: int, x_getter) -> List[List[object]]:
    if k <= 1 or len(objects) < k:
        return []

    xs = [float(x_getter(obj)) for obj in objects]
    centers = [_quantile(xs, (idx + 0.5) / k, xs[0]) for idx in range(k)]
    clusters: List[List[object]] = [[] for _ in range(k)]
    for _ in range(18):
        clusters = [[] for _ in range(k)]
        for obj in objects:
            x = float(x_getter(obj))
            best = min(range(k), key=lambda idx: abs(x - centers[idx]))
            clusters[best].append(obj)
        if any(not cluster for cluster in clusters):
            return []
        new_centers = [_median([float(x_getter(obj)) for obj in cluster], centers[idx]) for idx, cluster in enumerate(clusters)]
        if max(abs(new_centers[idx] - centers[idx]) for idx in range(k)) < 0.5:
            centers = new_centers
            break
        centers = new_centers

    return [cluster for _center, cluster in sorted(zip(centers, clusters), key=lambda pair: pair[0])]


def _multi_column_candidate(line: TextLine, page_w: float, page_h: float) -> bool:
    text = _clean_extracted_text(line.text)
    if len(text) < 3:
        return False
    if line.y < page_h * 0.07 or line.y > page_h * 0.94:
        return False
    if line.width <= 0 or line.height <= 0:
        return False
    if line.width > page_w * 0.52:
        return False
    if line.height > page_h * 0.10:
        return False
    semantic_type = (getattr(line, "semantic_type", "") or "").strip().lower().replace("-", " ")
    layout_type = (getattr(line, "_layout_region_type", "") or "").strip().lower().replace("-", " ")
    if semantic_type in {"table", "figure", "isolate formula", "formula"}:
        return False
    if layout_type in {"table", "figure", "isolate formula", "formula"}:
        return False
    return True


def _build_multi_column_profile(candidates: List[TextLine], page_w: float, page_h: float, column_count: int) -> Optional[dict]:
    min_total = 18 if column_count == 3 else 12
    if len(candidates) < min_total:
        return None

    clusters = _cluster_objects_by_x(candidates, column_count, lambda line: line.x)
    if len(clusters) != column_count:
        return None

    min_cluster = max(4 if column_count == 3 else 5, int(len(candidates) * (0.12 if column_count == 3 else 0.16)))
    if min(len(cluster) for cluster in clusters) < min_cluster:
        return None

    starts = [_median([line.x for line in cluster], 0.0) for cluster in clusters]
    gaps = [starts[idx + 1] - starts[idx] for idx in range(column_count - 1)]
    min_gap = page_w * (0.13 if column_count == 3 else 0.22)
    if min(gaps) < min_gap:
        return None
    min_span = page_w * (0.42 if column_count == 3 else 0.24)
    if starts[-1] - starts[0] < min_span:
        return None

    start_tolerance = page_w * (0.058 if column_count == 3 else 0.08)
    core_ratio = 0.45 if column_count == 3 else 0.55
    cores: List[List[TextLine]] = []
    for idx, cluster in enumerate(clusters):
        core = [line for line in cluster if abs(line.x - starts[idx]) <= start_tolerance]
        if len(core) < max(3, int(len(cluster) * core_ratio)):
            return None
        cores.append(core)

    ends = [_median([line.x + line.width for line in cluster], starts[idx]) for idx, cluster in enumerate(clusters)]
    gutters: List[float] = []
    for idx in range(column_count - 1):
        visual_gap = starts[idx + 1] - ends[idx]
        if visual_gap < max(8.0, page_w * 0.018):
            return None
        gutter = (ends[idx] + starts[idx + 1]) / 2.0
        if not (starts[idx] + page_w * 0.06 <= gutter <= starts[idx + 1] - page_w * 0.015):
            return None
        gutters.append(gutter)
    if column_count == 2 and not (page_w * 0.34 <= gutters[0] <= page_w * 0.66):
        return None

    min_ys = [min(line.y for line in core) for core in cores]
    max_ys = [max(line.y for line in core) for core in cores]
    balance = min(len(cluster) for cluster in clusters) / max(1, max(len(cluster) for cluster in clusters))
    columns = [
        {
            "index": idx,
            "start": starts[idx],
            "end": ends[idx],
            "count": len(clusters[idx]),
            "core_count": len(cores[idx]),
        }
        for idx in range(column_count)
    ]
    return {
        "column_count": column_count,
        "columns": columns,
        "gutters": gutters,
        "start_y": min(min_ys),
        "balanced_start_y": max(min_ys),
        "end_y": max(max_ys),
        "start_tolerance": start_tolerance,
        "balance": balance,
        "score": column_count * 100.0 + balance * 20.0 + (min(gaps) / max(page_w, 1.0)) * 100.0,
    }


def _detect_multi_column_profile(page_lines: List[TextLine], page_w: float, page_h: float) -> Optional[dict]:
    candidates = [line for line in page_lines if _multi_column_candidate(line, page_w, page_h)]
    if len(candidates) < 12:
        return None

    profiles = [
        profile
        for profile in (
            _build_multi_column_profile(candidates, page_w, page_h, 3),
            _build_multi_column_profile(candidates, page_w, page_h, 2),
        )
        if profile is not None
    ]
    if not profiles:
        return None
    return sorted(profiles, key=lambda profile: profile["score"], reverse=True)[0]


def _detect_two_column_profile(page_lines: List[TextLine], page_w: float, page_h: float) -> Optional[dict]:
    profile = _build_multi_column_profile(
        [line for line in page_lines if _multi_column_candidate(line, page_w, page_h)],
        page_w,
        page_h,
        2,
    )
    if not profile:
        return None
    return {
        **profile,
        "gutter": profile["gutters"][0],
        "left_start": profile["columns"][0]["start"],
        "right_start": profile["columns"][1]["start"],
        "left_count": profile["columns"][0]["count"],
        "right_count": profile["columns"][1]["count"],
    }


def _column_index_for_x(center_x: float, gutters: List[float]) -> int:
    for idx, gutter in enumerate(gutters):
        if center_x < gutter:
            return idx
    return len(gutters)


def _multi_column_role_for_index(idx: int, column_count: int) -> str:
    if column_count == 2:
        return "left" if idx == 0 else "right"
    return f"col{idx + 1}"


def apply_multi_column_reading_order(
    lines: List[TextLine],
    page_info: dict,
    logger: Logger,
) -> List[TextLine]:
    """
    Reorder scientific/newsletter pages column-major when independent text
    columns are clearly present. This is intentionally branch-agnostic: digital
    text extraction and scan OCR both emit positioned TextLine objects, so the
    same geometric reading-order rule applies to both.
    """
    if not lines:
        return lines

    lines_by_page: Dict[int, List[TextLine]] = {}
    for line in lines:
        lines_by_page.setdefault(int(line.page), []).append(line)

    profiles_by_page: Dict[int, dict] = {}
    repeated_edge_keys: Dict[str, int] = {}
    for page, page_lines in lines_by_page.items():
        info = page_info.get(page, {}) if page_info else {}
        page_w = float(info.get("width", 595.0) or 595.0)
        page_h = float(info.get("height", 842.0) or 842.0)
        profile = _detect_multi_column_profile(page_lines, page_w, page_h)
        if not profile:
            continue
        profiles_by_page[page] = profile
        for line in page_lines:
            if line.y > page_h * 0.09 and line.y + line.height < page_h * 0.93:
                continue
            key = _norm_table_key(re.sub(r"\d", "#", _clean_extracted_text(line.text)))
            if 3 <= len(key) <= 120:
                repeated_edge_keys[key] = repeated_edge_keys.get(key, 0) + 1

    def is_running_edge_line(line: TextLine, page_w: float, page_h: float) -> bool:
        if line.y > page_h * 0.09 and line.y + line.height < page_h * 0.93:
            return False
        text = _clean_extracted_text(line.text)
        key = _norm_table_key(re.sub(r"\d", "#", text))
        if repeated_edge_keys.get(key, 0) >= 2:
            return True
        compact = _norm_table_key(text)
        return (
            ("QXP" in compact and "PAGE" in compact)
            or ("CANCERCONTROL" in compact and bool(re.search(r"\b20\d{2}\b", text)))
            or bool(re.search(r"\bPAGE\s+\d+\b", text, re.IGNORECASE))
        )

    reordered_pages = []
    removed_running = 0
    result: List[TextLine] = []
    for page in sorted(lines_by_page):
        page_lines = lines_by_page[page]
        info = page_info.get(page, {}) if page_info else {}
        page_w = float(info.get("width", 595.0) or 595.0)
        page_h = float(info.get("height", 842.0) or 842.0)
        profile = profiles_by_page.get(page)
        if not profile:
            result.extend(sorted(page_lines, key=lambda line: (line.order, line.y, line.x)))
            continue

        column_count = int(profile.get("column_count", 2) or 2)
        columns = list(profile.get("columns") or [])
        gutters = [float(gutter) for gutter in (profile.get("gutters") or [])]
        start_y = float(profile["start_y"])
        balanced_start_y = float(profile["balanced_start_y"])
        end_y = float(profile["end_y"])
        start_tolerance = float(profile["start_tolerance"])
        y_tol = max(_median([line.height for line in page_lines if line.height > 0], 10.0), 6.0)
        full_width_threshold = page_w * (0.58 if column_count == 2 else 0.46)

        top: List[TextLine] = []
        column_lines: List[List[TextLine]] = [[] for _ in range(column_count)]
        bottom: List[TextLine] = []
        other_body: List[TextLine] = []

        for line in page_lines:
            if is_running_edge_line(line, page_w, page_h):
                removed_running += 1
                continue
            setattr(line, "_multi_column_page", True)
            setattr(line, "_multi_column_count", column_count)
            if column_count == 2:
                setattr(line, "_digital_two_column_page", True)
            center_x = line.x + line.width / 2.0
            col_idx = _column_index_for_x(center_x, gutters)
            col_idx = min(max(col_idx, 0), max(0, column_count - 1))
            column = columns[col_idx] if col_idx < len(columns) else {"start": line.x}
            is_full_width = (
                line.width > full_width_threshold
                or (
                    columns
                    and line.x <= float(columns[0]["start"]) + start_tolerance
                    and line.x + line.width >= float(columns[-1]["start"]) - start_tolerance
                )
            )
            near_column_start = abs(line.x - float(column.get("start", line.x))) <= start_tolerance
            before_balanced_columns = line.y < balanced_start_y - y_tol
            if line.y < start_y - y_tol or (is_full_width and line.y < balanced_start_y - y_tol):
                setattr(line, "_multi_column_role", "top")
                setattr(line, "_multi_column_index", -1)
                if column_count == 2:
                    setattr(line, "_digital_column", "top")
                top.append(line)
            elif line.y > max(end_y + y_tol, page_h * 0.94):
                setattr(line, "_multi_column_role", "bottom")
                setattr(line, "_multi_column_index", -1)
                if column_count == 2:
                    setattr(line, "_digital_column", "bottom")
                bottom.append(line)
            elif before_balanced_columns and line.width > page_w * (0.45 if column_count == 2 else 0.36):
                setattr(line, "_multi_column_role", "top")
                setattr(line, "_multi_column_index", -1)
                if column_count == 2:
                    setattr(line, "_digital_column", "top")
                top.append(line)
            elif (
                before_balanced_columns
                and columns
                and line.x > float(columns[0]["start"]) + page_w * 0.04
                and line.width < page_w * (0.24 if column_count == 2 else 0.19)
            ):
                setattr(line, "_multi_column_role", "top")
                setattr(line, "_multi_column_index", -1)
                if column_count == 2:
                    setattr(line, "_digital_column", "top")
                top.append(line)
            elif not is_full_width and (near_column_start or line.y >= balanced_start_y - y_tol):
                role = _multi_column_role_for_index(col_idx, column_count)
                setattr(line, "_multi_column_role", role)
                setattr(line, "_multi_column_index", col_idx)
                if column_count == 2:
                    setattr(line, "_digital_column", role)
                column_lines[col_idx].append(line)
            elif line.y < balanced_start_y - y_tol:
                setattr(line, "_multi_column_role", "top")
                setattr(line, "_multi_column_index", -1)
                if column_count == 2:
                    setattr(line, "_digital_column", "top")
                top.append(line)
            else:
                setattr(line, "_multi_column_role", "body")
                setattr(line, "_multi_column_index", -1)
                if column_count == 2:
                    setattr(line, "_digital_column", "body")
                other_body.append(line)

        ordered = sorted(top, key=lambda line: (line.y, line.x))
        for group in column_lines:
            ordered.extend(sorted(group, key=lambda line: (line.y, line.x)))
        ordered.extend(sorted(other_body, key=lambda line: (line.y, line.x)))
        ordered.extend(sorted(bottom, key=lambda line: (line.y, line.x)))
        base_order = page * 100000
        for offset, line in enumerate(ordered):
            line.order = base_order + offset
        result.extend(ordered)
        counts = ",".join(f"C{idx + 1}={int(col.get('count', 0))}" for idx, col in enumerate(columns))
        gutter_text = "/".join(f"{gutter:.1f}" for gutter in gutters)
        reordered_pages.append(f"{page}({column_count}col,{counts},gutter={gutter_text})")

    if reordered_pages:
        logger.log("Applied multi-column reading order on page(s): " + ", ".join(reordered_pages))
    if removed_running:
        logger.log(f"Removed {removed_running} repeated article running header/footer line(s)")
    return sorted(result, key=lambda line: (line.page, line.order, line.y, line.x))


def apply_digital_two_column_reading_order(
    lines: List[TextLine],
    page_info: dict,
    logger: Logger,
) -> List[TextLine]:
    return apply_multi_column_reading_order(lines, page_info, logger)


def _word_inline_sort_key(word: dict) -> Tuple[float, float]:
    x0, y0, _x1, _y1 = _word_bbox_for_assignment(word)
    return (x0, y0)


def split_lines_crossing_multi_column_gutters(
    lines: List[TextLine],
    page_info: dict,
    logger: Logger,
) -> List[TextLine]:
    """
    Split OCR lines that accidentally join adjacent columns.

    Scientific/newsletter scans often have multiple text baselines at the same
    y. OCR can merge them into one long line. Reading-order sorting cannot
    repair that unless the line is first split at detected column gutters.
    """
    if not lines:
        return lines

    lines_by_page: Dict[int, List[TextLine]] = {}
    for line in lines:
        lines_by_page.setdefault(int(line.page), []).append(line)

    split_count = 0
    result: List[TextLine] = []
    for page in sorted(lines_by_page):
        page_lines = lines_by_page[page]
        info = page_info.get(page, {}) if page_info else {}
        page_w = float(info.get("width", 595.0) or 595.0)
        page_h = float(info.get("height", 842.0) or 842.0)
        profile = _detect_multi_column_profile(page_lines, page_w, page_h)
        if not profile:
            result.extend(page_lines)
            continue

        column_count = int(profile.get("column_count", 2) or 2)
        gutters = [float(gutter) for gutter in (profile.get("gutters") or [])]
        min_line_width = page_w * (0.46 if column_count == 2 else 0.33)
        for line in page_lines:
            words = [
                word for word in (getattr(line, "word_items", []) or [])
                if _clean_extracted_text(str(word.get("text") or ""))
            ]
            if (
                len(words) < 4
                or line.width < min_line_width
                or not any(line.x < gutter < line.x + line.width for gutter in gutters)
            ):
                result.append(line)
                continue

            grouped_words: List[List[dict]] = [[] for _ in range(column_count)]
            crossing_word = False
            for word in words:
                wx0, wy0, wx1, wy1 = _word_bbox_for_assignment(word)
                if any(wx0 < gutter < wx1 for gutter in gutters):
                    crossing_word = True
                    break
                idx = _column_index_for_x((wx0 + wx1) / 2.0, gutters)
                idx = min(max(idx, 0), max(0, column_count - 1))
                grouped_words[idx].append(word)

            nonempty = [(idx, group) for idx, group in enumerate(grouped_words) if group]
            if crossing_word or len(nonempty) < 2:
                result.append(line)
                continue

            min_gap = max(3.0, line.height * 0.20)
            has_clear_gap = True
            for (_left_idx, left_group), (_right_idx, right_group) in zip(nonempty, nonempty[1:]):
                left_end = max(_word_bbox_for_assignment(word)[2] for word in left_group)
                right_start = min(_word_bbox_for_assignment(word)[0] for word in right_group)
                if right_start - left_end < min_gap:
                    has_clear_gap = False
                    break
            if not has_clear_gap:
                result.append(line)
                continue

            fragments: List[TextLine] = []
            for col_idx, group in nonempty:
                fragment = _line_fragment_from_words(line, sorted(group, key=_word_inline_sort_key))
                if fragment is None:
                    fragments = []
                    break
                setattr(fragment, "_column_split_fragment", True)
                setattr(fragment, "_multi_column_page", True)
                setattr(fragment, "_multi_column_count", column_count)
                setattr(fragment, "_multi_column_index", col_idx)
                role = _multi_column_role_for_index(col_idx, column_count)
                setattr(fragment, "_multi_column_role", role)
                if column_count == 2:
                    setattr(fragment, "_digital_two_column_page", True)
                    setattr(fragment, "_digital_column", role)
                fragments.append(fragment)
            if not fragments:
                result.append(line)
                continue

            result.extend(fragments)
            split_count += 1

    if split_count:
        logger.log(f"Split {split_count} OCR line(s) crossing detected multi-column gutters")
    return result


def split_lines_crossing_two_column_gutters(
    lines: List[TextLine],
    page_info: dict,
    logger: Logger,
) -> List[TextLine]:
    return split_lines_crossing_multi_column_gutters(lines, page_info, logger)


def split_digital_dual_header_spans(lines: List[TextLine], page_info: dict, logger: Logger) -> List[TextLine]:
    """
    Split digital-PDF header lines that PyMuPDF emits as one visual line even
    though the underlying spans belong to left/right administrative headers.
    """
    if not lines:
        return lines

    split_count = 0
    result: List[TextLine] = []
    for line in lines:
        info = (page_info or {}).get(int(line.page), {})
        page_w = float(info.get("width", 595.0) or 595.0)
        page_h = float(info.get("height", 842.0) or 842.0)
        if line.y > page_h * 0.18 or line.width < page_w * 0.48:
            result.append(line)
            continue
        spans = [
            span for span in (line.spans or [])
            if _clean_extracted_text(span.text) and float(getattr(span, "width", 0.0) or 0.0) > 0
        ]
        if len(spans) < 2:
            result.append(line)
            continue

        page_center = page_w / 2.0
        left_spans: List[TextSpan] = []
        right_spans: List[TextSpan] = []
        for span in spans:
            center_x = float(span.x) + float(span.width) / 2.0
            if center_x < page_center:
                left_spans.append(span)
            else:
                right_spans.append(span)
        if not left_spans or not right_spans:
            result.append(line)
            continue

        left_end = max(float(span.x) + float(span.width) for span in left_spans)
        right_start = min(float(span.x) for span in right_spans)
        left_text = _clean_extracted_text("".join(span.text for span in left_spans))
        right_text = _clean_extracted_text("".join(span.text for span in right_spans))
        header_pair = _looks_like_header_emphasis_line(left_text) and _looks_like_header_emphasis_line(right_text)
        if right_start - left_end < max(18.0, page_w * 0.045) and not header_pair:
            result.append(line)
            continue

        left_fragment = _line_fragment_from_spans(line, left_spans)
        right_fragment = _line_fragment_from_spans(line, right_spans)
        if left_fragment is None or right_fragment is None:
            result.append(line)
            continue
        setattr(left_fragment, "_span_split_fragment", True)
        setattr(right_fragment, "_span_split_fragment", True)
        result.extend([left_fragment, right_fragment])
        split_count += 1

    if split_count:
        logger.log(f"Split {split_count} digital dual-header line(s) by span geometry")
    return result


def filter_publication_running_edge_lines(lines: List[TextLine], page_info: dict, logger: Logger) -> List[TextLine]:
    """
    Remove journal/newsletter running headers and footers from the body flow.

    The signal is deliberately geometric plus publication-specific: text must
    sit at the page edge and the document must expose cues such as QXP source
    strings, "Page N", date/time stamps, or repeated digit-masked journal
    mastheads. Administrative letterheads are not removed by repetition alone.
    """
    if not lines:
        return lines

    def is_edge(line: TextLine) -> bool:
        info = (page_info or {}).get(int(line.page), {})
        page_h = float(info.get("height", 842.0) or 842.0)
        return line.y <= page_h * 0.095 or line.y + line.height >= page_h * 0.925

    edge_lines = [line for line in lines if is_edge(line)]
    if not edge_lines:
        return lines

    masked_counts: Dict[str, int] = {}
    cue_count = 0
    for line in edge_lines:
        text = _clean_extracted_text(line.text)
        compact = _norm_table_key(text)
        masked = _norm_table_key(re.sub(r"\d", "#", text))
        if 3 <= len(masked) <= 140:
            masked_counts[masked] = masked_counts.get(masked, 0) + 1
        if (
            "QXP" in compact
            or bool(re.search(r"\bPAGE\s+\d+\b", text, re.IGNORECASE))
            or bool(re.search(r"\b\d{1,2}/\d{1,2}/20\d{2}\b", text))
        ):
            cue_count += 1

    if cue_count < 2:
        return lines

    def looks_like_repeated_publication_label(text: str, masked: str) -> bool:
        if masked_counts.get(masked, 0) < 2:
            return False
        stripped = _clean_extracted_text(text)
        if not stripped or len(stripped) > 70 or re.search(r"[.:;]", stripped):
            return False
        words = re.findall(r"[A-Za-zÀ-ỹ]+", stripped)
        if not (1 <= len(words) <= 6):
            return False
        letters = [ch for ch in stripped if ch.isalpha()]
        if len(letters) < 6:
            return False
        upper = sum(1 for ch in letters if ch.upper() == ch)
        return upper / max(1, len(letters)) >= 0.82

    def is_publication_noise(line: TextLine) -> bool:
        if not is_edge(line):
            return False
        text = _clean_extracted_text(line.text)
        compact = _norm_table_key(text)
        masked = _norm_table_key(re.sub(r"\d", "#", text))
        if bool(re.fullmatch(r"page\s+\d+", text.strip(), re.IGNORECASE)):
            return True
        if "QXP" in compact:
            return True
        if "CANCERCONTROL" in compact and bool(re.search(r"\b20\d{2}\b", text)):
            return True
        if masked_counts.get(masked, 0) >= 2 and bool(re.search(r"\b20\d{2}\b", text)):
            return True
        if looks_like_repeated_publication_label(text, masked):
            return True
        return False

    kept: List[TextLine] = []
    removed = 0
    for line in lines:
        if is_publication_noise(line):
            removed += 1
            continue
        kept.append(line)
    if removed:
        logger.log(f"Removed {removed} publication running header/footer line(s)")
    return kept


_LAYOUT_TEXT_TYPES = {
    "plain text",
    "text",
    "title",
    "figure_caption",
    "figure caption",
    "table_caption",
    "table caption",
    "table_footnote",
    "table footnote",
    "formula_caption",
    "formula caption",
    "footnote",
    "page-header",
    "page header",
    "page-footer",
    "page footer",
    "abandon",
}
_LAYOUT_OBJECT_TYPES = {"table", "figure", "isolate_formula", "formula"}


def _layout_region_kind(region: dict) -> str:
    return str(region.get("type") or "").strip().lower().replace("-", " ")


def _layout_region_bbox_pdf(region: dict) -> Optional[Tuple[float, float, float, float]]:
    bbox = region.get("bbox_pdf")
    if not bbox or len(bbox) < 4:
        return None
    try:
        x0, y0, x1, y1 = (float(v) for v in bbox[:4])
    except (TypeError, ValueError):
        return None
    if x1 <= x0 or y1 <= y0:
        return None
    return (x0, y0, x1, y1)


def _layout_region_items(regions: List[dict]) -> List[dict]:
    items = []
    for idx, region in enumerate(regions or []):
        bbox = _layout_region_bbox_pdf(region)
        if not bbox:
            continue
        kind = _layout_region_kind(region)
        x0, y0, x1, y1 = bbox
        items.append({
            "idx": idx,
            "region": region,
            "kind": kind,
            "bbox": bbox,
            "x0": x0,
            "y0": y0,
            "x1": x1,
            "y1": y1,
            "width": x1 - x0,
            "height": y1 - y0,
            "cx": (x0 + x1) / 2.0,
            "cy": (y0 + y1) / 2.0,
        })
    return items


def _layout_center_inside(bbox: Tuple[float, float, float, float],
                          region_bbox: Tuple[float, float, float, float],
                          pad: float = 2.0) -> bool:
    x0, y0, x1, y1 = bbox
    rx0, ry0, rx1, ry1 = region_bbox
    cx = (x0 + x1) / 2.0
    cy = (y0 + y1) / 2.0
    return rx0 - pad <= cx <= rx1 + pad and ry0 - pad <= cy <= ry1 + pad


def _layout_multi_column_profile_from_items(items: List[dict], page_w: float, page_h: float) -> Optional[dict]:
    candidates = [
        item for item in items
        if item["kind"] not in {"page header", "page footer"}
        and page_h * 0.04 <= item["y0"] <= page_h * 0.96
        and page_w * 0.08 <= item["width"] <= page_w * 0.58
        and item["height"] >= max(18.0, page_h * 0.018)
    ]
    if len(candidates) < 4:
        return None

    profiles = []
    for column_count in (3, 2):
        if len(candidates) < column_count * 2:
            continue
        clusters = _cluster_objects_by_x(candidates, column_count, lambda item: item["x0"])
        if len(clusters) != column_count:
            continue
        if min(len(cluster) for cluster in clusters) < 2:
            continue
        starts = [_median([item["x0"] for item in cluster], 0.0) for cluster in clusters]
        gaps = [starts[idx + 1] - starts[idx] for idx in range(column_count - 1)]
        if min(gaps) < page_w * (0.12 if column_count == 3 else 0.20):
            continue
        if starts[-1] - starts[0] < page_w * (0.38 if column_count == 3 else 0.22):
            continue
        ends = [_median([item["x1"] for item in cluster], starts[idx]) for idx, cluster in enumerate(clusters)]
        gutters = []
        valid = True
        for idx in range(column_count - 1):
            if starts[idx + 1] - ends[idx] < max(8.0, page_w * 0.016):
                valid = False
                break
            gutter = (ends[idx] + starts[idx + 1]) / 2.0
            if not (starts[idx] + page_w * 0.05 <= gutter <= starts[idx + 1] - page_w * 0.015):
                valid = False
                break
            gutters.append(gutter)
        if not valid:
            continue
        if column_count == 2 and not (page_w * 0.34 <= gutters[0] <= page_w * 0.66):
            continue
        balance = min(len(cluster) for cluster in clusters) / max(1, max(len(cluster) for cluster in clusters))
        profiles.append({
            "column_count": column_count,
            "columns": [
                {
                    "index": idx,
                    "start": starts[idx],
                    "end": ends[idx],
                    "count": len(clusters[idx]),
                }
                for idx in range(column_count)
            ],
            "gutters": gutters,
            "start_y": min(min(item["y0"] for item in cluster) for cluster in clusters),
            "balanced_start_y": max(min(item["y0"] for item in cluster) for cluster in clusters),
            "end_y": max(max(item["y1"] for item in cluster) for cluster in clusters),
            "balance": balance,
            "score": column_count * 100.0 + balance * 20.0 + (min(gaps) / max(page_w, 1.0)) * 100.0,
        })

    if not profiles:
        return None
    return sorted(profiles, key=lambda profile: profile["score"], reverse=True)[0]


def _layout_two_column_profile_from_items(items: List[dict], page_w: float, page_h: float) -> Optional[dict]:
    profile = _layout_multi_column_profile_from_items(items, page_w, page_h)
    if not profile or int(profile.get("column_count", 0) or 0) != 2:
        return None
    return {
        **profile,
        "gutter": profile["gutters"][0],
        "left_count": profile["columns"][0]["count"],
        "right_count": profile["columns"][1]["count"],
    }


def _column_major_items_inside_layout_row(row: List[dict], page_w: float) -> List[dict]:
    """Order a row of layout regions without reversing stacked regions inside a column.

    Layout detectors often emit large text boxes for newspaper/scientific
    columns. A tall left-column box can overlap several right-column boxes, so
    naive row-major sorting by (x, y) treats all boxes as one row and lets tiny
    x jitter reverse the right-column stack. Split the row into geometric
    x-clusters first, then sort top-to-bottom inside each cluster.
    """
    if len(row) < 3:
        return sorted(row, key=lambda i: (i["x0"], i["y0"]))

    ordered = sorted(row, key=lambda i: (i["x0"], i["y0"]))
    gaps = [
        (
            max(0.0, float(ordered[idx]["x0"]) - float(ordered[idx - 1]["x0"])),
            idx,
        )
        for idx in range(1, len(ordered))
    ]
    split_threshold = max(32.0, page_w * 0.075)
    split_indices = [idx for gap, idx in gaps if gap >= split_threshold]
    if not split_indices:
        return ordered

    clusters: List[List[dict]] = []
    start = 0
    for idx in split_indices:
        clusters.append(ordered[start:idx])
        start = idx
    clusters.append(ordered[start:])

    if len(clusters) > 4 or any(not cluster for cluster in clusters):
        return ordered

    result: List[dict] = []
    for cluster in clusters:
        result.extend(sorted(cluster, key=lambda i: (i["y0"], i["x0"])))
    return result


def _row_major_layout_items(items: List[dict], page_w: float, page_h: float) -> List[dict]:
    if not items:
        return []
    y_tol = max(page_h * 0.018, 12.0)
    rows: List[List[dict]] = []
    for item in sorted(items, key=lambda i: (i["y0"], i["x0"])):
        placed = False
        for row in rows:
            row_y0 = min(r["y0"] for r in row)
            row_y1 = max(r["y1"] for r in row)
            overlap = max(0.0, min(row_y1, item["y1"]) - max(row_y0, item["y0"]))
            min_h = max(min(item["height"], row_y1 - row_y0), 1.0)
            if overlap / min_h >= 0.25 or abs(item["y0"] - row_y0) <= y_tol:
                row.append(item)
                placed = True
                break
        if not placed:
            rows.append([item])
    ordered: List[dict] = []
    for row in sorted(rows, key=lambda r: min(i["y0"] for i in r)):
        ordered.extend(_column_major_items_inside_layout_row(row, page_w))
    return ordered


def _sort_layout_items_for_reading_order(items: List[dict], page_w: float, page_h: float) -> List[dict]:
    profile = _layout_multi_column_profile_from_items(items, page_w, page_h)
    if not profile:
        return _row_major_layout_items(items, page_w, page_h)

    column_count = int(profile.get("column_count", 2) or 2)
    gutters = [float(gutter) for gutter in (profile.get("gutters") or [])]
    start_y = float(profile["start_y"])
    balanced_start_y = float(profile["balanced_start_y"])
    end_y = float(profile["end_y"])
    y_tol = max(page_h * 0.015, 10.0)

    top: List[dict] = []
    column_items: List[List[dict]] = [[] for _ in range(column_count)]
    bottom: List[dict] = []
    other: List[dict] = []
    for item in items:
        full_width = item["width"] > page_w * (0.60 if column_count == 2 else 0.48)
        if item["y1"] < start_y - y_tol or (full_width and item["y0"] < balanced_start_y - y_tol):
            top.append(item)
        elif item["y0"] > end_y + y_tol and full_width:
            bottom.append(item)
        elif not full_width:
            col_idx = _column_index_for_x(float(item["cx"]), gutters)
            col_idx = min(max(col_idx, 0), max(0, column_count - 1))
            column_items[col_idx].append(item)
        else:
            other.append(item)

    ordered = sorted(top, key=lambda i: (i["y0"], i["x0"]))
    for group in column_items:
        ordered.extend(sorted(group, key=lambda i: (i["y0"], i["x0"])))
    ordered.extend(sorted(other, key=lambda i: (i["y0"], i["x0"])))
    ordered.extend(sorted(bottom, key=lambda i: (i["y0"], i["x0"])))
    return ordered


def prepare_layout_region_reading_orders(
    layout_regions_by_page: Dict[int, List[dict]],
    page_info: dict,
    logger: Optional[Logger] = None,
) -> Dict[int, List[dict]]:
    ordered_by_page: Dict[int, List[dict]] = {}
    applied_pages = []
    for page, regions in (layout_regions_by_page or {}).items():
        info = (page_info or {}).get(page, {})
        page_w = float(info.get("width", 595.0) or 595.0)
        page_h = float(info.get("height", 842.0) or 842.0)
        items = _layout_region_items(regions)
        if not items:
            continue
        ordered = _sort_layout_items_for_reading_order(items, page_w, page_h)
        for rank, item in enumerate(ordered):
            order_key = rank * 10000
            item["region"]["_layout_order_rank"] = rank
            item["region"]["_layout_order_key"] = order_key
        ordered_by_page[page] = ordered
        applied_pages.append(page)
    if logger and applied_pages:
        logger.log(
            "Prepared layout-region reading order for page(s): "
            + ", ".join(str(p) for p in sorted(applied_pages))
        )
    return ordered_by_page


def _layout_item_for_bbox(page: int,
                          bbox: Tuple[float, float, float, float],
                          layout_regions_by_page: Dict[int, List[dict]]) -> Optional[dict]:
    best_item = None
    best_score = 0.0
    for item in _layout_region_items((layout_regions_by_page or {}).get(page, [])):
        region_order = item["region"].get("_layout_order_key")
        if region_order is None:
            continue
        overlap = _bbox_overlap_ratio(bbox, item["bbox"])
        center = _layout_center_inside(bbox, item["bbox"], pad=3.0)
        score = overlap + (0.35 if center else 0.0)
        if score > best_score:
            best_score = score
            best_item = item
    return best_item if best_score >= 0.20 else None


def _layout_order_for_bbox(page: int,
                           bbox: Tuple[float, float, float, float],
                           layout_regions_by_page: Dict[int, List[dict]],
                           default_order: int) -> int:
    item = _layout_item_for_bbox(page, bbox, layout_regions_by_page)
    if item is None:
        return int(default_order)
    return int(item["region"].get("_layout_order_key", default_order))


def apply_layout_region_reading_order(
    lines: List[TextLine],
    layout_regions_by_page: Dict[int, List[dict]],
    page_info: dict,
    logger: Logger,
) -> List[TextLine]:
    """
    Use detected layout blocks as the primary reading-order units.

    This follows the common OCR pipeline principle used by Docstrum/XY-cut and
    modern layout detectors: first segment the page into regions, then order
    text inside each region. It prevents global Y-sorting from interleaving
    two-column articles, captions, figures, and table-adjacent text.
    """
    if not lines or not layout_regions_by_page:
        return lines

    ordered_items_by_page = prepare_layout_region_reading_orders(layout_regions_by_page, page_info)
    if not ordered_items_by_page:
        return lines

    lines_by_page: Dict[int, List[TextLine]] = {}
    for line in lines:
        lines_by_page.setdefault(int(line.page), []).append(line)

    applied_pages = []
    result: List[TextLine] = []
    for page in sorted(lines_by_page):
        page_lines = lines_by_page[page]
        ordered_items = ordered_items_by_page.get(page)
        if not ordered_items:
            result.extend(sorted(page_lines, key=lambda line: (line.order, line.y, line.x)))
            continue

        assigned: Dict[int, List[TextLine]] = {int(item["region"]["_layout_order_rank"]): [] for item in ordered_items}
        unassigned: List[TextLine] = []
        for line in page_lines:
            bbox = _line_bbox(line)
            best_item = None
            best_score = 0.0
            for item in ordered_items:
                overlap = _bbox_overlap_ratio(bbox, item["bbox"])
                center = _layout_center_inside(bbox, item["bbox"], pad=max(2.0, line.height * 0.25))
                score = overlap + (0.35 if center else 0.0)
                if score > best_score:
                    best_score = score
                    best_item = item
            if best_item is not None and best_score >= 0.20:
                rank = int(best_item["region"]["_layout_order_rank"])
                setattr(line, "_layout_region_order", rank)
                setattr(line, "_layout_region_type", best_item["kind"])
                assigned.setdefault(rank, []).append(line)
            else:
                unassigned.append(line)

        assigned_count = sum(len(v) for v in assigned.values())
        if assigned_count < max(4, int(len(page_lines) * 0.25)):
            result.extend(sorted(page_lines, key=lambda line: (line.order, line.y, line.x)))
            continue

        page_result: List[TextLine] = []
        for item in ordered_items:
            rank = int(item["region"]["_layout_order_rank"])
            group = assigned.get(rank, [])
            if not group:
                continue
            base_order = int(item["region"].get("_layout_order_key", rank * 10000))
            # Preserve the geometry/column reading order computed before this
            # step. Layout detectors sometimes merge a two-column article into
            # one semantic region; sorting that region by (y, x) would
            # re-interleave the columns and undo the column detector.
            group_sorted = sorted(group, key=lambda line: (line.order, line.y, line.x))
            for offset, line in enumerate(group_sorted):
                line.order = base_order + offset
            page_result.extend(group_sorted)

        if unassigned:
            y_tol = max(float(page_info.get(page, {}).get("height", 842.0) or 842.0) * 0.01, 8.0)
            positioned_items = sorted(
                ordered_items,
                key=lambda item: int(item["region"].get("_layout_order_rank", item["idx"]))
            )

            def fallback_unassigned_order(line: TextLine, offset: int) -> int:
                if not positioned_items:
                    return int(line.order)
                candidate_rank = -1
                for item in positioned_items:
                    rank = int(item["region"].get("_layout_order_rank", item["idx"]))
                    if line.y >= float(item["y0"]) - y_tol:
                        candidate_rank = rank
                    else:
                        break
                if candidate_rank < 0:
                    return -10000 + offset
                return candidate_rank * 10000 + 9000 + offset

            for offset, line in enumerate(sorted(unassigned, key=lambda line: (line.order, line.y, line.x))):
                line.order = fallback_unassigned_order(line, offset)
            page_result.extend(sorted(unassigned, key=lambda line: (line.order, line.y, line.x)))

        result.extend(page_result)
        applied_pages.append(f"{page}({assigned_count}/{len(page_lines)} lines)")

    if applied_pages:
        logger.log("Applied layout-region reading order on page(s): " + ", ".join(applied_pages))
    return sorted(result, key=lambda line: (line.page, line.order, line.y, line.x))


def _source_pdf_for_layout(pdf_path: str) -> str:
    return pdf_path


def analyze_layout_regions_for_pdf(pdf_path: str, page_info: dict, logger: Logger) -> Dict[int, List[dict]]:
    """Run DocLayout ONNX on demand when the canonical sidecar lacks regions."""
    source_pdf = _source_pdf_for_layout(pdf_path)
    if not source_pdf or not os.path.exists(source_pdf):
        return {}
    try:
        from PIL import Image
        from scanindex.core.kie.json_utils import decorate_layout_regions
        from scanindex.core.tables import layout_analyzer as la
    except Exception as e:
        logger.log(f"DocLayout on-demand analysis unavailable: {e}")
        return {}

    try:
        primary = la.get_analyzer() if la.is_available() else None
        auxiliary = (
            la.get_doclaynet_analyzer()
            if os.environ.get("OCRTOOL_DISABLE_DOCLAYNET_LAYOUT") != "1"
            and hasattr(la, "is_doclaynet_available")
            and la.is_doclaynet_available()
            else None
        )
    except Exception as e:
        logger.log(f"DocLayout on-demand init failed: {e}")
        return {}

    if not primary and not auxiliary:
        logger.log("DocLayout on-demand analysis skipped: ONNX model/runtime unavailable")
        return {}

    regions_by_page: Dict[int, List[dict]] = {}
    try:
        doc = fitz.open(source_pdf)
        dpi = 240
        mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        for page_idx, page in enumerate(doc):
            page_num = page_idx + 1
            pix = page.get_pixmap(matrix=mat, annots=True)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_w = float(page.rect.width)
            page_h = float(page.rect.height)
            scale_x = page_w / max(float(pix.width), 1.0)
            scale_y = page_h / max(float(pix.height), 1.0)

            primary_regions = []
            auxiliary_regions = []
            if primary:
                primary_regions = decorate_layout_regions(
                    primary.analyze_page(image),
                    page_index=page_idx,
                    scale_x=scale_x,
                    scale_y=scale_y,
                )
            if auxiliary:
                auxiliary_regions = decorate_layout_regions(
                    auxiliary.analyze_page(image),
                    page_index=page_idx,
                    scale_x=scale_x,
                    scale_y=scale_y,
                )
            if hasattr(la, "merge_auxiliary_layout_regions"):
                regions = la.merge_auxiliary_layout_regions(primary_regions, auxiliary_regions)
            else:
                regions = list(primary_regions or [])
            if regions:
                regions_by_page[page_num] = regions
        doc.close()
    except Exception as e:
        logger.log(f"DocLayout on-demand analysis failed: {e}")
        return {}

    if regions_by_page:
        total = sum(len(v) for v in regions_by_page.values())
        logger.log(f"DocLayout on-demand analysis found {total} regions on {len(regions_by_page)} pages")
    return regions_by_page


def save_layout_regions_to_companion(
    companion_path: Optional[Path],
    ocr_data: Optional[dict],
    layout_regions_by_page: Dict[int, List[dict]],
    logger: Logger,
) -> None:
    if not companion_path or not ocr_data or not layout_regions_by_page:
        return
    pages = ocr_data.get("pages") or []
    changed = 0
    for page_idx, page in enumerate(pages, 1):
        if page.get("layout_regions"):
            continue
        regions = layout_regions_by_page.get(page_idx)
        if not regions:
            continue
        page["layout_regions"] = regions
        changed += 1
    if not changed:
        return
    try:
        from scanindex.core.canonical_io import save_canonical

        saved_path = save_canonical(companion_path, ocr_data, profile="docx_export")
        logger.log(f"Updated canonical OCR sidecar with layout regions: {saved_path}")
    except Exception as e:
        logger.log(f"Could not update canonical OCR sidecar with layout regions: {e}")


def _sha256_file(path: str) -> Optional[str]:
    if not path or not os.path.exists(path):
        return None
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _ocr_lines_hash(pdf_lines: List[TextLine]) -> str:
    h = hashlib.sha256()
    for line in sorted(pdf_lines or [], key=lambda l: (l.page, l.order, l.y, l.x, l.text)):
        payload = [
            int(line.page),
            round(float(line.x), 2),
            round(float(line.y), 2),
            round(float(line.width), 2),
            round(float(line.height), 2),
            line.source_line_id or "",
            line.text or "",
        ]
        h.update(json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8"))
        h.update(b"\n")
    return h.hexdigest()


def _ocr_lines_hash_by_page(pdf_lines: List[TextLine]) -> Dict[str, str]:
    by_page: Dict[int, List[TextLine]] = {}
    for line in pdf_lines or []:
        by_page.setdefault(int(line.page), []).append(line)
    return {
        str(page): _ocr_lines_hash(lines)
        for page, lines in sorted(by_page.items())
    }


def _page_visual_hashes_from_companion(companion_data: Optional[dict]) -> Dict[str, str]:
    result: Dict[str, str] = {}
    for idx, page in enumerate((companion_data or {}).get("pages") or [], 1):
        if not isinstance(page, dict):
            continue
        visual = page.get("visual_source") if isinstance(page.get("visual_source"), dict) else {}
        sha = str(visual.get("sha256") or "").strip()
        if sha:
            result[str(idx)] = sha
    return result


def _continued_table_flow_enabled() -> bool:
    return os.environ.get("OCRTOOL_ENABLE_CONTINUED_TABLE_FLOW", "1").strip() != "0"


def _table_options_hash(layout_regions_by_page: Dict[int, List[dict]]) -> str:
    payload = {
        "cache_schema": "docx_table_cache_v7",
        "continued_table_flow": _continued_table_flow_enabled(),
        "layout_region_pages": sorted(int(page) for page in (layout_regions_by_page or {}).keys()),
        "engines": ["doclayout", "gmft_onnx", "docling_tableformer"],
    }
    raw = json.dumps(payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def _current_table_cache_meta(
    pdf_path: str,
    pdf_lines: List[TextLine],
    layout_regions_by_page: Dict[int, List[dict]],
    companion_data: Optional[dict],
) -> dict:
    ocr_pipeline = ((companion_data or {}).get("pipeline") or {}).get("ocr") or {}
    source_mode = ocr_pipeline.get("source_mode")
    if source_mode not in {"scan", "digital", "mixed"}:
        source_mode = "digital" if ocr_pipeline.get("engine") == "digital_pdf_text" else "scan"
    page_hashes = _page_visual_hashes_from_companion(companion_data)
    file_hash = _sha256_file(pdf_path)
    return {
        "status": "complete",
        "completed_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "source_mode": source_mode,
        "visual_source_sha256": file_hash,
        "page_visual_sha256": page_hashes or {"__file__": file_hash},
        "ocr_lines_hash": _ocr_lines_hash(pdf_lines),
        "page_ocr_lines_hash": _ocr_lines_hash_by_page(pdf_lines),
        "options_hash": _table_options_hash(layout_regions_by_page),
        "engines": ["doclayout", "gmft_onnx", "docling_tableformer"],
        "cache_layers": {
            "raw_page": "docx_raw_tables_v1",
            "composed_groups": "docx_continued_table_groups_v1",
            "render_structures": "docx_table_structures_v1",
        },
    }


def _table_cache_valid(companion_data: Optional[dict], expected: dict, logger: Logger) -> bool:
    if not companion_data:
        return False
    meta = (companion_data.get("pipeline") or {}).get("table_extraction") or {}
    if not meta:
        logger.log("Table cache missing; regenerating")
        return False

    page_hash_expected = expected.get("page_visual_sha256") or {}
    page_hash_actual = meta.get("page_visual_sha256") or {}
    if page_hash_expected and page_hash_actual != page_hash_expected:
        logger.log("Table cache invalid (page visual fingerprint mismatch); regenerating")
        return False

    page_text_expected = expected.get("page_ocr_lines_hash") or {}
    page_text_actual = meta.get("page_ocr_lines_hash") or {}
    if page_text_expected and page_text_actual != page_text_expected:
        logger.log("Table cache invalid (page OCR text hash mismatch); regenerating")
        return False

    for key in ("visual_source_sha256", "ocr_lines_hash", "options_hash"):
        if meta.get(key) != expected.get(key):
            logger.log(f"Table cache invalid ({key} mismatch); regenerating")
            return False
    if meta.get("status") != "complete":
        logger.log("Table cache invalid (status is not complete); regenerating")
        return False
    return True


def _bbox_record_to_tuple(value) -> Tuple[float, float, float, float]:
    try:
        vals = [float(v) for v in list(value or [])[:4]]
    except Exception:
        vals = []
    if len(vals) != 4:
        return (0.0, 0.0, 0.0, 0.0)
    return (vals[0], vals[1], vals[2], vals[3])


def _normalise_cell_grid(cells, row_count: int, col_count: int) -> List[List[str]]:
    rows = []
    raw_rows = cells if isinstance(cells, list) else []
    for r in range(row_count):
        raw = raw_rows[r] if r < len(raw_rows) and isinstance(raw_rows[r], list) else []
        row = [str(raw[c]) if c < len(raw) and raw[c] is not None else "" for c in range(col_count)]
        rows.append(row)
    return rows


def _normalise_bbox_grid(cell_bboxes, row_count: int, col_count: int) -> List[List[Tuple[float, float, float, float]]]:
    rows = []
    raw_rows = cell_bboxes if isinstance(cell_bboxes, list) else []
    for r in range(row_count):
        raw = raw_rows[r] if r < len(raw_rows) and isinstance(raw_rows[r], list) else []
        row = [_bbox_record_to_tuple(raw[c] if c < len(raw) else None) for c in range(col_count)]
        rows.append(row)
    return rows


def _table_region_to_cache(table: TableRegion, index: int) -> dict:
    row_count = int(getattr(table, "row_count", 0) or len(getattr(table, "cells", []) or []))
    col_count = int(getattr(table, "col_count", 0) or max((len(row) for row in (getattr(table, "cells", []) or [])), default=0))
    x0, y0, x1, y1 = _table_bbox(table)
    render_options = {
        "continued_table_flow": bool(getattr(table, "enable_blank_continuation_vmerge", False)),
    }
    if getattr(table, "disable_vertical_merge", False):
        render_options["disable_vertical_merge"] = True
    record = {
        "id": getattr(table, "id", None) or f"tbl_p{int(getattr(table, 'page', 1)) - 1}_{index}",
        "page_index": max(0, int(getattr(table, "page", 1)) - 1),
        "source_engine": getattr(table, "source", None) or ("pymupdf_native" if getattr(table, "native_table", False) else "unknown"),
        "bbox": [round(float(x0), 2), round(float(y0), 2), round(float(x1), 2), round(float(y1), 2)],
        "row_count": row_count,
        "col_count": col_count,
        "cells": _normalise_cell_grid(getattr(table, "cells", []), row_count, col_count),
        "cell_bboxes": _normalise_bbox_grid(getattr(table, "cell_bboxes", []), row_count, col_count),
        "source_pages": [int(p) for p in (getattr(table, "source_pages", None) or [getattr(table, "page", 1)])],
        "row_source_pages": [int(p) for p in (getattr(table, "row_source_pages", None) or [])],
        "skip_render": bool(getattr(table, "skip_render", False)),
        "render_options": render_options,
    }
    if getattr(table, "native_table", False):
        record["native_table"] = True
    if hasattr(table, "horizontal_text_spans"):
        record["horizontal_text_spans"] = getattr(table, "horizontal_text_spans")
    return record


def _table_region_from_cache(record: dict, logger: Logger) -> Optional[TableRegion]:
    if not isinstance(record, dict):
        return None
    try:
        row_count = int(record.get("row_count") or 0)
        col_count = int(record.get("col_count") or 0)
        page = int(record.get("page_index", 0)) + 1
        bbox = _bbox_record_to_tuple(record.get("bbox"))
    except Exception:
        return None
    if row_count <= 0 or col_count <= 0 or page <= 0:
        return None
    cells = _normalise_cell_grid(record.get("cells"), row_count, col_count)
    cell_bboxes = _normalise_bbox_grid(record.get("cell_bboxes"), row_count, col_count)
    if not any(any(any(bx) for bx in row) for row in cell_bboxes):
        logger.log(f"Table cache entry {record.get('id') or '?'} has no usable cell bboxes; regenerating")
        return None
    table = TableRegion(
        page=page,
        y_top=float(bbox[1]),
        y_bottom=float(bbox[3]),
        cells=cells,
        row_count=row_count,
        col_count=col_count,
        cell_bboxes=cell_bboxes,
    )
    setattr(table, "id", record.get("id") or "")
    setattr(table, "source", record.get("source_engine") or "cached")
    setattr(table, "x_left", float(bbox[0]))
    setattr(table, "x_right", float(bbox[2]))
    setattr(table, "source_pages", [int(p) for p in (record.get("source_pages") or [page])])
    setattr(table, "row_source_pages", [int(p) for p in (record.get("row_source_pages") or [])])
    setattr(table, "skip_render", bool(record.get("skip_render", False)))
    if record.get("native_table"):
        setattr(table, "native_table", True)
    render_options = record.get("render_options") if isinstance(record.get("render_options"), dict) else {}
    if render_options.get("continued_table_flow"):
        setattr(table, "enable_blank_continuation_vmerge", True)
    if render_options.get("disable_vertical_merge"):
        setattr(table, "disable_vertical_merge", True)
    if isinstance(record.get("horizontal_text_spans"), dict):
        spans = {}
        for key, value in record.get("horizontal_text_spans", {}).items():
            try:
                spans[int(key)] = value
            except Exception:
                continue
        setattr(table, "horizontal_text_spans", spans)
    return table


def load_table_structures_from_companion(
    companion_data: Optional[dict],
    expected_meta: dict,
    logger: Logger,
) -> Optional[List[TableRegion]]:
    if not _table_cache_valid(companion_data, expected_meta, logger):
        return None
    tables: List[TableRegion] = []
    saw_table_field = False
    for page in (companion_data or {}).get("pages") or []:
        if "table_structures" in page or "table_raw_cache" in page:
            saw_table_field = True
        records = page.get("table_structures")
        if records is None and "table_raw_cache" in page:
            records = page.get("table_raw_cache")
        for record in records or []:
            table = _table_region_from_cache(record, logger)
            if table is None:
                return None
            tables.append(table)
    if not tables:
        if saw_table_field:
            logger.log("Loaded empty table structure cache from canonical sidecar")
            return []
        logger.log("Table cache missing table_structures; regenerating")
        return None
    logger.log(f"Loaded {len(tables)} table structures from canonical cache")
    return sorted(tables, key=lambda t: (t.page, t.y_top, t.y_bottom))


def _continued_table_group_cache_records(table_regions: List[TableRegion], table_meta: dict) -> List[dict]:
    page_hashes = table_meta.get("page_visual_sha256") or {}
    text_hashes = table_meta.get("page_ocr_lines_hash") or {}
    records = []
    seen = set()
    for idx, table in enumerate(table_regions or []):
        source_pages = [int(p) for p in (getattr(table, "source_pages", None) or [getattr(table, "page", 1)])]
        source_pages = sorted({p for p in source_pages if p > 0})
        if len(source_pages) <= 1 and not getattr(table, "skip_render", False):
            continue
        key = tuple(source_pages) or (int(getattr(table, "page", 1)),)
        table_id = getattr(table, "id", None) or f"tbl_p{int(getattr(table, 'page', 1)) - 1}_{idx}"
        record_key = (key, table_id)
        if record_key in seen:
            continue
        seen.add(record_key)
        records.append({
            "id": f"group_{len(records)}",
            "table_id": table_id,
            "source_pages": list(key),
            "page_visual_sha256": {str(p): page_hashes.get(str(p)) for p in key if page_hashes.get(str(p))},
            "page_ocr_lines_hash": {str(p): text_hashes.get(str(p)) for p in key if text_hashes.get(str(p))},
            "compose_options_hash": table_meta.get("options_hash"),
            "status": "complete",
        })
    return records


def save_table_structures_to_companion(
    companion_path: Optional[Path],
    companion_data: Optional[dict],
    table_regions: List[TableRegion],
    table_meta: dict,
    logger: Logger,
) -> None:
    if not companion_path or not companion_data:
        return
    pages = companion_data.get("pages") or []
    by_page: Dict[int, List[dict]] = {}
    for idx, table in enumerate(table_regions or []):
        page_index = max(0, int(getattr(table, "page", 1)) - 1)
        by_page.setdefault(page_index, []).append(_table_region_to_cache(table, idx))
    for page_idx, page in enumerate(pages):
        if not isinstance(page, dict):
            continue
        records = by_page.get(page_idx, [])
        page["table_raw_cache"] = records
        page["table_structures"] = records
        page["table_extraction"] = {
            "status": "complete",
            "page_index": page_idx,
            "visual_source_sha256": (table_meta.get("page_visual_sha256") or {}).get(str(page_idx + 1)),
            "ocr_lines_hash": (table_meta.get("page_ocr_lines_hash") or {}).get(str(page_idx + 1)),
            "options_hash": table_meta.get("options_hash"),
        }
    pipeline = companion_data.setdefault("pipeline", {})
    pipeline["table_extraction"] = dict(table_meta)
    pipeline.setdefault("docx_export", {})["continued_table_groups"] = _continued_table_group_cache_records(
        table_regions,
        table_meta,
    )
    try:
        from scanindex.core.canonical_io import save_canonical

        saved_path = save_canonical(companion_path, companion_data, profile="docx_export")
        logger.log(f"Updated canonical OCR sidecar with table structures: {saved_path}")
    except Exception as e:
        logger.log(f"Could not update canonical OCR sidecar with table structures: {e}")


def get_lines_in_rect(rect: Tuple[float, float, float, float], lines: List[TextLine]) -> List[TextLine]:
    """
    Find all lines that fall within the rect (visual coords).
    Uses strict center-point inclusion for X axis to prevent duplicates across columns.
    rect: (x0, y0, x1, y1)
    """
    rx0, ry0, rx1, ry1 = rect
    found = []
    
    for line in lines:
        # Check intersection/containment
        # Broad logic: y center of line is within y range of rect
        ly = line.y + line.height/2
        
        # Relaxed Y check slightly to catch lines just on the edge
        # But cell top/bottoms are usually precise.
        if ry0 <= ly <= ry1:
             # Check x: STRICT center point inclusion
             # This ensures a line is assigned to exactly one column (unless columns overlap)
             lx_center = line.x + line.width / 2
             
             if rx0 <= lx_center <= rx1:
                 found.append(line)
    
    # Sort: Y then X using robust reading order (clustering)
    return sort_lines_reading_order(found)


def _norm_table_key(text: str) -> str:
    text = unicodedata.normalize("NFD", (text or "").upper())
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", "", text)


def _table_text(table: TableRegion) -> str:
    return " ".join(str(cell) for row in getattr(table, "cells", []) for cell in row)


def _looks_like_table_header_text(text: str) -> bool:
    compact = _norm_table_key(text)
    return (
        ("STT" in compact or "SOTTT" in compact)
        and (
            "NOIDUNG" in compact
            or "TEN" in compact
            or "TENCO" in compact
            or "CONGTAC" in compact
            or "SOLUONG" in compact
            or "GHICHU" in compact
        )
    )


def _row_looks_like_table_header(row: List[str]) -> bool:
    texts = [_clean_extracted_text(str(cell)) for cell in row if _clean_extracted_text(str(cell))]
    if len(texts) < 2:
        return False
    if any(len(text) > 90 for text in texts):
        return False
    avg_len = sum(len(text) for text in texts) / len(texts)
    if avg_len > 45:
        return False
    sentence_like = sum(1 for text in texts if re.search(r"[.;:]\s+\S", text))
    return sentence_like == 0


def _table_has_header(table: TableRegion) -> bool:
    head_rows = getattr(table, "cells", [])[:3]
    if _looks_like_table_header_text(" ".join(str(c) for row in head_rows for c in row)):
        return True
    return any(_row_looks_like_table_header(row) for row in head_rows[:2])


def _table_column_intervals(table: TableRegion) -> List[Tuple[float, float]]:
    cols = getattr(table, "col_count", 0) or 0
    if cols <= 0:
        return []

    best: List[Tuple[float, float]] = []
    for row_boxes in getattr(table, "cell_bboxes", []) or []:
        if len(row_boxes) < cols:
            continue
        intervals = []
        seen = set()
        ok = True
        for c in range(cols):
            bx = row_boxes[c]
            if not any(bx) or bx[2] <= bx[0]:
                ok = False
                break
            key = tuple(round(v, 1) for v in bx)
            if key in seen:
                ok = False
                break
            seen.add(key)
            intervals.append((float(bx[0]), float(bx[2])))
        if ok and len(intervals) > len(best):
            best = intervals
    return best


def _table_bbox(table: TableRegion) -> Tuple[float, float, float, float]:
    x_left = getattr(table, "x_left", None)
    x_right = getattr(table, "x_right", None)
    if x_left is None or x_right is None:
        xs = []
        for row in getattr(table, "cell_bboxes", []) or []:
            for bx in row:
                if any(bx):
                    xs.extend([bx[0], bx[2]])
        x_left = min(xs) if xs else 0.0
        x_right = max(xs) if xs else 0.0
    return (float(x_left), float(table.y_top), float(x_right), float(table.y_bottom))


def _bbox_intersection_ratio(a: Tuple[float, float, float, float],
                             b: Tuple[float, float, float, float]) -> float:
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0, iy0 = max(ax0, bx0), max(ay0, by0)
    ix1, iy1 = min(ax1, bx1), min(ay1, by1)
    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0
    area = max((ax1 - ax0) * (ay1 - ay0), 1e-6)
    return ((ix1 - ix0) * (iy1 - iy0)) / area


def _line_in_bbox(line: TextLine, bbox: Tuple[float, float, float, float]) -> bool:
    x0, y0, x1, y1 = bbox
    lx0, ly0, lx1, ly1 = _line_bbox_for_assignment(line)
    line_w = max(lx1 - lx0, 1e-6)
    line_h = max(ly1 - ly0, 1e-6)
    cx = line.x + line.width / 2
    cy = line.y + line.height / 2
    pad = max(line_h * 0.75, 2.0)
    if lx1 < x0 - pad or lx0 > x1 + pad or ly1 < y0 - pad or ly0 > y1 + pad:
        return False
    overlap_x = _axis_overlap_for_assignment(lx0, lx1, x0, x1)
    overlap_y = _axis_overlap_for_assignment(ly0, ly1, y0, y1)
    x_ok = x0 <= cx <= x1 or overlap_x / line_w >= 0.15 or min(abs(lx1 - x0), abs(lx0 - x1)) <= pad
    y_ok = y0 <= cy <= y1 or overlap_y / line_h >= 0.35
    return x_ok and y_ok


def _line_overlaps_table_area(line: TextLine, table: TableRegion) -> bool:
    """2D table membership check used when suppressing separately rendered table OCR."""
    x0, y0, x1, y1 = _table_bbox(table)
    lx0, ly0, lx1, ly1 = _line_bbox_for_assignment(line)
    line_w = max(lx1 - lx0, 1e-6)
    line_h = max(ly1 - ly0, 1e-6)
    cx = (lx0 + lx1) / 2.0
    cy = (ly0 + ly1) / 2.0
    pad_x = max(2.0, line_h * 0.35)
    pad_y = max(2.0, line_h * 0.45)

    if not (y0 - pad_y <= cy <= y1 + pad_y):
        return False
    overlap_x = _axis_overlap_for_assignment(lx0, lx1, x0 - pad_x, x1 + pad_x)
    overlap_y = _axis_overlap_for_assignment(ly0, ly1, y0 - pad_y, y1 + pad_y)
    if x0 - pad_x <= cx <= x1 + pad_x and overlap_y / line_h >= 0.35:
        return True
    return overlap_x / line_w >= 0.25 and overlap_y / line_h >= 0.45


def _cluster_y_positions(values: List[float], threshold: float = 42.0) -> List[float]:
    if not values:
        return []
    values = sorted(values)
    clusters = [[values[0]]]
    for value in values[1:]:
        if value - clusters[-1][-1] <= threshold:
            clusters[-1].append(value)
        else:
            clusters.append([value])
    return [min(cluster) for cluster in clusters]


def _is_table_row_key_line(line: TextLine, col_idx: int) -> bool:
    if col_idx > 1:
        return False
    key = _norm_table_key(line.text)
    if not key:
        return False
    if key in {"S", "T", "TT", "STT", "TEN", "TENCO", "TENCOQUAN", "QUAN"}:
        return False
    if bool(re.fullmatch(r"\d{1,3}", key)):
        return True

    words = [w for w in re.split(r"\s+", _clean_extracted_text(line.text)) if w]
    if not words:
        return False
    if len(words) > 8 or len(line.text) > 90:
        return False
    # Row labels normally live in the leftmost columns and are short noun
    # phrases, not sentence continuations. This is geometric/content-shape
    # based and does not depend on a specific agency name.
    return not line.text.strip().endswith((".", ";", ":"))


def _text_from_table_cell_lines(cell_lines: List[TextLine],
                                page: int,
                                x0: float,
                                x1: float,
                                logger: Logger) -> str:
    cell_lines = sort_lines_reading_order(cell_lines)
    if not cell_lines:
        return ""
    try:
        paras = merge_raw_paragraphs(cell_lines, {page: (x0, x1)}, logger)
        text = "\n".join(p[0] for p in paras if p[0])
    except Exception:
        text = "\n".join(l.text for l in cell_lines)
    return clean_ocr_cell_text(text)


def _line_bbox_for_assignment(line: TextLine) -> Tuple[float, float, float, float]:
    return (
        float(line.x),
        float(line.y),
        float(line.x + line.width),
        float(line.y + line.height),
    )


def _bbox_area_for_assignment(bbox: Tuple[float, float, float, float]) -> float:
    return max(0.0, bbox[2] - bbox[0]) * max(0.0, bbox[3] - bbox[1])


def _axis_overlap_for_assignment(a0: float, a1: float, b0: float, b1: float) -> float:
    return max(0.0, min(a1, b1) - max(a0, b0))


def _cell_line_assignment_score(
    line: TextLine,
    cell_bbox: Tuple[float, float, float, float],
    table_bbox: Tuple[float, float, float, float],
) -> Optional[float]:
    return _text_bbox_cell_assignment_score(_line_bbox_for_assignment(line), cell_bbox, table_bbox)


def _text_bbox_cell_assignment_score(
    text_bbox: Tuple[float, float, float, float],
    cell_bbox: Tuple[float, float, float, float],
    table_bbox: Tuple[float, float, float, float],
) -> Optional[float]:
    bx0, by0, bx1, by1 = cell_bbox
    if bx1 <= bx0 or by1 <= by0:
        return None

    lx0, ly0, lx1, ly1 = text_bbox
    line_w = max(lx1 - lx0, 1e-6)
    line_h = max(ly1 - ly0, 1e-6)
    cell_w = max(bx1 - bx0, 1e-6)
    cell_h = max(by1 - by0, 1e-6)
    lx_center = (lx0 + lx1) / 2.0
    ly_center = (ly0 + ly1) / 2.0

    table_x0, table_y0, table_x1, table_y1 = table_bbox
    table_pad = max(line_h * 0.65, 2.0)
    if (
        lx1 < table_x0 - table_pad
        or lx0 > table_x1 + table_pad
        or ly1 < table_y0 - table_pad
        or ly0 > table_y1 + table_pad
    ):
        return None
    if not (table_y0 <= ly_center <= table_y1):
        return None
    if not (table_x0 - table_pad <= lx_center <= table_x1 + table_pad):
        return None

    overlap_x = _axis_overlap_for_assignment(lx0, lx1, bx0, bx1)
    overlap_y = _axis_overlap_for_assignment(ly0, ly1, by0, by1)
    line_area = max(_bbox_area_for_assignment((lx0, ly0, lx1, ly1)), 1e-6)
    overlap_area = overlap_x * overlap_y
    line_cover = overlap_area / line_area
    x_cover = overlap_x / line_w
    y_cover = overlap_y / line_h

    center_inside = bx0 <= lx_center <= bx1 and by0 <= ly_center <= by1

    x_gap = max(bx0 - lx1, lx0 - bx1, 0.0)
    y_gap = max(by0 - ly1, ly0 - by1, 0.0)
    max_x_gap = max(line_h * 1.25, min(cell_w, table_x1 - table_x0) * 0.035)
    max_y_gap = max(line_h * 0.45, cell_h * 0.12)

    if center_inside:
        base = 4.0 + line_cover + x_cover + y_cover
    elif line_cover >= 0.18 and y_cover >= 0.35:
        base = 3.0 + 2.0 * line_cover + x_cover + y_cover
    elif overlap_x > 0.0 and y_cover >= 0.55:
        base = 2.0 + x_cover + y_cover
    elif overlap_y > 0.0 and x_cover >= 0.55:
        base = 1.8 + x_cover + y_cover
    elif y_cover >= 0.55 and x_gap <= max_x_gap:
        base = 1.2 + y_cover + max(0.0, 1.0 - x_gap / max(max_x_gap, 1e-6))
    elif x_cover >= 0.55 and y_gap <= max_y_gap:
        base = 1.0 + x_cover + max(0.0, 1.0 - y_gap / max(max_y_gap, 1e-6))
    else:
        return None

    cell_cx = (bx0 + bx1) / 2.0
    cell_cy = (by0 + by1) / 2.0
    norm_dist = (
        abs(lx_center - cell_cx) / max(cell_w, line_w, 1.0)
        + abs(ly_center - cell_cy) / max(cell_h, line_h, 1.0)
    )
    return base - norm_dist * 0.05


def _cell_bbox_key_for_assignment(bbox: Tuple[float, float, float, float]) -> Tuple[float, float, float, float]:
    return tuple(round(float(v), 2) for v in bbox[:4])


def _iter_table_cell_bboxes(table: TableRegion):
    rows = int(getattr(table, "row_count", 0) or 0)
    cols = int(getattr(table, "col_count", 0) or 0)
    for r in range(rows):
        for c in range(cols):
            if r >= len(table.cell_bboxes) or c >= len(table.cell_bboxes[r]):
                continue
            bbox = tuple(float(v) for v in table.cell_bboxes[r][c][:4])
            if any(bbox) and bbox[2] > bbox[0] and bbox[3] > bbox[1]:
                yield r, c, bbox


def _best_cell_for_text_bbox(
    text_bbox: Tuple[float, float, float, float],
    page: int,
    table_regions: List[TableRegion],
    table_bboxes: Dict[int, Tuple[float, float, float, float]],
) -> Optional[Tuple[int, int, int]]:
    best_key = None
    best_score = None
    for t_idx, table in enumerate(table_regions):
        if getattr(table, "skip_render", False) or table.page != page:
            continue
        table_bbox = table_bboxes.get(t_idx)
        if table_bbox is None:
            continue
        for r, c, bbox in _iter_table_cell_bboxes(table):
            score = _text_bbox_cell_assignment_score(text_bbox, bbox, table_bbox)
            if score is None:
                continue
            if best_score is None or score > best_score:
                best_score = score
                best_key = (t_idx, r, c)
    return best_key


def _word_bbox_for_assignment(word: dict) -> Tuple[float, float, float, float]:
    x = float(word.get("x", 0.0) or 0.0)
    y = float(word.get("y", 0.0) or 0.0)
    w = float(word.get("w", 0.0) or 0.0)
    h = float(word.get("h", 0.0) or 0.0)
    return (x, y, x + w, y + h)


def _text_from_word_items(words: List[dict]) -> str:
    parts = [_clean_extracted_text(str(word.get("text") or "")) for word in words]
    return " ".join(part for part in parts if part)


def _line_fragment_from_words(parent: TextLine, words: List[dict]) -> Optional[TextLine]:
    if not words:
        return None
    text = _text_from_word_items(words)
    if not text:
        return None
    boxes = [_word_bbox_for_assignment(word) for word in words]
    x0 = min(b[0] for b in boxes)
    y0 = min(b[1] for b in boxes)
    x1 = max(b[2] for b in boxes)
    y1 = max(b[3] for b in boxes)
    fragment = TextLine(
        text=text,
        x=x0,
        y=y0,
        width=max(0.0, x1 - x0),
        height=max(0.0, y1 - y0),
        page=parent.page,
        font_size=parent.font_size,
        spans=[TextSpan(text=text, font_size=parent.font_size, y=y0)],
        block_id=parent.block_id,
        paragraph_id=parent.paragraph_id,
        content_type=parent.content_type,
        fg_gray=parent.fg_gray,
        confidence=parent.confidence,
        semantic_type=parent.semantic_type,
        order=parent.order,
        source_line_id=parent.source_line_id,
        kie_labels=set(getattr(parent, "kie_labels", set())),
    )
    setattr(fragment, "word_items", list(words))
    return fragment


def _first_alpha_char(text: str) -> str:
    for ch in _clean_extracted_text(text or ""):
        if ch.isalpha():
            return ch
    return ""


def _starts_like_continuation_tail(text: str) -> bool:
    """
    Detect a text fragment that is more likely the tail of the row above than
    the beginning of a new table row. This is intentionally content-shape
    based: lowercase starts and closing punctuation are continuation signals;
    list markers, headings, and row numbers are not.
    """
    cleaned = _clean_extracted_text(text or "")
    if not cleaned:
        return False
    if _looks_like_table_index_marker(cleaned) or _looks_like_list_item(cleaned):
        return False
    if re.match(r"^(?:\d+(?:\.\d+)*|[IVXLCDM]{1,6})[.)]?\s+", cleaned, flags=re.IGNORECASE):
        return False
    first = _first_alpha_char(cleaned)
    if first and first.islower():
        return True
    return bool(re.match(r"^[,.;:)\]}]", cleaned))


def _ends_like_open_continuation(text: str) -> bool:
    cleaned = _clean_extracted_text(text or "")
    if not cleaned:
        return False
    if _hard_terminal_punctuation(cleaned):
        return False
    return not bool(re.search(r"[.?!;:]\s*$", cleaned))


def _lines_joined_text(lines: List[TextLine]) -> str:
    return _clean_extracted_text(" ".join(_clean_extracted_text(line.text) for line in lines if line.text))


def _repair_boundary_continuation_cell_lines(
    cell_lines: Dict[Tuple[int, int, int], List[TextLine]],
    table_regions: List[TableRegion],
    logger: Logger,
) -> int:
    """
    Move tiny top-of-cell continuation fragments back to the row above.

    Scanned ruled tables often have OCR bboxes that straddle a horizontal rule.
    If the detector places the next row's top boundary a few pixels too high,
    a tail such as "giấy theo quy..." can be mapped to the row below. The
    heuristic uses geometry plus lexical continuation shape, never a fixed
    phrase.
    """
    moved = 0
    for t_idx, table in enumerate(table_regions):
        if getattr(table, "skip_render", False):
            continue
        rows = int(getattr(table, "row_count", 0) or 0)
        cols = int(getattr(table, "col_count", 0) or 0)
        if rows < 2 or cols <= 0:
            continue
        for r in range(1, rows):
            for c in range(cols):
                key = (t_idx, r, c)
                current_lines = sort_lines_reading_order(cell_lines.get(key, []))
                if not current_lines:
                    continue
                if r >= len(table.cell_bboxes) or c >= len(table.cell_bboxes[r]):
                    continue
                if r - 1 >= len(table.cell_bboxes) or c >= len(table.cell_bboxes[r - 1]):
                    continue
                current_bbox = tuple(float(v) for v in table.cell_bboxes[r][c][:4])
                previous_bbox = tuple(float(v) for v in table.cell_bboxes[r - 1][c][:4])
                if not any(current_bbox) or not any(previous_bbox):
                    continue
                top = current_bbox[1]
                candidate = current_lines[0]
                candidate_text = _clean_extracted_text(candidate.text)
                if not _starts_like_continuation_tail(candidate_text):
                    continue
                line_top = float(candidate.y)
                line_bottom = float(candidate.y + candidate.height)
                line_h = max(float(candidate.height), 1.0)
                near_top_boundary = line_top <= top + max(line_h * 1.6, 14.0)
                straddles_or_hugs = line_top <= top + max(line_h * 0.55, 5.0) or line_bottom <= top + max(line_h * 2.2, 22.0)
                if not (near_top_boundary and straddles_or_hugs):
                    continue
                x_overlap = _axis_overlap_for_assignment(
                    float(candidate.x),
                    float(candidate.x + candidate.width),
                    previous_bbox[0],
                    previous_bbox[2],
                )
                if x_overlap / max(float(candidate.width), 1.0) < 0.45:
                    continue
                previous_text = _lines_joined_text(cell_lines.get((t_idx, r - 1, c), []))
                if previous_text and not _ends_like_open_continuation(previous_text):
                    continue
                cell_lines[key] = [line for line in cell_lines.get(key, []) if id(line) != id(candidate)]
                cell_lines.setdefault((t_idx, r - 1, c), []).append(candidate)
                moved += 1
    if moved:
        logger.log(f"Repaired {moved} table-boundary continuation fragment(s)")
    return moved


def _line_fragment_from_spans(parent: TextLine, spans: List[TextSpan]) -> Optional[TextLine]:
    spans = [span for span in spans if _clean_extracted_text(span.text)]
    if not spans:
        return None
    text = _clean_extracted_text("".join(span.text for span in spans))
    if not text:
        return None
    x0 = min(float(span.x) for span in spans)
    x1 = max(float(span.x) + float(span.width) for span in spans)
    y0 = min(float(span.y) for span in spans)
    font_size = _median([float(span.font_size) for span in spans if span.font_size], parent.font_size)
    height = max(parent.height, font_size * 1.15)
    copied_spans = [
        TextSpan(
            text=span.text,
            font_size=span.font_size,
            y=span.y,
            is_superscript=span.is_superscript,
            x=span.x,
            width=span.width,
            fg_gray=span.fg_gray,
            has_space_after=span.has_space_after,
        )
        for span in spans
    ]
    fragment = TextLine(
        text=text,
        x=x0,
        y=y0,
        width=max(0.0, x1 - x0),
        height=max(0.0, height),
        page=parent.page,
        font_size=font_size,
        spans=copied_spans,
        block_id=parent.block_id,
        paragraph_id=parent.paragraph_id,
        content_type=parent.content_type,
        fg_gray=parent.fg_gray,
        confidence=parent.confidence,
        semantic_type=parent.semantic_type,
        order=parent.order,
        source_line_id=parent.source_line_id,
        kie_labels=set(getattr(parent, "kie_labels", set())),
    )
    return fragment


def _word_sort_key_for_assignment(item) -> Tuple[int, float, float]:
    idx, word = item
    try:
        order = int(word.get("order", idx) or idx)
    except (TypeError, ValueError):
        order = idx
    return (
        order,
        float(word.get("y", 0.0) or 0.0),
        float(word.get("x", 0.0) or 0.0),
    )


def _residual_fragments_from_unassigned_words(parent: TextLine,
                                              words: List[dict],
                                              assigned_word_ids: Set[int]) -> List[TextLine]:
    """Keep non-table words when one OCR line straddles a table and body text."""
    residuals: List[TextLine] = []
    current: List[dict] = []
    for _idx, word in sorted(enumerate(words), key=_word_sort_key_for_assignment):
        if id(word) in assigned_word_ids:
            if current:
                fragment = _line_fragment_from_words(parent, current)
                if fragment is not None:
                    setattr(fragment, "_table_residual_fragment", True)
                    residuals.append(fragment)
                current = []
            continue
        current.append(word)
    if current:
        fragment = _line_fragment_from_words(parent, current)
        if fragment is not None:
            setattr(fragment, "_table_residual_fragment", True)
            residuals.append(fragment)
    return residuals


def assign_ocr_lines_to_table_cells_by_geometry(
    table_regions: List[TableRegion],
    pdf_lines: List[TextLine],
    logger: Logger,
    candidate_lines: Optional[List[TextLine]] = None,
    rebuild_cells: bool = True,
    preserve_header_rows: bool = False,
) -> set:
    """
    Assign OCR text inside table bboxes to exactly one best cell by geometry.
    Scoring uses text bbox overlap, row/column band overlap, and small boundary
    gaps. If OCR word boxes are available, split a line by word only when its
    words truly fall into different cells.
    """
    if not table_regions or not pdf_lines:
        return set()

    candidates = candidate_lines if candidate_lines is not None else pdf_lines
    table_bboxes: Dict[int, Tuple[float, float, float, float]] = {}
    duplicate_cell_targets: Dict[Tuple[int, Tuple[float, float, float, float]], List[Tuple[int, int, int]]] = {}
    for t_idx, table in enumerate(table_regions):
        if getattr(table, "skip_render", False):
            continue
        table_bboxes[t_idx] = _table_bbox(table)
        for r, c, bbox in _iter_table_cell_bboxes(table):
            duplicate_cell_targets.setdefault((t_idx, _cell_bbox_key_for_assignment(bbox)), []).append((t_idx, r, c))

    cell_lines: Dict[Tuple[int, int, int], List[TextLine]] = {}
    assigned_source_ids = set()
    residual_fragment_count = 0

    def add_to_cell(key: Tuple[int, int, int], line: TextLine):
        t_idx, r, c = key
        bbox = table_regions[t_idx].cell_bboxes[r][c]
        targets = duplicate_cell_targets.get((t_idx, _cell_bbox_key_for_assignment(tuple(bbox[:4]))), [key])
        for target in targets:
            cell_lines.setdefault(target, []).append(line)

    for line in candidates:
        words = [word for word in getattr(line, "word_items", []) or [] if _clean_extracted_text(str(word.get("text") or ""))]
        if words:
            grouped_words: Dict[Tuple[int, int, int], List[dict]] = {}
            assigned_word_ids: Set[int] = set()
            for word in words:
                key = _best_cell_for_text_bbox(_word_bbox_for_assignment(word), line.page, table_regions, table_bboxes)
                if key is not None:
                    grouped_words.setdefault(key, []).append(word)
                    assigned_word_ids.add(id(word))
            if grouped_words:
                assigned_source_ids.add(id(line))
                all_words_assigned = len(assigned_word_ids) == len(words)
                if len(grouped_words) == 1 and all_words_assigned:
                    add_to_cell(next(iter(grouped_words.keys())), line)
                else:
                    for key, group in grouped_words.items():
                        group.sort(key=lambda item: int(item.get("order", 0) or 0))
                        fragment = _line_fragment_from_words(line, group)
                        if fragment is not None:
                            add_to_cell(key, fragment)
                    if not all_words_assigned:
                        residuals = _residual_fragments_from_unassigned_words(line, words, assigned_word_ids)
                        if residuals:
                            setattr(line, "_table_residual_fragments", residuals)
                            residual_fragment_count += len(residuals)
                continue

        key = _best_cell_for_text_bbox(_line_bbox_for_assignment(line), line.page, table_regions, table_bboxes)
        if key is not None:
            assigned_source_ids.add(id(line))
            add_to_cell(key, line)

    if cell_lines:
        _repair_boundary_continuation_cell_lines(cell_lines, table_regions, logger)

    if rebuild_cells:
        for t_idx, table in enumerate(table_regions):
            if getattr(table, "skip_render", False):
                continue
            if getattr(table, "source_segments", None):
                continue
            rows = int(getattr(table, "row_count", 0) or 0)
            cols = int(getattr(table, "col_count", 0) or 0)
            existing_cells = getattr(table, "cells", []) or []
            preserve_until = -1
            if preserve_header_rows:
                for idx, row in enumerate(existing_cells[:4]):
                    if _is_numeric_header_row(list(row)):
                        preserve_until = idx
                        break
            rebuilt_cells = [[""] * cols for _ in range(rows)]
            for r in range(rows):
                for c in range(cols):
                    if preserve_until >= 0 and r <= preserve_until and r < len(existing_cells) and c < len(existing_cells[r]):
                        rebuilt_cells[r][c] = str(existing_cells[r][c])
                        continue
                    key = (t_idx, r, c)
                    lines_for_cell = _unique_lines_for_cell(cell_lines.get(key, []))
                    if lines_for_cell:
                        bbox = table.cell_bboxes[r][c]
                        x0 = min(float(bbox[0]), *(line.x for line in lines_for_cell))
                        x1 = max(float(bbox[2]), *(line.x + line.width for line in lines_for_cell))
                        rebuilt_cells[r][c] = _text_from_table_cell_lines(lines_for_cell, table.page, x0, x1, logger)
                    elif r < len(existing_cells) and c < len(existing_cells[r]):
                        rebuilt_cells[r][c] = str(existing_cells[r][c])
            table.cells = rebuilt_cells

    if assigned_source_ids:
        logger.log(f"Assigned {len(assigned_source_ids)} OCR table lines by geometry cell mapping")
    if residual_fragment_count:
        logger.log(f"Preserved {residual_fragment_count} non-table fragment(s) from OCR lines crossing table boundaries")
    return assigned_source_ids


def _unique_lines_for_cell(lines: List[TextLine]) -> List[TextLine]:
    seen = set()
    unique = []
    for line in lines:
        line_id = id(line)
        if line_id in seen:
            continue
        seen.add(line_id)
        unique.append(line)
    return sort_lines_reading_order(unique)


def _compact_cell_text_for_compare(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "").casefold()
    return re.sub(r"\s+", "", text)


def _merge_rescued_cell_text(existing: str, rebuilt: str, rescued_lines: List[TextLine]) -> str:
    existing = clean_ocr_cell_text(existing or "")
    rebuilt = clean_ocr_cell_text(rebuilt or "")
    if not existing:
        return rebuilt
    if not rebuilt:
        return existing

    compact_existing = _compact_cell_text_for_compare(existing)
    compact_rebuilt = _compact_cell_text_for_compare(rebuilt)
    if compact_existing and compact_existing in compact_rebuilt:
        return rebuilt
    return existing


def assign_orphan_ocr_lines_to_table_cells(
    table_regions: List[TableRegion],
    pdf_lines: List[TextLine],
    logger: Logger,
    candidate_lines: Optional[List[TextLine]] = None,
) -> set:
    """
    Recover OCR lines that are inside a detected table but missed by strict
    center-point cell mapping. Assignment is geometry-only: bbox overlap,
    shared row band, and small boundary gaps for text that sits on grid lines.
    """
    if not table_regions or not pdf_lines:
        return set()

    page_lines: Dict[int, List[TextLine]] = {}
    for line in pdf_lines:
        page_lines.setdefault(line.page, []).append(line)

    candidates = candidate_lines if candidate_lines is not None else pdf_lines
    assigned_ids = set()
    cell_lines: Dict[Tuple[int, int, int], List[TextLine]] = {}
    table_bboxes: Dict[int, Tuple[float, float, float, float]] = {}

    for t_idx, table in enumerate(table_regions):
        if getattr(table, "skip_render", False):
            continue
        rows = int(getattr(table, "row_count", 0) or 0)
        cols = int(getattr(table, "col_count", 0) or 0)
        if rows <= 0 or cols <= 0:
            continue
        table_bbox = _table_bbox(table)
        table_bboxes[t_idx] = table_bbox
        table_page_lines = page_lines.get(table.page, [])
        for r in range(rows):
            for c in range(cols):
                if r >= len(table.cell_bboxes) or c >= len(table.cell_bboxes[r]):
                    continue
                bbox = tuple(float(v) for v in table.cell_bboxes[r][c][:4])
                if not any(bbox):
                    continue
                lines = get_lines_in_rect(bbox, table_page_lines)
                if lines:
                    cell_lines[(t_idx, r, c)] = list(lines)
                    assigned_ids.update(id(line) for line in lines)

    rescued_by_cell: Dict[Tuple[int, int, int], List[TextLine]] = {}
    rescued_ids = set()
    for line in candidates:
        line_id = id(line)
        if line_id in assigned_ids:
            continue

        best_key = None
        best_score = None
        for t_idx, table in enumerate(table_regions):
            if getattr(table, "skip_render", False) or table.page != line.page:
                continue
            table_bbox = table_bboxes.get(t_idx)
            if table_bbox is None:
                continue
            rows = int(getattr(table, "row_count", 0) or 0)
            cols = int(getattr(table, "col_count", 0) or 0)
            for r in range(rows):
                for c in range(cols):
                    if r >= len(table.cell_bboxes) or c >= len(table.cell_bboxes[r]):
                        continue
                    bbox = tuple(float(v) for v in table.cell_bboxes[r][c][:4])
                    score = _cell_line_assignment_score(line, bbox, table_bbox)
                    if score is None:
                        continue
                    if best_score is None or score > best_score:
                        best_score = score
                        best_key = (t_idx, r, c)

        if best_key is not None:
            rescued_by_cell.setdefault(best_key, []).append(line)
            rescued_ids.add(line_id)

    for key, rescued_lines in rescued_by_cell.items():
        t_idx, r, c = key
        table = table_regions[t_idx]
        original_lines = cell_lines.get(key, [])
        merged_lines = _unique_lines_for_cell(original_lines + rescued_lines)
        if not merged_lines:
            continue
        bx0, _by0, bx1, _by1 = table.cell_bboxes[r][c]
        x0 = min(float(bx0), *(line.x for line in merged_lines))
        x1 = max(float(bx1), *(line.x + line.width for line in merged_lines))
        rebuilt = _text_from_table_cell_lines(merged_lines, table.page, x0, x1, logger)
        existing = ""
        if r < len(table.cells) and c < len(table.cells[r]):
            existing = str(table.cells[r][c])
        table.cells[r][c] = _merge_rescued_cell_text(existing, rebuilt, rescued_lines)

    if rescued_ids:
        logger.log(f"Assigned {len(rescued_ids)} OCR table lines by geometry-overlap fallback")
    return rescued_ids


def _build_table_rows_from_bbox(page: int,
                                bbox: Tuple[float, float, float, float],
                                col_intervals: List[Tuple[float, float]],
                                pdf_lines: List[TextLine],
                                logger: Logger) -> Tuple[List[List[str]], List[List[Tuple[float, float, float, float]]]]:
    x0, y0, x1, y1 = bbox
    col_count = len(col_intervals)
    if col_count <= 0:
        return [], []

    lines = [l for l in pdf_lines if l.page == page and _line_in_bbox(l, bbox)]
    if not lines:
        return [], []

    line_cols: Dict[int, List[TextLine]] = {i: [] for i in range(col_count)}
    key_y = []
    for line in lines:
        cx = line.x + line.width / 2
        col_idx = None
        for i, (cx0, cx1) in enumerate(col_intervals):
            if cx0 <= cx <= cx1:
                col_idx = i
                break
        if col_idx is None:
            nearest = min(range(col_count), key=lambda i: abs(cx - ((col_intervals[i][0] + col_intervals[i][1]) / 2)))
            if abs(cx - ((col_intervals[nearest][0] + col_intervals[nearest][1]) / 2)) < 80:
                col_idx = nearest
        if col_idx is None:
            continue
        line_cols[col_idx].append(line)
        if _is_table_row_key_line(line, col_idx):
            key_y.append(line.y)

    starts = _cluster_y_positions(key_y)
    if starts and starts[0] <= y0 + 45:
        starts[0] = y0
    else:
        starts = [y0] + starts
    starts = sorted(set(round(v, 2) for v in starts if y0 - 2 <= v <= y1 + 2))
    if not starts:
        starts = [y0]

    bands = []
    for idx, start in enumerate(starts):
        end = starts[idx + 1] if idx + 1 < len(starts) else y1
        if end - start >= 8:
            bands.append((float(start), float(end)))

    rows: List[List[str]] = []
    row_boxes: List[List[Tuple[float, float, float, float]]] = []
    for ry0, ry1 in bands:
        row = []
        boxes = []
        for col_idx, (cx0, cx1) in enumerate(col_intervals):
            cell_lines = [
                l for l in line_cols[col_idx]
                if ry0 <= l.y + l.height / 2 <= ry1
            ]
            row.append(_text_from_table_cell_lines(cell_lines, page, cx0, cx1, logger))
            boxes.append((cx0, ry0, cx1, ry1))
        if any(cell.strip() for cell in row):
            rows.append(row)
            row_boxes.append(boxes)

    return rows, row_boxes


def _is_continuation_table_row(row: List[str]) -> bool:
    if not row:
        return False
    key = _norm_table_key(" ".join(row[:2]))
    if re.fullmatch(r"\d{1,3}", key or ""):
        return False
    leading = _clean_extracted_text(" ".join(str(cell) for cell in row[:2]))
    if leading:
        words = [w for w in re.split(r"\s+", leading) if w]
        if len(words) <= 8 and len(leading) <= 90 and not leading.endswith((".", ";", ":")):
            return False
    return any((cell or "").strip() for cell in row[2:])


def _append_row_text(dst: str, src: str) -> str:
    dst = (dst or "").strip()
    src = (src or "").strip()
    if not src:
        return dst
    if not dst:
        return src
    if src in dst:
        return dst
    return dst + "\n" + src


def _normalize_table_key_cells(row: List[str]) -> List[str]:
    # Keep OCR content as-is. Earlier versions tried to synthesize a normalized
    # organization name from row fragments; that is too case-specific and can
    # invent text that was not recognized by OCR.
    if len(row) >= 2 and not _clean_extracted_text(str(row[0])):
        second = _clean_extracted_text(str(row[1]))
        match = re.match(r"^(\d{1,3})\s+(.+)$", second)
        if match and re.search(r"[A-Za-zÀ-Ỵà-ỵĐđ]", match.group(2)):
            row[0] = match.group(1)
            row[1] = match.group(2).strip()
    return row


def _merge_table_rows(base_rows: List[List[str]], incoming_rows: List[List[str]], col_count: int) -> List[List[str]]:
    rows = [list(row[:col_count]) + [""] * max(0, col_count - len(row)) for row in base_rows]
    for raw_row in incoming_rows:
        row = list(raw_row[:col_count]) + [""] * max(0, col_count - len(raw_row))
        row = _normalize_table_key_cells(row)
        if rows and _is_continuation_table_row(row):
            for c in range(col_count):
                rows[-1][c] = _append_row_text(rows[-1][c], row[c])
        else:
            rows.append(row)
    return rows


def _numeric_header_sequence(row: List[str], expected_count: int) -> bool:
    if expected_count <= 0 or len(row) < expected_count:
        return False
    values = [_norm_table_key(cell) for cell in row[:expected_count]]
    digits = [re.sub(r"\D+", "", value) for value in values]
    expected = [str(i) for i in range(1, expected_count + 1)]
    return digits == expected or (
        expected_count >= 5
        and digits[:5] in (["1", "2", "3", "4", "5"], ["", "2", "3", "4", "5"])
    )


def _is_numeric_header_row(row: List[str]) -> bool:
    values = [_norm_table_key(cell) for cell in row[:5]]
    digits = [re.sub(r"\D+", "", value) for value in values]
    if len(digits) < 5:
        return False
    if digits[:5] in (["1", "2", "3", "4", "5"], ["", "2", "3", "4", "5"]):
        return True
    # OCR sometimes reads column number 4 as 5 in tiny header rows.
    return digits[0] in ("", "1") and digits[1:3] == ["2", "3"] and digits[3] in ("4", "5") and digits[4] == "5"


def _collapse_spaced_acronym(text: str) -> str:
    stripped = _clean_extracted_text(text)
    compact = _norm_table_key(stripped)
    if compact == "STT":
        return "STT"
    return stripped


def _remove_prefix_by_norm(text: str, prefix: str) -> str:
    text = _clean_extracted_text(text)
    prefix_norm = _norm_table_key(prefix)
    if not text or not prefix_norm:
        return text
    words = text.split()
    for idx in range(1, len(words) + 1):
        if _norm_table_key(" ".join(words[:idx])) == prefix_norm:
            return " ".join(words[idx:]).strip()
    return text


def _extract_group_header_text(header_cells: List[str]) -> str:
    candidates = []
    for cell in header_cells:
        text = _clean_extracted_text(cell)
        norm = _unaccent_upper(text)
        if "CONG TAC" not in norm:
            continue
        # Use text that is already present in OCR. Stop before common lower-level
        # header phrases so a merged cell can be split into group/subheader.
        stop_positions = []
        for stop in ("SO LUONG", "GHI CHU", "STT", "TEN "):
            pos = norm.find(stop, norm.find("CONG TAC") + len("CONG TAC"))
            if pos > 0:
                stop_positions.append(pos)
        start = norm.find("CONG TAC")
        end = min(stop_positions) if stop_positions else len(text)
        phrase = text[start:end].strip(" :-")
        if phrase:
            candidates.append(phrase)
    if not candidates:
        return ""
    return min(candidates, key=len)


def _normalize_continued_table_header_rows(rows: List[List[str]], col_count: int) -> List[List[str]]:
    if col_count < 2 or not rows:
        return rows

    numeric_idx = None
    for idx, row in enumerate(rows[:4]):
        if _is_numeric_header_row(row):
            numeric_idx = idx
            break
    if numeric_idx is None:
        return rows

    header_rows = [
        list(row[:col_count]) + [""] * max(0, col_count - len(row))
        for row in rows[:numeric_idx]
    ]
    body_rows = [
        list(row[:col_count]) + [""] * max(0, col_count - len(row))
        for row in rows[numeric_idx + 1:]
    ]
    number_row = [str(i) for i in range(1, col_count + 1)]

    if not header_rows:
        return [number_row] + body_rows

    primary = [_collapse_spaced_acronym(cell) for cell in header_rows[0]]
    group_title = _extract_group_header_text([cell for row in header_rows for cell in row])

    if group_title and len(header_rows) == 1:
        lower = list(primary)
        for c in range(2, col_count):
            lower[c] = _remove_prefix_by_norm(lower[c], group_title)
        top = [""] * col_count
        top[0] = primary[0]
        top[1] = primary[1]
        for c in range(2, col_count):
            top[c] = group_title if c < col_count - 1 or not lower[c] else ""
        return [top, lower, number_row] + body_rows

    normalized_headers = []
    for row in header_rows:
        normalized_headers.append([_collapse_spaced_acronym(cell) for cell in row])
    if len(normalized_headers) >= 2:
        for c in range(min(2, col_count)):
            top = _clean_extracted_text(normalized_headers[0][c])
            lower = _clean_extracted_text(normalized_headers[1][c])
            combined = _collapse_spaced_acronym(f"{top} {lower}".strip())
            word_count = len([w for w in combined.split() if w])
            if (
                top and lower
                and word_count <= 4
                and len(top) <= 24
                and len(lower) <= 24
                and not re.search(r"[.;:]", combined)
            ):
                normalized_headers[0][c] = combined
                normalized_headers[1][c] = combined
    return normalized_headers + [number_row] + body_rows


def _repair_continued_table_headers(rows: List[List[str]], col_count: int) -> List[List[str]]:
    return _normalize_continued_table_header_rows(rows, col_count)


def _matching_layout_bbox_for_table(table: TableRegion,
                                    layout_regions_by_page: Dict[int, List[dict]]) -> Optional[Tuple[float, float, float, float]]:
    best_bbox = None
    best_overlap = 0.0
    table_bbox = _table_bbox(table)
    for region in (layout_regions_by_page or {}).get(table.page, []):
        if region.get("type") != "table":
            continue
        bbox = region.get("bbox_pdf")
        if not bbox or len(bbox) < 4:
            continue
        layout_bbox = tuple(float(v) for v in bbox[:4])
        overlap = _bbox_intersection_ratio(table_bbox, layout_bbox)
        if overlap > best_overlap:
            best_overlap = overlap
            best_bbox = layout_bbox
    return best_bbox if best_overlap >= 0.45 else None


def _median_interval_width(intervals: List[Tuple[float, float]]) -> float:
    widths = sorted(max(0.0, x1 - x0) for x0, x1 in intervals if x1 > x0)
    if not widths:
        return 0.0
    return widths[len(widths) // 2]


def _append_empty_column(table: TableRegion, x0: float, x1: float):
    if x1 <= x0:
        return
    for row in table.cells:
        row.append("")
    for r_idx, row_boxes in enumerate(table.cell_bboxes):
        ys = [(bbox[1], bbox[3]) for bbox in row_boxes if any(bbox) and bbox[3] > bbox[1]]
        if ys:
            y0 = min(y[0] for y in ys)
            y1 = max(y[1] for y in ys)
        else:
            row_height = max((table.y_bottom - table.y_top) / max(table.row_count, 1), 12.0)
            y0 = table.y_top + r_idx * row_height
            y1 = y0 + row_height
        row_boxes.append((x0, y0, x1, y1))
    table.col_count += 1
    setattr(table, "x_right", max(float(getattr(table, "x_right", x1) or x1), x1))


def _maybe_append_group_grid_column(table: TableRegion) -> bool:
    """
    Some structure engines emit only the visible logical columns for a grouped
    header, while DOCX needs one extra narrow grid column so the top-level group
    cell can span over a lower "Ghi chú/Note" header without shifting body
    cells. Detect this from the header geometry/text pattern rather than from a
    specific document.
    """
    if table.row_count < 3 or table.col_count < 5 or table.col_count > 8:
        return False

    rows = table.cells
    first = list(rows[0][:table.col_count]) + [""] * max(0, table.col_count - len(rows[0]))
    second = list(rows[1][:table.col_count]) + [""] * max(0, table.col_count - len(rows[1]))
    third = list(rows[2][:table.col_count]) + [""] * max(0, table.col_count - len(rows[2]))

    grouped = [_clean_extracted_text(cell) for cell in first[2:]]
    if len(grouped) < 3 or any(not cell for cell in grouped):
        return False
    group_keys = {_norm_table_key(cell) for cell in grouped if cell}
    if len(group_keys) != 1:
        return False

    trailing_lower_header = _clean_extracted_text(second[-1])
    if not trailing_lower_header:
        return False

    lower_header_count = sum(1 for cell in second[2:] if _clean_extracted_text(cell))
    if lower_header_count < max(2, table.col_count - 3):
        return False

    body_rows = rows[3:]
    if body_rows:
        trailing_body_filled = sum(
            1
            for row in body_rows
            if len(row) >= table.col_count and _clean_extracted_text(row[-1])
        )
        if trailing_body_filled / max(len(body_rows), 1) > 0.35:
            return False

    numbered = [_clean_extracted_text(cell) for cell in third]
    numeric_labels = sum(1 for cell in numbered if re.fullmatch(r"\d{1,3}", cell or ""))
    if numeric_labels and numeric_labels < max(2, table.col_count - 2):
        return False

    intervals = _table_column_intervals(table)
    if len(intervals) != table.col_count:
        return False
    widths = [max(0.0, x1 - x0) for x0, x1 in intervals if x1 > x0]
    if not widths:
        return False
    narrow_width = max(6.0, min(sorted(widths)[len(widths) // 2] * 0.12, 12.0))
    x0 = intervals[-1][1]
    x1 = x0 + narrow_width

    group_text = grouped[0]
    _append_empty_column(table, x0, x1)
    if table.cells and table.cells[0]:
        table.cells[0][-1] = group_text
    return True


def _append_empty_row(table: TableRegion, col_intervals: List[Tuple[float, float]], y0: float, y1: float):
    if y1 <= y0 or len(col_intervals) != table.col_count:
        return
    table.cells.append([""] * table.col_count)
    table.cell_bboxes.append([(x0, y0, x1, y1) for x0, x1 in col_intervals])
    table.row_count += 1
    table.y_bottom = max(float(table.y_bottom), y1)


def _normalize_grouped_header_rows(table: TableRegion) -> bool:
    """
    Normalize generic two-level table headers when the detector places the
    group label in the visual center column and leaves leading group cells
    empty. The content is only moved, never invented.
    """
    if table.row_count < 2 or table.col_count < 4:
        return False
    rows = table.cells
    first = list(rows[0][:table.col_count]) + [""] * max(0, table.col_count - len(rows[0]))
    second = list(rows[1][:table.col_count]) + [""] * max(0, table.col_count - len(rows[1]))

    prefix = 2 if table.col_count >= 5 else 1
    if any(_clean_extracted_text(first[c]) for c in range(min(prefix, len(first)))):
        return False

    fixed_headers = [_clean_extracted_text(second[c]) for c in range(min(prefix, len(second)))]
    if len([h for h in fixed_headers if h and len(h) <= 40]) < min(prefix, len(fixed_headers)):
        return False

    group_cells = [
        (idx, _clean_extracted_text(cell))
        for idx, cell in enumerate(first[prefix:], prefix)
        if _clean_extracted_text(cell)
    ]
    if len(group_cells) != 1:
        return False

    group_idx, group_text = group_cells[0]
    if group_idx <= prefix:
        return False

    for c in range(prefix):
        first[c] = _collapse_spaced_acronym(second[c])
    for c in range(prefix, table.col_count):
        first[c] = ""
    first[prefix] = group_text
    rows[0] = first
    return True


def repair_split_table_header_fragments(table_regions: List[TableRegion], logger: Logger) -> List[TableRegion]:
    """
    Normalize multi-row table headers whose fixed columns are OCR-split across
    rows and whose group label lands in only one covered grid cell.

    The signal is structural: a numeric column-index row below two header rows.
    We combine short vertical header fragments and record an explicit
    horizontal span for a group label whose adjacent top cell is blank but both
    lower cells have subheaders.
    """
    repaired_tables = 0
    for table in table_regions or []:
        if getattr(table, "skip_render", False):
            continue
        cols = int(getattr(table, "col_count", 0) or 0)
        if cols < 4 or int(getattr(table, "row_count", 0) or 0) < 3 or len(table.cells) < 3:
            continue

        numeric_idx = None
        for idx, row in enumerate(table.cells[:4]):
            normalized_row = list(row[:cols]) + [""] * max(0, cols - len(row))
            if _is_numeric_header_row(normalized_row):
                numeric_idx = idx
                break
        if numeric_idx is None or numeric_idx < 2:
            continue

        rows = [
            list(row[:cols]) + [""] * max(0, cols - len(row))
            for row in table.cells
        ]
        top = rows[0]
        lower = rows[1]
        numeric = rows[numeric_idx]
        changed = False

        for c in range(cols):
            top_text = _clean_extracted_text(str(top[c]))
            lower_text = _clean_extracted_text(str(lower[c]))
            numeric_text = _clean_extracted_text(str(numeric[c])) if c < len(numeric) else ""
            if not top_text or not lower_text or not re.fullmatch(r"\d{1,3}", numeric_text or ""):
                continue
            top_key = _norm_table_key(top_text)
            lower_key = _norm_table_key(lower_text)
            if lower_key and lower_key in top_key:
                combined = top_text
            elif top_key and top_key in lower_key:
                combined = lower_text
            else:
                combined = _collapse_spaced_acronym(f"{top_text} {lower_text}")
            words = [word for word in re.split(r"\s+", combined) if word]
            if len(words) > 5 or len(combined) > 45:
                continue
            top[c] = combined
            lower[c] = ""
            changed = True

        spans = dict(getattr(table, "horizontal_text_spans", {}) or {})
        for c in range(1, cols):
            group_text = _clean_extracted_text(str(top[c]))
            if not group_text or _clean_extracted_text(str(top[c - 1])):
                continue
            left_lower = _clean_extracted_text(str(lower[c - 1]))
            this_lower = _clean_extracted_text(str(lower[c]))
            if not left_lower or not this_lower:
                continue
            if len(group_text) < 4 or len(group_text) > 80:
                continue
            if c + 1 < cols and _clean_extracted_text(str(top[c + 1])) == group_text:
                continue
            spans.setdefault(0, (c - 1, c, group_text))
            changed = True
            break

        if changed:
            table.cells = rows
            if spans:
                setattr(table, "horizontal_text_spans", spans)
            repaired_tables += 1

    if repaired_tables:
        logger.log(f"Repaired split header fragments in {repaired_tables} table(s)")
    return table_regions


def _has_sparse_group_header(table: TableRegion) -> bool:
    if table.row_count < 2 or table.col_count < 5:
        return False
    first = list(table.cells[0][:table.col_count]) + [""] * max(0, table.col_count - len(table.cells[0]))
    second = list(table.cells[1][:table.col_count]) + [""] * max(0, table.col_count - len(table.cells[1]))
    top_nonempty = _row_nonempty_count(first)
    if top_nonempty <= 0 or top_nonempty > table.col_count - 2:
        return False
    if not any(not _clean_extracted_text(cell) for cell in first[2:]):
        return False
    lower_labels = sum(1 for cell in second if _clean_extracted_text(cell))
    return lower_labels >= max(3, table.col_count - 2)


def postprocess_table_layout_grids(table_regions: List[TableRegion],
                                   layout_regions_by_page: Dict[int, List[dict]],
                                   logger: Logger) -> List[TableRegion]:
    """
    Use document-layout table boxes to repair detector grids that stop just
    inside the actual ruled table border. This covers narrow empty edge columns
    and blank trailing rows without relying on document-specific text.
    """
    added_cols = 0
    added_rows = 0
    normalized_headers = 0
    added_grid_cols = 0

    for table in table_regions:
        if getattr(table, "skip_render", False):
            continue
        if getattr(table, "source", "") == "docling_tableformer":
            # Docling TableFormer already runs its own table-structure matching.
            # The generic layout-gap repair below was designed for detector
            # grids that stop inside ruled borders; applying it to Docling can
            # create extra columns/rows from harmless bbox padding.
            continue
        if table.row_count <= 0 or table.col_count <= 0:
            continue

        layout_bbox = _matching_layout_bbox_for_table(table, layout_regions_by_page)
        intervals = _table_column_intervals(table)
        if layout_bbox and len(intervals) == table.col_count:
            median_width = _median_interval_width(intervals)
            right_gap = layout_bbox[2] - intervals[-1][1]
            header_like = _table_has_header(table) or _looks_like_table_header_text(_table_text(table))
            if (
                header_like
                and _has_sparse_group_header(table)
                and table.col_count >= 4
                and right_gap >= 10.0
                and right_gap <= max(36.0, median_width * 0.65)
            ):
                _append_empty_column(table, intervals[-1][1], layout_bbox[2])
                intervals = _table_column_intervals(table)
                added_cols += 1

            bottom_gap = layout_bbox[3] - table.y_bottom
            if (
                header_like
                and table.row_count <= 2
                and table.col_count >= 3
                and len(intervals) == table.col_count
                and 8.0 <= bottom_gap <= 40.0
            ):
                _append_empty_row(table, intervals, table.y_bottom, layout_bbox[3])
                added_rows += 1

        if _maybe_append_group_grid_column(table):
            added_grid_cols += 1

        if _normalize_grouped_header_rows(table):
            normalized_headers += 1

    if added_cols:
        logger.log(f"Added {added_cols} narrow trailing table column(s) from layout grid")
    if added_rows:
        logger.log(f"Added {added_rows} trailing blank table row(s) from layout grid")
    if added_grid_cols:
        logger.log(f"Added {added_grid_cols} grouped-header grid column(s)")
    if normalized_headers:
        logger.log(f"Normalized {normalized_headers} grouped table header row(s)")
    return table_regions


def repair_wrapped_group_header_tokens(table_regions: List[TableRegion], logger: Logger) -> List[TableRegion]:
    """
    Repair multi-row group headers when the OCR/model drops a continuation
    token into the lower subheader row.

    Example pattern by geometry, not by document text:
      top:   [Group A,] [Group A,]
      lower: [Cases]   [II %]
    becomes:
      top:   [Group A, II] [Group A, II]
      lower: [Cases]       [%]
    """
    repaired = 0
    for table in table_regions or []:
        if getattr(table, "skip_render", False):
            continue
        if table.row_count < 2 or table.col_count < 3:
            continue
        if len(table.cells) < 2:
            continue

        first = list(table.cells[0][:table.col_count]) + [""] * max(0, table.col_count - len(table.cells[0]))
        second = list(table.cells[1][:table.col_count]) + [""] * max(0, table.col_count - len(table.cells[1]))
        changed = False

        for c in range(0, table.col_count - 1):
            top = _clean_extracted_text(first[c])
            top_next = _clean_extracted_text(first[c + 1])
            lower_next = _clean_extracted_text(second[c + 1])
            if not top or _norm_table_key(top) != _norm_table_key(top_next):
                continue
            match = re.match(r"^([IVXLCDM]{1,8})\s+(.+)$", lower_next, re.IGNORECASE)
            if not match:
                continue
            suffix = match.group(1).upper()
            remainder = _clean_extracted_text(match.group(2))
            if not remainder or len(remainder) > 12:
                continue
            if not (top.endswith((",", "/", "-", "&")) or re.search(r"\b[IVXLCDM]+\s*[,/\\-]\s*$", top, re.IGNORECASE)):
                continue
            combined = _clean_extracted_text(f"{top} {suffix}")
            first[c] = combined
            first[c + 1] = combined
            second[c + 1] = remainder
            changed = True

        if changed:
            for c in range(table.col_count):
                if _norm_table_key(first[c]) == _norm_table_key(second[c]) and _clean_extracted_text(first[c]):
                    second[c] = ""
            table.cells[0] = first
            table.cells[1] = second
            repaired += 1

    if repaired:
        logger.log(f"Repaired wrapped continuation token(s) in {repaired} grouped table header(s)")
    return table_regions


def _dummy_table_bboxes(rows: int,
                        col_intervals: List[Tuple[float, float]],
                        y_top: float,
                        row_height: float = 24.0) -> List[List[Tuple[float, float, float, float]]]:
    out = []
    for r in range(rows):
        ry0 = y_top + r * row_height
        ry1 = ry0 + row_height
        out.append([(x0, ry0, x1, ry1) for x0, x1 in col_intervals])
    return out


def repair_continued_tables(table_regions: List[TableRegion],
                            layout_regions_by_page: Dict[int, List[dict]],
                            pdf_lines: List[TextLine],
                            page_info: dict,
                            logger: Logger) -> List[TableRegion]:
    """
    Keep table render units page-local.

    Earlier versions synthesized one DOCX table across physical PDF pages
    when a table appeared to continue on the next page. That preserves logical
    continuity only when every page has the same grid, but it is fragile for
    scanned administrative documents: continuation pages often have different
    detected schemas, merged cells, missing narrow columns, or partial rows.
    For visual fidelity, the source page boundary is a hard rendering boundary.

    We still use layout table regions as a recovery signal: if layout detected
    a table region but the structure recognizer missed it entirely, build a
    page-local fallback table from the most recent page schema. Detected tables
    are never marked ``skip_render`` here.
    """
    if not table_regions:
        return table_regions

    detected = list(table_regions)
    layout_items = []
    for page, regions in (layout_regions_by_page or {}).items():
        for region in regions:
            if region.get("type") != "table":
                continue
            bbox = region.get("bbox_pdf")
            if not bbox or len(bbox) < 4:
                continue
            layout_items.append({
                "page": page,
                "bbox": tuple(float(v) for v in bbox[:4]),
                "detected": None,
            })

    for item in layout_items:
        best = None
        best_overlap = 0.0
        for table in detected:
            if table.page != item["page"]:
                continue
            overlap = _bbox_intersection_ratio(_table_bbox(table), item["bbox"])
            if overlap > best_overlap:
                best = table
                best_overlap = overlap
        if best is not None and best_overlap >= 0.45:
            item["detected"] = best

    # Include detected tables when no layout region was available.
    for table in detected:
        if any(item.get("detected") is table for item in layout_items):
            continue
        layout_items.append({"page": table.page, "bbox": _table_bbox(table), "detected": table})

    layout_items.sort(key=lambda i: (i["page"], i["bbox"][1]))
    if not layout_items:
        return table_regions

    repaired: List[TableRegion] = list(table_regions)
    fallback_count = 0
    last_col_intervals: List[Tuple[float, float]] = []

    for item in layout_items:
        table = item.get("detected")
        if table is not None:
            table.y_top = item["bbox"][1]
            table.y_bottom = item["bbox"][3]
            intervals = _table_column_intervals(table)
            if len(intervals) >= 2:
                last_col_intervals = intervals
            continue

        if len(last_col_intervals) < 2:
            continue
        rows, boxes = _build_table_rows_from_bbox(
            item["page"], item["bbox"], last_col_intervals, pdf_lines, logger
        )
        if not rows:
            continue
        fallback = TableRegion(
            page=item["page"],
            y_top=item["bbox"][1],
            y_bottom=item["bbox"][3],
            cells=rows,
            row_count=len(rows),
            col_count=len(last_col_intervals),
            cell_bboxes=boxes,
        )
        repaired.append(fallback)
        fallback_count += 1

    if fallback_count:
        logger.log(f"Added {fallback_count} page-local fallback table segment(s)")
    return repaired


def _horizontal_overlap_ratio(a: Tuple[float, float, float, float],
                              b: Tuple[float, float, float, float]) -> float:
    aw = max(a[2] - a[0], 1e-6)
    bw = max(b[2] - b[0], 1e-6)
    overlap = max(0.0, min(a[2], b[2]) - max(a[0], b[0]))
    return overlap / max(min(aw, bw), 1e-6)


def _table_is_full_page_segment(table: TableRegion, page_info: dict) -> bool:
    page_height = float(page_info.get(table.page, {}).get("height", 842) or 842)
    height = max(0.0, float(table.y_bottom) - float(table.y_top))
    return height >= page_height * 0.35


def _regrid_table_columns(table: TableRegion,
                          col_intervals: List[Tuple[float, float]],
                          logger: Logger) -> bool:
    if len(col_intervals) <= int(getattr(table, "col_count", 0) or 0):
        return False

    row_bands: List[Tuple[float, float]] = []
    for row_boxes in getattr(table, "cell_bboxes", []) or []:
        y0, y1 = _row_y_bounds(row_boxes)
        if y1 > y0 and y1 - y0 >= 4:
            row_bands.append((float(y0), float(y1)))

    if not row_bands:
        return False

    old_shape = (int(getattr(table, "row_count", 0) or 0), int(getattr(table, "col_count", 0) or 0))
    new_cols = len(col_intervals)
    table.row_count = len(row_bands)
    table.col_count = new_cols
    table.cells = [[""] * new_cols for _ in row_bands]
    table.cell_bboxes = [
        [(float(x0), y0, float(x1), y1) for x0, x1 in col_intervals]
        for y0, y1 in row_bands
    ]
    setattr(table, "disable_vertical_merge", True)
    logger.log(
        f"Regridded page-local table on page {table.page}: "
        f"{old_shape[0]}x{old_shape[1]} -> {table.row_count}x{table.col_count}"
    )
    return True


def stabilize_page_local_table_grids(table_regions: List[TableRegion],
                                     layout_regions_by_page: Dict[int, List[dict]],
                                     page_info: dict,
                                     logger: Logger) -> List[TableRegion]:
    """
    Normalize collapsed continuation-page grids without merging pages.

    A table may span several PDF pages visually while each page still needs its
    own DOCX table. Structure engines sometimes collapse a narrow serial-number
    or score column on one page even though adjacent pages with the same ruled
    border expose the full grid. Use nearby page-local tables as schema
    references, but only when the horizontal border is strongly consistent.
    """
    ordered = [
        table for table in sorted(table_regions, key=lambda t: (t.page, t.y_top))
        if not getattr(table, "skip_render", False)
    ]
    fixed = 0

    for table in ordered:
        cols = int(getattr(table, "col_count", 0) or 0)
        if cols <= 0 or cols >= 4:
            continue
        if not _table_is_full_page_segment(table, page_info):
            continue

        table_bbox = _matching_layout_bbox_for_table(table, layout_regions_by_page) or _table_bbox(table)
        best = None
        best_score = float("inf")

        for candidate in ordered:
            if candidate is table or not _table_is_full_page_segment(candidate, page_info):
                continue
            cand_cols = int(getattr(candidate, "col_count", 0) or 0)
            if cand_cols <= cols or cand_cols > 8:
                continue
            page_delta = abs(int(candidate.page) - int(table.page))
            if page_delta <= 0 or page_delta > 3:
                continue

            cand_intervals = _table_column_intervals(candidate)
            if len(cand_intervals) != cand_cols:
                continue
            cand_bbox = _matching_layout_bbox_for_table(candidate, layout_regions_by_page) or _table_bbox(candidate)
            if _horizontal_overlap_ratio(table_bbox, cand_bbox) < 0.82:
                continue

            edge_delta = abs(table_bbox[0] - cand_bbox[0]) + abs(table_bbox[2] - cand_bbox[2])
            if edge_delta > max(48.0, (table_bbox[2] - table_bbox[0]) * 0.14):
                continue
            score = edge_delta + page_delta * 4.0
            if score < best_score:
                best_score = score
                best = cand_intervals

        if best and _regrid_table_columns(table, best, logger):
            fixed += 1

    if fixed:
        logger.log(f"Stabilized {fixed} page-local table grid(s) from adjacent page schemas")
    return table_regions


def trim_empty_trailing_table_columns(table_regions: List[TableRegion], logger: Logger) -> List[TableRegion]:
    """
    Remove detector-only right-edge columns that contain no text anywhere.

    This is a geometry/schema cleanup, not a document-specific rule. Structure
    recognizers sometimes create a narrow extra column at the right border when
    a merged header touches the table edge. Keeping that empty column prevents
    continued-page schema matching, so drop only columns that are empty in every
    row and only from the trailing edge.
    """
    trimmed = 0
    for table in table_regions or []:
        if getattr(table, "skip_render", False):
            continue
        cols = int(getattr(table, "col_count", 0) or 0)
        rows = [list(row) for row in (getattr(table, "cells", []) or [])]
        boxes = [list(row) for row in (getattr(table, "cell_bboxes", []) or [])]
        if cols <= 1 or not rows:
            continue

        remove_count = 0
        while cols - remove_count > 1:
            col_idx = cols - remove_count - 1
            has_text = any(
                col_idx < len(row) and _clean_extracted_text(str(row[col_idx]))
                for row in rows
            )
            if has_text:
                break
            remove_count += 1

        if not remove_count:
            continue

        new_cols = cols - remove_count
        table.cells = [row[:new_cols] for row in rows]
        table.cell_bboxes = [row[:new_cols] for row in boxes]
        table.col_count = new_cols
        trimmed += remove_count

    if trimmed:
        logger.log(f"Trimmed {trimmed} empty trailing table column(s)")
    return table_regions


def _continued_table_geometry_signal(prev: TableRegion,
                                     nxt: TableRegion,
                                     layout_regions_by_page: Dict[int, List[dict]],
                                     page_info: dict,
                                     pdf_lines: List[TextLine]) -> bool:
    if getattr(prev, "skip_render", False) or getattr(nxt, "skip_render", False):
        return False
    if int(nxt.page) != int(prev.page) + 1:
        return False
    if not _table_starts_near_page_top(nxt, page_info):
        return False
    if not _table_ends_near_page_bottom(prev, page_info):
        return False
    if _table_header_signal(nxt, pdf_lines):
        return False

    prev_bbox = _matching_layout_bbox_for_table(prev, layout_regions_by_page) or _table_bbox(prev)
    next_bbox = _matching_layout_bbox_for_table(nxt, layout_regions_by_page) or _table_bbox(nxt)
    if _horizontal_overlap_ratio(prev_bbox, next_bbox) < 0.82:
        return False
    edge_delta = abs(prev_bbox[0] - next_bbox[0]) + abs(prev_bbox[2] - next_bbox[2])
    return edge_delta <= max(48.0, (prev_bbox[2] - prev_bbox[0]) * 0.14)


def stabilize_continuation_table_schemas_from_geometry(
    table_regions: List[TableRegion],
    layout_regions_by_page: Dict[int, List[dict]],
    page_info: dict,
    pdf_lines: List[TextLine],
    logger: Logger,
) -> List[TableRegion]:
    """
    Expand continuation-page segments to the previous page's ruled schema.

    Real multi-page tables often lose a narrow index/notes column on pages that
    start mid-row or have no repeated header. The continuation decision is made
    from page geometry and table borders; then OCR text is assigned again to the
    normalized grid by cell area, so repeated text in different cells is kept.
    """
    ordered = [
        table for table in sorted(table_regions or [], key=lambda t: (int(t.page), float(t.y_top)))
        if not getattr(table, "skip_render", False)
    ]
    fixed = 0
    prev: Optional[TableRegion] = None
    for table in ordered:
        if prev is not None and _continued_table_geometry_signal(
            prev, table, layout_regions_by_page, page_info, pdf_lines
        ):
            prev_cols = int(getattr(prev, "col_count", 0) or 0)
            cols = int(getattr(table, "col_count", 0) or 0)
            if 3 <= cols < prev_cols <= 8:
                ratios = _canonical_column_ratios([prev], layout_regions_by_page)
                if len(ratios) == prev_cols:
                    intervals = _intervals_from_ratios(table, ratios, layout_regions_by_page)
                    if len(intervals) == prev_cols and _regrid_table_columns(table, intervals, logger):
                        fixed += 1
        prev = table

    if fixed:
        logger.log(f"Stabilized {fixed} continuation table schema(s) from previous page geometry")
    return table_regions


def _line_column_index(line: TextLine, col_intervals: List[Tuple[float, float]]) -> Optional[int]:
    if not col_intervals:
        return None
    lx0, _, lx1, _ = _line_bbox_for_assignment(line)
    line_w = max(lx1 - lx0, 1e-6)
    center_x = (lx0 + lx1) / 2.0
    best_idx = None
    best_score = 0.0
    for idx, (x0, x1) in enumerate(col_intervals):
        overlap = max(0.0, min(lx1, x1) - max(lx0, x0))
        score = overlap / line_w
        if score > best_score:
            best_score = score
            best_idx = idx
        if x0 <= center_x <= x1 and best_score < 0.15:
            best_idx = idx
            best_score = 0.15
    return best_idx if best_score >= 0.10 else None


def _table_lines_in_y_band(table: TableRegion,
                           pdf_lines: List[TextLine],
                           y0: float,
                           y1: float,
                           col_intervals: List[Tuple[float, float]]) -> List[TextLine]:
    if y1 <= y0 or not col_intervals:
        return []
    x0 = min(x for x, _ in col_intervals)
    x1 = max(x for _, x in col_intervals)
    out = []
    for line in pdf_lines:
        if int(line.page) != int(table.page):
            continue
        lx0, ly0, lx1, ly1 = _line_bbox_for_assignment(line)
        cy = (ly0 + ly1) / 2.0
        if cy < y0 or cy > y1:
            continue
        if lx1 < x0 - 4.0 or lx0 > x1 + 4.0:
            continue
        if not _clean_extracted_text(line.text):
            continue
        out.append(line)
    return out


def _row_is_right_only_fragment(row: List[str]) -> bool:
    if len(row) < 3:
        return False
    first = _clean_extracted_text(str(row[0]))
    second = _clean_extracted_text(str(row[1]))
    last = _clean_extracted_text(str(row[-1]))
    middle = [_clean_extracted_text(str(cell)) for cell in row[2:-1]]
    return not first and not second and not last and any(middle)


def _gap_has_row_anchor(lines: List[TextLine], col_intervals: List[Tuple[float, float]]) -> bool:
    if len(lines) < 2:
        return False
    has_leading_marker = False
    has_left_text = False
    has_score = False
    has_body_text = False
    last_col = len(col_intervals) - 1
    for line in lines:
        text = _clean_extracted_text(line.text)
        if not text:
            continue
        col_idx = _line_column_index(line, col_intervals)
        if col_idx is None:
            continue
        if col_idx == 0 and _looks_like_table_index_marker(text):
            has_leading_marker = True
        elif col_idx == 1 and len(text.split()) >= 3:
            has_left_text = True
        elif col_idx == last_col and is_numeric_cell(text):
            has_score = True
        elif 1 < col_idx < last_col and len(text.split()) >= 2:
            has_body_text = True
    return (has_leading_marker or has_left_text) and (has_score or has_body_text)


def _row_boxes_from_band(col_intervals: List[Tuple[float, float]],
                         y0: float,
                         y1: float) -> List[Tuple[float, float, float, float]]:
    return [(float(x0), float(y0), float(x1), float(y1)) for x0, x1 in col_intervals]


def repair_table_row_gaps_from_ocr(table_regions: List[TableRegion],
                                   pdf_lines: List[TextLine],
                                   logger: Logger) -> List[TableRegion]:
    """
    Repair row bands when the structure recognizer leaves a ruled-row gap.

    The detector can miss the middle of a tall row in scanned tables. If OCR
    lines inside that vertical gap contain row anchors (index/left descriptor or
    score), use the text geometry to restore the row band before cell assignment.
    """
    repaired = 0
    inserted = 0
    for table in table_regions:
        if getattr(table, "skip_render", False) or int(getattr(table, "col_count", 0) or 0) < 3:
            continue
        col_intervals = _table_column_intervals(table)
        rows = [list(row) for row in (getattr(table, "cells", []) or [])]
        boxes = [list(row) for row in (getattr(table, "cell_bboxes", []) or [])]
        if len(col_intervals) < int(getattr(table, "col_count", 0) or 0) or not rows or len(rows) != len(boxes):
            continue

        out_rows: List[List[str]] = []
        out_boxes: List[List[Tuple[float, float, float, float]]] = []
        r = 0
        changed = False
        while r < len(rows):
            row = rows[r]
            row_boxes = boxes[r]
            y0, y1 = _row_y_bounds(row_boxes)
            if r + 1 >= len(rows) or y1 <= y0:
                out_rows.append(row)
                out_boxes.append(row_boxes)
                r += 1
                continue

            next_y0, next_y1 = _row_y_bounds(boxes[r + 1])
            gap = next_y0 - y1
            gap_lines = _table_lines_in_y_band(table, pdf_lines, y1, next_y0, col_intervals) if gap >= 18.0 else []
            if gap_lines and _gap_has_row_anchor(gap_lines, col_intervals):
                current_fragment = _row_is_right_only_fragment(row)
                next_fragment = _row_is_right_only_fragment(rows[r + 1])
                if current_fragment:
                    new_bottom = next_y1 if next_fragment and next_y1 > next_y0 else next_y0
                    out_rows.append(row)
                    out_boxes.append(_row_boxes_from_band(col_intervals, y0, new_bottom))
                    r += 2 if next_fragment else 1
                    repaired += 1
                    changed = True
                    continue
                if next_fragment:
                    out_rows.append(row)
                    out_boxes.append(row_boxes)
                    out_rows.append([""] * int(getattr(table, "col_count", 0) or 0))
                    out_boxes.append(_row_boxes_from_band(col_intervals, y1, next_y1))
                    r += 2
                    inserted += 1
                    changed = True
                    continue

                out_rows.append(row)
                out_boxes.append(row_boxes)
                out_rows.append([""] * int(getattr(table, "col_count", 0) or 0))
                out_boxes.append(_row_boxes_from_band(col_intervals, y1, next_y0))
                r += 1
                inserted += 1
                changed = True
                continue

            out_rows.append(row)
            out_boxes.append(row_boxes)
            r += 1

        if changed:
            table.cells = out_rows
            table.cell_bboxes = out_boxes
            table.row_count = len(out_rows)

    if repaired:
        logger.log(f"Repaired {repaired} OCR-anchored table row gap(s)")
    if inserted:
        logger.log(f"Inserted {inserted} OCR-anchored missing table row(s)")
    return table_regions


def _looks_like_table_index_marker(text: str) -> bool:
    cleaned = _clean_extracted_text(text)
    if not cleaned:
        return False
    key = _norm_table_key(cleaned)
    return bool(re.fullmatch(r"\d{1,3}|[IVXLCDM]{1,6}", key or ""))


def repair_shifted_leading_table_cells(table_regions: List[TableRegion], logger: Logger) -> List[TableRegion]:
    """
    Move long prose out of the narrow leading index column.

    OCR/table engines occasionally swap the first two columns on continuation
    rows with vertical spans: the prose goes into "Số TT" and the serial number
    goes into the content column. Detect this by content shape, not document
    text, so the fix stays generic.
    """
    fixes = 0
    for table in table_regions:
        if getattr(table, "skip_render", False) or int(getattr(table, "col_count", 0) or 0) < 4:
            continue
        rows = getattr(table, "cells", []) or []
        cols = int(getattr(table, "col_count", 0) or 0)
        for r_idx, row in enumerate(rows):
            normalized = list(row[:cols]) + [""] * max(0, cols - len(row))
            first = _clean_extracted_text(str(normalized[0]))
            second = _clean_extracted_text(str(normalized[1]))
            if not first or len(first) < 16 or _looks_like_table_index_marker(first):
                continue
            first_words = [w for w in re.split(r"\s+", first) if w]
            if len(first_words) < 3:
                continue
            if _looks_like_table_index_marker(second):
                normalized[0], normalized[1] = normalized[1], normalized[0]
                rows[r_idx] = normalized
                fixes += 1
            elif not second:
                normalized[1] = normalized[0]
                normalized[0] = ""
                rows[r_idx] = normalized
                fixes += 1

    if fixes:
        logger.log(f"Repaired {fixes} shifted leading table cell(s)")
    return table_regions


def _merge_cell_bbox(a: Tuple[float, float, float, float],
                     b: Tuple[float, float, float, float]) -> Tuple[float, float, float, float]:
    if not any(a):
        return tuple(float(v) for v in b)
    if not any(b):
        return tuple(float(v) for v in a)
    return (
        min(float(a[0]), float(b[0])),
        min(float(a[1]), float(b[1])),
        max(float(a[2]), float(b[2])),
        max(float(a[3]), float(b[3])),
    )


def _append_continuation_text(dst: str, src: str) -> str:
    dst = _clean_extracted_text(dst or "")
    src = _clean_extracted_text(src or "")
    if not src:
        return dst
    if not dst:
        return src
    if src in dst:
        return dst
    return dst + "\n" + src


def _row_leading_fragment_matches_anchor(fragment: str, anchor: str) -> bool:
    fragment_key = _norm_table_key(fragment)
    anchor_key = _norm_table_key(anchor)
    if not fragment_key or not anchor_key or len(fragment_key) > 12:
        return False
    return fragment_key in anchor_key and fragment_key != anchor_key


def _row_has_left_anchor(row: List[str]) -> bool:
    if not row:
        return False
    first = _clean_extracted_text(str(row[0])) if len(row) > 0 else ""
    second = _clean_extracted_text(str(row[1])) if len(row) > 1 else ""
    if _looks_like_table_index_marker(first) or _looks_like_table_index_marker(second):
        return True
    for value in (first, second):
        if not value or _starts_like_continuation_tail(value):
            continue
        words = [w for w in re.split(r"\s+", value) if w]
        if 1 <= len(words) <= 8 and len(value) <= 80:
            return True
    return False


def _row_middle_starts_like_continuation(row: List[str], cols: int) -> bool:
    for c in range(1, max(1, cols - 1)):
        if c >= len(row):
            continue
        text = _clean_extracted_text(str(row[c]))
        if text:
            return _starts_like_continuation_tail(text)
    return False


def fuse_page_local_continuation_rows(table_regions: List[TableRegion], logger: Logger) -> List[TableRegion]:
    """
    Merge detector-created continuation rows back into the logical row.

    This handles scanned rows that are cut by a physical page/long-cell band:
    the continuation row has no index/score, but carries more text for the
    middle descriptive columns. It keeps table segments page-local and changes
    only row structure inside each detected segment.
    """
    fused = 0
    suppressed = 0

    for table in table_regions:
        if getattr(table, "skip_render", False) or int(getattr(table, "col_count", 0) or 0) < 4:
            continue
        cols = int(getattr(table, "col_count", 0) or 0)
        rows = [list(row[:cols]) + [""] * max(0, cols - len(row)) for row in (getattr(table, "cells", []) or [])]
        boxes = [
            list(row[:cols]) + [(0.0, 0.0, 0.0, 0.0)] * max(0, cols - len(row))
            for row in (getattr(table, "cell_bboxes", []) or [])
        ]
        if not rows or len(boxes) != len(rows):
            continue

        out_rows: List[List[str]] = []
        out_boxes: List[List[Tuple[float, float, float, float]]] = []
        for row, row_boxes in zip(rows, boxes):
            first = _clean_extracted_text(str(row[0]))
            score = _clean_extracted_text(str(row[-1]))
            middle_has_text = any(_clean_extracted_text(str(cell)) for cell in row[1:-1])
            previous_anchor = None
            for previous in reversed(out_rows):
                prev_first = _clean_extracted_text(str(previous[0]))
                prev_second = _clean_extracted_text(str(previous[1]))
                if prev_first or prev_second:
                    previous_anchor = previous
                    break
            previous_row = out_rows[-1] if out_rows else None
            previous_first = _clean_extracted_text(str(previous_row[0])) if previous_row else ""
            previous_second = _clean_extracted_text(str(previous_row[1])) if previous_row and len(previous_row) > 1 else ""
            leading_matches_previous = (
                bool(first)
                and (
                    _row_leading_fragment_matches_anchor(first, previous_first)
                    or _row_leading_fragment_matches_anchor(first, previous_second)
                )
            )
            current_has_left_anchor = _row_has_left_anchor(row)
            is_continuation = (
                not score
                and middle_has_text
                and bool(previous_row)
                and _row_has_left_anchor(previous_row)
                and not current_has_left_anchor
                and (
                    not first
                    or (
                        leading_matches_previous
                        and _row_middle_starts_like_continuation(row, cols)
                    )
                )
            )
            if is_continuation:
                target = out_rows[-1]
                target_boxes = out_boxes[-1]
                for c in range(1, cols - 1):
                    target[c] = _append_continuation_text(str(target[c]), str(row[c]))
                    target_boxes[c] = _merge_cell_bbox(tuple(target_boxes[c]), tuple(row_boxes[c]))
                fused += 1
                continue

            if previous_anchor is not None:
                if (
                    _clean_extracted_text(str(row[0]))
                    and _clean_extracted_text(str(row[0])) == _clean_extracted_text(str(previous_anchor[0]))
                    and _clean_extracted_text(str(row[1]))
                    and _clean_extracted_text(str(row[1])) == _clean_extracted_text(str(previous_anchor[1]))
                ):
                    row[0] = ""
                    row[1] = ""
                    suppressed += 1

            out_rows.append(row)
            out_boxes.append([tuple(float(v) for v in bx[:4]) for bx in row_boxes])

        if len(out_rows) != len(rows) or suppressed:
            table.cells = out_rows
            table.cell_bboxes = out_boxes
            table.row_count = len(out_rows)

    if fused:
        logger.log(f"Fused {fused} page-local continuation table row(s)")
    if suppressed:
        logger.log(f"Suppressed {suppressed} repeated leading table cell pair(s)")
    return table_regions


_ROMAN_DIGITS = [
    (1000, "M"),
    (900, "CM"),
    (500, "D"),
    (400, "CD"),
    (100, "C"),
    (90, "XC"),
    (50, "L"),
    (40, "XL"),
    (10, "X"),
    (9, "IX"),
    (5, "V"),
    (4, "IV"),
    (1, "I"),
]


def _roman_to_int(text: str) -> Optional[int]:
    key = _norm_table_key(text)
    if not key or not re.fullmatch(r"[IVXLCDM]{1,8}", key):
        return None
    total = 0
    prev = 0
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    for ch in reversed(key):
        value = values.get(ch, 0)
        total += -value if value < prev else value
        prev = max(prev, value)
    return total if 0 < total < 100 else None


def _int_to_roman(value: int) -> str:
    if value <= 0:
        return ""
    out = []
    remaining = value
    for number, roman in _ROMAN_DIGITS:
        while remaining >= number:
            out.append(roman)
            remaining -= number
    return "".join(out)


def _uppercase_letter_ratio(text: str) -> float:
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return 0.0
    return sum(1 for ch in letters if ch.isupper()) / len(letters)


def _looks_like_section_heading_text(text: str) -> bool:
    text = _clean_extracted_text(text)
    if not text or _is_list_item_line(text):
        return False
    if len(text.split()) < 3 or len(text) > 170:
        return False
    if re.search(r"[.;:]", text):
        return False
    key = _norm_table_key(text)
    return key.startswith("CONGTAC") or key.startswith("TONG") or _uppercase_letter_ratio(text) >= 0.78


def _section_heading_middle_text(row: List[str], cols: int) -> str:
    if cols < 3:
        return ""
    score = _clean_extracted_text(str(row[-1])) if row else ""
    if not is_numeric_cell(score):
        return ""
    middle_values = [
        _clean_extracted_text(str(row[c]))
        for c in range(1, cols - 1)
        if c < len(row) and _clean_extracted_text(str(row[c]))
    ]
    if not middle_values:
        return ""
    deduped = []
    for value in middle_values:
        if deduped and _norm_table_key(deduped[-1]) == _norm_table_key(value):
            continue
        deduped.append(value)
    combined = _clean_extracted_text(" ".join(deduped))
    return combined if _looks_like_section_heading_text(combined) else ""


def repair_section_heading_rows(table_regions: List[TableRegion], logger: Logger) -> List[TableRegion]:
    """
    Normalize all-caps section rows that span the middle columns.

    These rows are visual section headers, not data rows with independent middle
    cells. OCR/table structure often splits one centered title across adjacent
    columns. Store that as explicit render-span metadata instead of duplicating
    the title into every covered cell.
    """
    refs = []
    normalized = 0
    for table in sorted(table_regions, key=lambda t: (int(t.page), float(t.y_top))):
        if getattr(table, "skip_render", False):
            continue
        cols = int(getattr(table, "col_count", 0) or 0)
        if cols < 3:
            continue
        rows = [list(row[:cols]) + [""] * max(0, cols - len(row)) for row in (getattr(table, "cells", []) or [])]
        heading_spans = dict(getattr(table, "horizontal_text_spans", {}) or {})
        changed = False
        for r_idx, row in enumerate(rows):
            combined = _section_heading_middle_text(row, cols)
            if not combined:
                continue
            row[1] = combined
            for c in range(2, cols - 1):
                row[c] = ""
            rows[r_idx] = row
            heading_spans[r_idx] = (1, cols - 2, combined)
            normalized += 1
            changed = True
            refs.append({
                "table": table,
                "row": r_idx,
                "marker": _clean_extracted_text(str(row[0])),
            })
        if changed:
            table.cells = rows
            setattr(table, "horizontal_text_spans", heading_spans)

    inferred = 0
    i = 0
    while i < len(refs):
        marker_value = _roman_to_int(refs[i]["marker"])
        if marker_value is not None:
            i += 1
            continue
        run_start = i
        while i < len(refs) and _roman_to_int(refs[i]["marker"]) is None:
            i += 1
        run_end = i
        blanks = refs[run_start:run_end]
        prev_value = None
        next_value = None
        for j in range(run_start - 1, -1, -1):
            prev_value = _roman_to_int(refs[j]["marker"])
            if prev_value is not None:
                break
        for j in range(run_end, len(refs)):
            next_value = _roman_to_int(refs[j]["marker"])
            if next_value is not None:
                break

        start_value = prev_value or 0
        can_infer = False
        if next_value is not None and next_value - start_value == len(blanks) + 1:
            can_infer = True
        elif prev_value is None and next_value is not None and next_value > len(blanks):
            start_value = next_value - len(blanks) - 1
            can_infer = True

        if not can_infer:
            continue
        for offset, ref in enumerate(blanks, start=1):
            value = start_value + offset
            marker = _int_to_roman(value)
            if not marker:
                continue
            table = ref["table"]
            rows = getattr(table, "cells", []) or []
            if ref["row"] < len(rows) and rows[ref["row"]]:
                rows[ref["row"]][0] = marker
                ref["marker"] = marker
                inferred += 1

    if normalized:
        logger.log(f"Normalized {normalized} section heading row span(s)")
    if inferred:
        logger.log(f"Inferred {inferred} missing section marker(s)")
    return table_regions


def _table_starts_near_page_top(table: TableRegion, page_info: dict) -> bool:
    page_height = float(page_info.get(table.page, {}).get("height", 842) or 842)
    return float(table.y_top) <= page_height * 0.16


def _table_ends_near_page_bottom(table: TableRegion, page_info: dict) -> bool:
    page_height = float(page_info.get(table.page, {}).get("height", 842) or 842)
    return float(table.y_bottom) >= page_height * 0.84


def _table_header_signal(table: TableRegion, pdf_lines: List[TextLine]) -> bool:
    if _table_has_header(table):
        return True
    bbox = _table_bbox(table)
    lines_in_region = [
        line for line in pdf_lines
        if line.page == table.page and _line_in_bbox(line, bbox)
    ]
    return _looks_like_table_header_text(" ".join(line.text for line in lines_in_region[:20]))


def _continued_table_reference_intervals(
    tables: List[TableRegion],
    col_count: int,
) -> List[Tuple[float, float]]:
    candidates: List[List[Tuple[float, float]]] = []
    for table in tables:
        intervals = _table_column_intervals(table)
        if len(intervals) == col_count:
            candidates.append(intervals)
    if not candidates:
        return []
    return max(candidates, key=lambda intervals: intervals[-1][1] - intervals[0][0])


def _table_column_widths_from_intervals(
    table: TableRegion,
    intervals: List[Tuple[float, float]],
    layout_regions_by_page: Optional[Dict[int, List[dict]]] = None,
) -> List[float]:
    col_count = len(intervals)
    if col_count <= 0:
        return []
    table_bbox = (
        _matching_layout_bbox_for_table(table, layout_regions_by_page or {})
        or _table_bbox(table)
    )
    boundaries = [float(table_bbox[0])]
    for col_idx in range(col_count - 1):
        boundaries.append((float(intervals[col_idx][1]) + float(intervals[col_idx + 1][0])) / 2.0)
    boundaries.append(float(table_bbox[2]))
    widths = [max(0.0, boundaries[i + 1] - boundaries[i]) for i in range(col_count)]
    if any(width <= 1.0 for width in widths):
        return []
    return widths


def _canonical_column_ratios(
    tables: List[TableRegion],
    layout_regions_by_page: Dict[int, List[dict]],
) -> List[float]:
    if not tables:
        return []
    col_count = int(getattr(tables[0], "col_count", 0) or 0)
    ratio_sets: List[List[float]] = []
    for table in tables:
        if int(getattr(table, "col_count", 0) or 0) != col_count:
            continue
        intervals = _table_column_intervals(table)
        if len(intervals) != col_count:
            continue
        widths = _table_column_widths_from_intervals(table, intervals, layout_regions_by_page)
        total = sum(widths)
        if total <= 1.0:
            continue
        ratio_sets.append([width / total for width in widths])
    if not ratio_sets:
        return []
    ratios = [_median([ratio_set[c] for ratio_set in ratio_sets]) for c in range(col_count)]
    total = sum(ratios)
    if total <= 0:
        return []
    return [ratio / total for ratio in ratios]


def _intervals_from_ratios(
    table: TableRegion,
    ratios: List[float],
    layout_regions_by_page: Dict[int, List[dict]],
) -> List[Tuple[float, float]]:
    table_bbox = _matching_layout_bbox_for_table(table, layout_regions_by_page) or _table_bbox(table)
    x = float(table_bbox[0])
    right = float(table_bbox[2])
    total_width = max(right - x, 1.0)
    intervals: List[Tuple[float, float]] = []
    for idx, ratio in enumerate(ratios):
        if idx == len(ratios) - 1:
            next_x = right
        else:
            next_x = x + total_width * max(float(ratio), 0.0)
        intervals.append((x, max(next_x, x + 1.0)))
        x = next_x
    return intervals


def _normalize_table_rows_to_intervals(
    table: TableRegion,
    col_intervals: List[Tuple[float, float]],
) -> Tuple[List[List[str]], List[List[Tuple[float, float, float, float]]]]:
    col_count = len(col_intervals)
    rows = [
        list(row[:col_count]) + [""] * max(0, col_count - len(row))
        for row in (getattr(table, "cells", []) or [])
    ]
    out_boxes: List[List[Tuple[float, float, float, float]]] = []
    for r_idx, row in enumerate(rows):
        source_boxes = (
            table.cell_bboxes[r_idx]
            if r_idx < len(getattr(table, "cell_bboxes", []) or [])
            else []
        )
        y0, y1 = _row_y_bounds(source_boxes)
        if y1 <= y0:
            row_height = max((float(table.y_bottom) - float(table.y_top)) / max(len(rows), 1), 12.0)
            y0 = float(table.y_top) + r_idx * row_height
            y1 = y0 + row_height
        out_boxes.append([(x0, y0, x1, y1) for x0, x1 in col_intervals])
    return rows, out_boxes


def _continued_table_groups(table_regions: List[TableRegion],
                            layout_regions_by_page: Dict[int, List[dict]],
                            page_info: dict,
                            pdf_lines: List[TextLine]) -> List[List[TableRegion]]:
    ordered = sorted(
        [table for table in table_regions if not getattr(table, "skip_render", False)],
        key=lambda table: (table.page, table.y_top),
    )
    groups: List[List[TableRegion]] = []
    current: List[TableRegion] = []
    for table in ordered:
        if (
            current
            and _continued_tables_are_compatible(
                current[-1], table, layout_regions_by_page, page_info, pdf_lines
            )
        ):
            current.append(table)
            continue
        if current:
            groups.append(current)
        current = [table]
    if current:
        groups.append(current)
    return groups


def _continued_tables_are_compatible(prev: TableRegion,
                                     nxt: TableRegion,
                                     layout_regions_by_page: Dict[int, List[dict]],
                                     page_info: dict,
                                     pdf_lines: List[TextLine]) -> bool:
    if getattr(prev, "skip_render", False) or getattr(nxt, "skip_render", False):
        return False
    if int(nxt.page) != int(prev.page) + 1:
        return False
    if int(getattr(prev, "col_count", 0) or 0) != int(getattr(nxt, "col_count", 0) or 0):
        return False
    if int(getattr(prev, "col_count", 0) or 0) < 3:
        return False
    if not _table_starts_near_page_top(nxt, page_info):
        return False
    if not _table_ends_near_page_bottom(prev, page_info):
        return False

    prev_bbox = _matching_layout_bbox_for_table(prev, layout_regions_by_page) or _table_bbox(prev)
    next_bbox = _matching_layout_bbox_for_table(nxt, layout_regions_by_page) or _table_bbox(nxt)
    if _horizontal_overlap_ratio(prev_bbox, next_bbox) < 0.86:
        return False
    edge_delta = abs(prev_bbox[0] - next_bbox[0]) + abs(prev_bbox[2] - next_bbox[2])
    if edge_delta > max(36.0, (prev_bbox[2] - prev_bbox[0]) * 0.10):
        return False

    # A repeated header is a strong signal that the next page should stay as a
    # separate table. Continuation pages in scanned forms usually start with a
    # body row or a blank-leading row inside an existing rowspan.
    if _table_header_signal(nxt, pdf_lines):
        return False
    return True


def stabilize_continued_table_column_schemas(table_regions: List[TableRegion],
                                             layout_regions_by_page: Dict[int, List[dict]],
                                             page_info: dict,
                                             pdf_lines: List[TextLine],
                                             logger: Logger) -> List[TableRegion]:
    """
    Keep page-local continuation segments on the same DOCX column schema.

    Structure recognizers may choose slightly different x boundaries on each
    physical page. Rendering each segment with its own tblGrid makes vertical
    borders appear to jump even when the PDF ruled table is continuous. Use the
    median column-width ratio across compatible continuation pages, then map it
    back to each page's own layout bbox so OCR geometry still stays local.
    """
    fixed = 0
    for group in _continued_table_groups(table_regions, layout_regions_by_page, page_info, pdf_lines):
        if len(group) <= 1:
            continue
        ratios = _canonical_column_ratios(group, layout_regions_by_page)
        if len(ratios) != int(getattr(group[0], "col_count", 0) or 0):
            continue
        for table in group:
            intervals = _intervals_from_ratios(table, ratios, layout_regions_by_page)
            rows, boxes = _normalize_table_rows_to_intervals(table, intervals)
            if not rows or not boxes:
                continue
            table.cells = rows
            table.cell_bboxes = boxes
            table.row_count = len(rows)
            table.col_count = len(intervals)
            table_bbox = _matching_layout_bbox_for_table(table, layout_regions_by_page) or _table_bbox(table)
            setattr(table, "x_left", float(table_bbox[0]))
            setattr(table, "x_right", float(table_bbox[2]))
            fixed += 1
    if fixed:
        logger.log(f"Stabilized {fixed} continued table segment column schema(s)")
    return table_regions


def compose_continued_tables_for_word_flow(table_regions: List[TableRegion],
                                           layout_regions_by_page: Dict[int, List[dict]],
                                           page_info: dict,
                                           pdf_lines: List[TextLine],
                                           logger: Logger) -> List[TableRegion]:
    """
    Build one DOCX table for physical pages that are clearly one ruled table.

    The repair is intentionally after page-local structure/OCR cleanup. We do
    not rebuild later pages from the first page's grid; each page keeps its
    detected rows, then a compatible column schema is applied so Word can flow
    the same logical table across pages and vertical rowspans can continue.
    """
    synthetic_tables: List[TableRegion] = []
    merged_groups = 0
    for group in _continued_table_groups(table_regions, layout_regions_by_page, page_info, pdf_lines):
        if len(group) <= 1:
            continue
        col_count = int(getattr(group[0], "col_count", 0) or 0)
        col_intervals = _continued_table_reference_intervals(group, col_count)
        if len(col_intervals) != col_count:
            continue

        combined_rows: List[List[str]] = []
        combined_boxes: List[List[Tuple[float, float, float, float]]] = []
        row_source_pages: List[int] = []
        for source in group:
            rows, boxes = _normalize_table_rows_to_intervals(source, col_intervals)
            combined_rows.extend(rows)
            combined_boxes.extend(boxes)
            row_source_pages.extend([int(source.page)] * len(rows))

        if not combined_rows:
            continue

        first = group[0]
        layout_bbox = _matching_layout_bbox_for_table(first, layout_regions_by_page) or _table_bbox(first)
        synthetic = TableRegion(
            page=first.page,
            y_top=first.y_top,
            y_bottom=first.y_bottom,
            cells=combined_rows,
            row_count=len(combined_rows),
            col_count=col_count,
            cell_bboxes=combined_boxes,
        )
        setattr(synthetic, "x_left", float(layout_bbox[0]))
        setattr(synthetic, "x_right", float(layout_bbox[2]))
        setattr(synthetic, "source_segments", list(group))
        setattr(synthetic, "source_pages", [int(table.page) for table in group])
        setattr(synthetic, "row_source_pages", row_source_pages)
        setattr(synthetic, "enable_blank_continuation_vmerge", True)
        setattr(synthetic, "disable_vertical_merge", True)
        for source in group:
            setattr(source, "skip_render", True)
        synthetic_tables.append(synthetic)
        merged_groups += 1

    if merged_groups:
        logger.log(f"Composed {merged_groups} continued table group(s) for Word flow")
    return table_regions + synthetic_tables


def filter_false_positive_tables(table_regions: List[TableRegion],
                                 layout_regions_by_page: Dict[int, List[dict]],
                                 logger: Logger) -> List[TableRegion]:
    """
    Remove detector artifacts that are just wrapped text blocks. These are
    especially damaging because the paragraph text disappears into a 1-column
    DOCX table and creates extra tables compared with the ground truth.
    """
    kept: List[TableRegion] = []
    removed = 0
    for table in table_regions:
        if getattr(table, "skip_render", False):
            kept.append(table)
            continue
        if (getattr(table, "col_count", 0) or 0) <= 1 and (getattr(table, "row_count", 0) or 0) > 1:
            removed += 1
            continue
        kept.append(table)

    if removed:
        logger.log(f"Removed {removed} false-positive 1-column table(s)")
    return kept


def _row_y_bounds(row_boxes: List[Tuple[float, float, float, float]]) -> Tuple[float, float]:
    ys0 = [bbox[1] for bbox in row_boxes if any(bbox)]
    ys1 = [bbox[3] for bbox in row_boxes if any(bbox)]
    if not ys0 or not ys1:
        return (0.0, 0.0)
    return (min(ys0), max(ys1))


def _copy_table_slice(table: TableRegion, start: int, end: int) -> Optional[TableRegion]:
    cells = [list(row) for row in table.cells[start:end]]
    boxes = [list(row) for row in table.cell_bboxes[start:end]]
    if not cells:
        return None
    y_top, _ = _row_y_bounds(boxes[0])
    _, y_bottom = _row_y_bounds(boxes[-1])
    if y_top <= 0:
        y_top = table.y_top
    if y_bottom <= 0:
        y_bottom = table.y_bottom
    out = TableRegion(
        page=table.page,
        y_top=y_top,
        y_bottom=y_bottom,
        cells=cells,
        row_count=len(cells),
        col_count=table.col_count,
        cell_bboxes=boxes,
    )
    for attr in ("x_left", "x_right", "disable_vertical_merge"):
        if hasattr(table, attr):
            setattr(out, attr, getattr(table, attr))
    return out


def _row_signature(row: List[str]) -> str:
    parts = []
    for cell in row:
        key = _norm_table_key(cell)
        key = re.sub(r"\d+", "#", key)
        if key:
            parts.append(key)
    return "|".join(parts)


def _row_nonempty_count(row: List[str]) -> int:
    return sum(1 for cell in row if _clean_extracted_text(str(cell)))


def _row_similarity(a: List[str], b: List[str]) -> float:
    from difflib import SequenceMatcher
    return SequenceMatcher(None, _row_signature(a), _row_signature(b)).ratio()


def split_stacked_tables(table_regions: List[TableRegion], logger: Logger) -> List[TableRegion]:
    """
    Split two logical tables that the detector merged into one tall region.
    This happens when two ruled tables with the same grid are stacked on one
    page and a caption row sits between them.
    """
    out: List[TableRegion] = []
    split_count = 0
    for table in table_regions:
        if getattr(table, "skip_render", False) or table.row_count < 9 or table.col_count < 3:
            out.append(table)
            continue

        split_start = None
        for idx in range(3, table.row_count - 2):
            if _row_nonempty_count(table.cells[idx]) < 1 or _row_nonempty_count(table.cells[idx + 1]) < 2:
                continue
            previous_is_caption = _row_nonempty_count(table.cells[idx - 1]) <= 2
            first_header_repeat = _row_similarity(table.cells[idx], table.cells[0])
            second_header_repeat = _row_similarity(table.cells[idx + 1], table.cells[1])
            paired_repeat = (first_header_repeat + second_header_repeat) / 2.0
            if previous_is_caption and paired_repeat >= 0.55:
                split_start = idx
                break

        if split_start is None:
            out.append(table)
            continue

        first_end = split_start
        if split_start > 0 and _row_nonempty_count(table.cells[split_start - 1]) <= 2:
            first_end = split_start - 1

        first = _copy_table_slice(table, 0, first_end)
        second = _copy_table_slice(table, split_start, table.row_count)
        if first is None or second is None:
            out.append(table)
            continue
        out.extend([first, second])
        split_count += 1

    if split_count:
        logger.log(f"Split {split_count} stacked table(s)")
    return out


def detect_tables_img2table(pdf_path: str, logger: Logger, page_info: dict, pdf_lines: List[TextLine]) -> List[TableRegion]:
    """
    Detect tables using img2table by processing the PDF directly.
    Uses pdf_lines to accurately populate cell content.
    """
    if not IMG2TABLE_AVAILABLE:
        logger.log("img2table not available, skipping table detection")
        return []

    tables = []
    try:
        # Use simple import inside function to avoid circular dependency if any
        from img2table.document import PDF
        
        # Initialize PDF object
        # Note: img2table uses 200 DPI by default for PDF conversion
        pdf = PDF(src=pdf_path)
        
        # Check total pixels for safety
        max_pixels_safe = 40_000_000 # 40MP (Safe for 3x Zoom A4, but blocks 66MP)
        should_skip = False
        
        for p_idx in range(len(page_info)):
            p_num = p_idx + 1
            if p_num in page_info:
                w_pt = page_info[p_num].get("width", 0)
                h_pt = page_info[p_num].get("height", 0)
                # Convert to 200 DPI pixels
                w_px = w_pt * (200/72)
                h_px = h_pt * (200/72)
                total_px = w_px * h_px
                
                if total_px > max_pixels_safe:
                    logger.log(f"  img2table: Page {p_num} is too large (~{total_px/1e6:.1f}MP). Skipping img2table to prevent hang.")
                    should_skip = True
                    break
        
        if should_skip:
            return []

        logger.log("  img2table: Starting extraction (this may take time for large files)...")

        # Extract tables with optimized parameters
        # borderless_tables=False: avoid false positives on text-only layouts
        # (e.g. "Noi nhan" + signature blocks detected as borderless table)
        extracted_tables = pdf.extract_tables(
            implicit_rows=False,
            borderless_tables=False,
            min_confidence=50
        )
        
        # We need to calculate scale factors. 
        # img2table converts PDF to images at 200 DPI.
        # Coordinates in extracted_tables are in pixels relative to that 200 DPI image.
        # We need to convert them to PDF Points (72 DPI).
        # Scale Factor = 72 / 200 = 0.36
        
        scale_x = 72.0 / 200.0
        scale_y = 72.0 / 200.0
        
        # Iterate through results
        # extracted_tables is a dict {page_idx: [Table, ...]} where page_idx is 0-based
        for page_idx in sorted(extracted_tables.keys()):
            page_tables = extracted_tables[page_idx]
            page_num = page_idx + 1 # Convert to 1-based for our system
            
            # Get lines for this page once
            pdf_lines_page = [l for l in pdf_lines if l.page == page_num]
            
            for i, t in enumerate(page_tables):
                # t is img2table Table object
                
                # Scale BBox to Points
                # t.bbox is (x1, y1, x2, y2)
                x1_view = t.bbox.x1 * scale_x
                y1_view = t.bbox.y1 * scale_y
                x2_view = t.bbox.x2 * scale_x
                y2_view = t.bbox.y2 * scale_y
                
                # Restore BBox extraction for PDF Verification
                # We need BBoxes to know WHERE the row is.
                
                # Hybrid Approach:
                # 1. Text from DataFrame (Reliable)
                # 2. BBoxes from Content (Required for verification)
                
                # Dimensions
                if t.df is not None:
                    # Prepare DF
                    df_clean = t.df.fillna("")
                    rows, cols = df_clean.shape
                elif t.content:
                    rows = max(t.content.keys()) + 1
                    cols = max(len(row) for row in t.content.values()) if t.content else 0
                else:
                    rows, cols = 0, 0
                
                # Initialize grids
                cells_grid = [[""] * cols for _ in range(rows)]
                vis_cell_bboxes = [[(0.0, 0.0, 0.0, 0.0)] * cols for _ in range(rows)]
                
                try:
                    # 1. Populate BBoxes from Content (First pass needed for geometric text)
                    if t.content:
                        for r_idx, row_obj in t.content.items():
                            if r_idx < rows:
                                if isinstance(row_obj, dict):
                                    iterator = row_obj.items()
                                else:
                                    iterator = enumerate(row_obj)
                                
                                for c_idx, cell_obj in iterator:
                                    if c_idx < cols:
                                        cb = cell_obj.bbox
                                        scaled_bbox = (
                                            cb.x1 * scale_x,
                                            cb.y1 * scale_y,
                                            cb.x2 * scale_x,
                                            cb.y2 * scale_y
                                        )
                                        vis_cell_bboxes[r_idx][c_idx] = scaled_bbox

                    # 2. Populate Text: Prioritize Geometric Extraction with Smart Merging
                    for r in range(rows):
                        for c in range(cols):
                            bbox = vis_cell_bboxes[r][c]
                            if any(bbox) and pdf_lines_page:
                                # Get all lines in this cell
                                cell_lines = get_lines_in_rect(bbox, pdf_lines_page)
                                if cell_lines:
                                    # Merge them using the same logic as paragraphs
                                    # Use cell boundaries for merge logic
                                    # But merge_raw_paragraphs now expects margin_map.
                                    # Create a fake margin map for this single cell context.
                                    # The "Page" of these lines is p_idx+1.
                                    # cell_base_x = bbox[0]
                                    # cell_max_right = bbox[2]
                                    
                                    cell_margin_map = {page_num: (bbox[0], bbox[2])}
                                    
                                    merged_paras = merge_raw_paragraphs(cell_lines, cell_margin_map, logger)
                                    
                                    # Join paragraphs with newline (standard cell behavior)
                                    # But user requested: "Nội dung trong một ô là một paragraph" (Content in a cell is A paragraph).
                                    # "Qua ô khác thì áp dụng thuật toán lại từ đầu." (Move to next cell, apply again).
                                    # If the algorithm yields multiple paragraphs (e.g. explicitly split), we join them with newlines or spaces?
                                    # "Nội dung trong một ô là một paragraph" implies result should be ONE paragraph?
                                    # Or maybe it means "Treat content as paragraphs and merge accordingly".
                                    # Let's join with "\n" if multiple paragraphs are detected.
                                    
                                    full_cell_text = "\n".join([p[0] for p in merged_paras])
                                    geo_text = clean_ocr_cell_text(full_cell_text)

                                    # Cross-validate with DataFrame text:
                                    # If geometric extraction has EXTRA text that DataFrame
                                    # doesn't (e.g. text from adjacent cell bleeding in),
                                    # prefer DataFrame text as it's more reliable
                                    df_text = ""
                                    if t.df is not None and r < len(df_clean) and c < len(df_clean.columns):
                                        df_text = str(df_clean.iloc[r, c]).strip()
                                        while df_text.startswith('|'):
                                            df_text = df_text[1:].strip()

                                    if df_text and geo_text and geo_text != df_text:
                                        # Cross-validate: if DataFrame text is a clean
                                        # substring of geometric text, the extra content
                                        # in geometric is likely bleed from adjacent cell.
                                        # Prefer DataFrame in that case.
                                        # Normalize all whitespace (PDF uses \xa0 non-breaking space)
                                        geo_flat = re.sub(r'[\xa0\s]+', ' ', geo_text).strip()
                                        df_flat = re.sub(r'[\xa0\s]+', ' ', df_text).strip()
                                        if df_flat and df_flat in geo_flat and geo_flat != df_flat:
                                            cells_grid[r][c] = df_text
                                        else:
                                            cells_grid[r][c] = geo_text
                                    else:
                                        cells_grid[r][c] = geo_text
                                    continue

                            # Fallback to img2table DF if geometric failed
                            if t.df is not None and r < len(df_clean) and c < len(df_clean.columns):
                                val = str(df_clean.iloc[r, c]).strip()
                                while val.startswith('|'):
                                    val = val[1:].strip()
                                cells_grid[r][c] = val

                                        
                except Exception as e:
                    logger.log(f"  Warning: Could not extract cell data for p{page_num} t{i}: {e}")

                tables.append(TableRegion(
                    page=page_num,
                    y_top=y1_view,
                    y_bottom=y2_view,
                    cells=cells_grid,
                    row_count=rows,
                    col_count=cols,
                    cell_bboxes=vis_cell_bboxes
                ))
                
                logger.log(f"  img2table (Hybrid Mode): p{page_num} Table {i+1} {rows}x{cols} @ y={y1_view:.1f}-{y2_view:.1f}")

    except Exception as e:
        logger.log(f"img2table error: {e}")
        import traceback
        traceback.print_exc()
        
    return tables


# ============================================================================
# TABLE ENGINE CONFIGURATION
# ============================================================================

def get_table_config() -> dict:
    """
    Read table extraction settings from settings.ini.
    Returns dict with 'engine' and 'device' keys.
    """
    import configparser
    from scanindex.infra.paths import get_base_dir
    config = configparser.ConfigParser()
    config.read(os.path.join(get_base_dir(), "settings.ini"))
    
    return {
        "engine": "hybrid",
        "device": "cpu"
    }



def _mark_table_source(tables: List[TableRegion], source: str) -> List[TableRegion]:
    for table in tables:
        setattr(table, "source", source)
    return tables


def _table_nonempty_ratio(table: TableRegion) -> float:
    cells = [str(cell).strip() for row in getattr(table, "cells", []) for cell in row]
    if not cells:
        return 0.0
    return sum(1 for cell in cells if cell) / len(cells)


def _has_stacked_header_repeat(table: TableRegion) -> bool:
    rows = getattr(table, "cells", []) or []
    if len(rows) < 8 or getattr(table, "col_count", 0) < 3:
        return False
    for idx in range(3, len(rows) - 2):
        if _row_nonempty_count(rows[idx]) < 1 or _row_nonempty_count(rows[idx + 1]) < 2:
            continue
        if _row_nonempty_count(rows[idx - 1]) > 2:
            continue
        paired = (_row_similarity(rows[idx], rows[0]) + _row_similarity(rows[idx + 1], rows[1])) / 2.0
        if paired >= 0.55:
            return True
    return False


def _fragmented_body_row_ratio(table: TableRegion) -> float:
    rows = getattr(table, "cells", []) or []
    cols = getattr(table, "col_count", 0) or 0
    if len(rows) < 6 or cols < 3:
        return 0.0

    start_idx = 1 if _table_has_header(table) else 0
    fragments = 0
    candidates = 0
    for row in rows[start_idx:]:
        row = list(row[:cols]) + [""] * max(0, cols - len(row))
        if _row_nonempty_count(row) == 0:
            continue
        candidates += 1
        leading_empty = (
            not _clean_extracted_text(row[0])
            and (cols < 2 or not _clean_extracted_text(row[1]))
        )
        has_body_text = any(_clean_extracted_text(cell) for cell in row[2:])
        if leading_empty and has_body_text:
            fragments += 1
    if not candidates:
        return 0.0
    return fragments / candidates


def _layout_table_bboxes(layout_regions_by_page: Optional[Dict[int, List[dict]]], page: int) -> List[Tuple[float, float, float, float]]:
    out = []
    for region in (layout_regions_by_page or {}).get(page, []):
        if region.get("type") != "table":
            continue
        bbox = region.get("bbox_pdf")
        if bbox and len(bbox) >= 4:
            out.append(tuple(float(v) for v in bbox[:4]))
    return out


def _candidate_table_score(table: TableRegion,
                           layout_bboxes: List[Tuple[float, float, float, float]]) -> float:
    rows = getattr(table, "row_count", 0) or 0
    cols = getattr(table, "col_count", 0) or 0
    if rows <= 0 or cols <= 0:
        return -100.0

    score = 0.0
    ratio = _table_nonempty_ratio(table)
    bbox = _table_bbox(table)

    # Strongly reject detector artifacts that turn paragraphs into very wide
    # or one-column tables. Real administrative tables in this pipeline have a
    # stable ruled grid and moderate column count.
    if cols == 1 and rows > 1:
        score -= 12.0
    if cols > 14:
        score -= 10.0 + (cols - 14) * 0.8
    elif cols > 10:
        score -= (cols - 10) * 0.7
    if rows <= 2 and cols > 12:
        score -= 8.0

    score += min(rows, 8) * 0.18
    score += min(cols, 8) * 0.35
    score += ratio * 1.2

    if _table_has_header(table):
        score += 1.6
    if _has_stacked_header_repeat(table):
        # A single detector region that contains a repeated header can be split
        # deterministically later. Prefer it over multiple partial detections
        # that may miss the first logical table's body rows.
        score += 12.0
    else:
        fragmentation = _fragmented_body_row_ratio(table)
        if fragmentation >= 0.25:
            score -= 1.0 + fragmentation * 4.0

    if layout_bboxes:
        best_overlap = max((_bbox_intersection_ratio(_table_bbox(table), lb) for lb in layout_bboxes), default=0.0)
        if best_overlap >= 0.35:
            score += 3.0 * best_overlap
        else:
            score -= 3.0

        # Penalize candidates that cover much more vertical area than the
        # layout table regions on the same page; this catches text-block false
        # positives while still allowing slight detector expansion.
        table_h = max(bbox[3] - bbox[1], 1.0)
        best_h = max((min(bbox[3], lb[3]) - max(bbox[1], lb[1]) for lb in layout_bboxes), default=0.0)
        if best_h > 0 and table_h > best_h * 1.8:
            score -= 2.0

    return score


def _candidate_set_score(tables: List[TableRegion],
                         layout_bboxes: List[Tuple[float, float, float, float]]) -> float:
    if not tables:
        return -2.0 * len(layout_bboxes)
    score = sum(_candidate_table_score(table, layout_bboxes) for table in tables)
    if layout_bboxes:
        matched = 0
        for lb in layout_bboxes:
            if max((_bbox_intersection_ratio(_table_bbox(table), lb) for table in tables), default=0.0) >= 0.35:
                matched += 1
        score += matched * 1.5
        score -= (len(layout_bboxes) - matched) * 1.2
        if len(tables) > len(layout_bboxes):
            score -= (len(tables) - len(layout_bboxes)) * 2.0
    return score


def _candidate_set_grid_cells(tables: List[TableRegion]) -> int:
    total = 0
    for table in tables:
        rows = max(0, int(getattr(table, "row_count", 0) or 0))
        cols = max(0, int(getattr(table, "col_count", 0) or 0))
        total += rows * cols
    return total


def _candidate_set_nonempty_cells(tables: List[TableRegion]) -> int:
    total = 0
    for table in tables:
        rows = int(getattr(table, "row_count", 0) or 0)
        cols = int(getattr(table, "col_count", 0) or 0)
        raw_cells = getattr(table, "cells", []) or []
        for r in range(rows):
            for c in range(cols):
                if r < len(raw_cells) and c < len(raw_cells[r]) and _clean_extracted_text(str(raw_cells[r][c])):
                    total += 1
    return total


def _candidate_set_layout_matches(tables: List[TableRegion],
                                  layout_bboxes: List[Tuple[float, float, float, float]]) -> int:
    if not layout_bboxes:
        return len(tables)
    matched = 0
    for lb in layout_bboxes:
        if max((_bbox_intersection_ratio(_table_bbox(table), lb) for table in tables), default=0.0) >= 0.35:
            matched += 1
    return matched


def _candidate_set_structurally_usable(tables: List[TableRegion]) -> bool:
    if not tables:
        return False
    for table in tables:
        rows = getattr(table, "row_count", 0) or 0
        cols = getattr(table, "col_count", 0) or 0
        if rows <= 0 or cols <= 0:
            return False
        if cols == 1 and rows > 1:
            return False
        if cols > 14:
            return False
    return True


def _select_table_candidates_for_page(page: int,
                                      candidate_sets: Dict[str, List[TableRegion]],
                                      layout_regions_by_page: Optional[Dict[int, List[dict]]],
                                      logger: Logger) -> List[TableRegion]:
    layout_bboxes = _layout_table_bboxes(layout_regions_by_page, page)
    scored = []
    for source, tables in candidate_sets.items():
        if not tables:
            continue
        score = _candidate_set_score(tables, layout_bboxes)
        scored.append((score, source, tables))
    if not scored:
        return []

    scored.sort(key=lambda item: item[0], reverse=True)
    best_score, best_source, best_tables = scored[0]

    # Conservative tie-break: when Img2Table and GMFT-ONNX find the same number
    # of tables with the same column grid, prefer the candidate with more row
    # boundaries if scores are close. This keeps extra header/number rows that
    # can be normalized later instead of losing them permanently.
    score_by_source = {source: (score, tables) for score, source, tables in scored}
    if "gmft_onnx" in score_by_source:
        onnx_score, onnx_tables = score_by_source["gmft_onnx"]
        if any(_has_stacked_header_repeat(table) for table in onnx_tables) and best_score - onnx_score <= 3.0:
            best_score, best_source, best_tables = onnx_score, "gmft_onnx", onnx_tables

    if best_source == "img2table" and "gmft_onnx" in score_by_source:
        onnx_score, onnx_tables = score_by_source["gmft_onnx"]
        if len(onnx_tables) == len(best_tables) and best_score - onnx_score <= 2.0:
            onnx_sorted = sorted(onnx_tables, key=lambda t: (t.y_top, t.y_bottom))
            best_sorted = sorted(best_tables, key=lambda t: (t.y_top, t.y_bottom))
            comparable = all(o.col_count == b.col_count for o, b in zip(onnx_sorted, best_sorted))
            row_better = any(o.row_count > b.row_count for o, b in zip(onnx_sorted, best_sorted))
            onnx_fragmented = max((_fragmented_body_row_ratio(t) for t in onnx_sorted), default=0.0) >= 0.25
            if comparable and row_better and not onnx_fragmented:
                best_score, best_source, best_tables = onnx_score, "gmft_onnx", onnx_tables

    if best_source == "rapidtable_slanet":
        established = [
            (score, source, tables)
            for score, source, tables in scored
            if source in {"gmft_onnx", "img2table", "legacy_gmft"}
        ]
        if established:
            est_score, est_source, est_tables = max(established, key=lambda item: item[0])
            rapid_matches = _candidate_set_layout_matches(best_tables, layout_bboxes)
            est_matches = _candidate_set_layout_matches(est_tables, layout_bboxes)
            rapid_cells = _candidate_set_grid_cells(best_tables)
            est_cells = _candidate_set_grid_cells(est_tables)
            rapid_is_richer_grid = rapid_cells >= max(est_cells * 1.4, est_cells + 6)
            if (
                est_matches >= rapid_matches
                and best_score - est_score <= 1.0
                and _candidate_set_structurally_usable(est_tables)
                and not rapid_is_richer_grid
            ):
                best_score, best_source, best_tables = est_score, est_source, est_tables
    elif "rapidtable_slanet" in score_by_source:
        rapid_score, rapid_tables = score_by_source["rapidtable_slanet"]
        rapid_matches = _candidate_set_layout_matches(rapid_tables, layout_bboxes)
        best_matches = _candidate_set_layout_matches(best_tables, layout_bboxes)
        rapid_cells = _candidate_set_grid_cells(rapid_tables)
        best_cells = _candidate_set_grid_cells(best_tables)
        if (
            rapid_matches >= best_matches
            and best_score - rapid_score <= 1.5
            and rapid_cells >= max(best_cells * 1.45, best_cells + 8)
        ):
            best_score, best_source, best_tables = rapid_score, "rapidtable_slanet", rapid_tables

    summary = ", ".join(f"{src}:{score:.2f}/{len(tbls)}" for score, src, tbls in scored)
    logger.log(f"  Page {page}: Selected {best_source} by quality score ({summary})")
    return best_tables


def detect_tables_hybrid(
    pdf_path: str, 
    logger: Logger, 
    page_info: dict, 
    pdf_lines: List[TextLine],
    device: str = "auto",
    layout_regions_by_page: Optional[Dict[int, List[dict]]] = None
) -> List[TableRegion]:
    """
    Hybrid detection: run available table engines, then choose page candidates
    by table-likeness and layout overlap rather than raw area.
    """
    logger.log("Running Hybrid Table Detection (GMFT-ONNX + Img2Table + RapidTable ensemble)...")
    
    import concurrent.futures

    # Define wrappers to handle availability and execution
    def _run_gmft_onnx():
        if GMFT_ONNX_AVAILABLE:
            try:
                logger.log("Hybrid: Starting GMFT-ONNX thread...")
                res = detect_tables_gmft_onnx(pdf_path, logger, page_info, pdf_lines, device)
                logger.log(f"Hybrid: GMFT-ONNX thread finished. Found {len(res)} tables.")
                return _mark_table_source(res, "gmft_onnx")
            except Exception as e:
                logger.log(f"Hybrid: GMFT-ONNX failed: {e}")
                return []
        logger.log("Hybrid: GMFT-ONNX not available")
        return []

    def _run_legacy_gmft():
        if GMFT_AVAILABLE and detect_tables_gmft is not None:
            try:
                logger.log("Hybrid: Starting legacy PyTorch GMFT thread...")
                res = detect_tables_gmft(pdf_path, logger, page_info, pdf_lines, device)
                logger.log(f"Hybrid: legacy PyTorch GMFT thread finished. Found {len(res)} tables.")
                return _mark_table_source(res, "legacy_gmft")
            except Exception as e:
                logger.log(f"Hybrid: legacy PyTorch GMFT failed: {e}")
                return []
        logger.log("Hybrid: legacy PyTorch GMFT not enabled")
        return []

    def _run_img2table():
        if IMG2TABLE_AVAILABLE:
            try:
                logger.log("Hybrid: Starting Img2Table thread...")
                res = detect_tables_img2table(pdf_path, logger, page_info, pdf_lines)
                logger.log(f"Hybrid: Img2Table thread finished. Found {len(res)} tables.")
                return _mark_table_source(res, "img2table")
            except Exception as e:
                logger.log(f"Hybrid: Img2Table failed: {e}")
                return []
        else:
            logger.log("Hybrid: Img2Table not available")
            return []

    def _run_rapidtable():
        if not RAPIDTABLE_AVAILABLE or detect_tables_rapidtable_slanet is None:
            logger.log("Hybrid: RapidTable SLANet+ not available")
            return []
        if not layout_regions_by_page:
            logger.log("Hybrid: RapidTable SLANet+ skipped because layout table boxes are unavailable")
            return []
        try:
            logger.log("Hybrid: Starting RapidTable SLANet+ thread...")
            lines_by_page: Dict[int, List[TextLine]] = {}
            for line in pdf_lines or []:
                lines_by_page.setdefault(line.page, []).append(line)

            def _resolve_cell_text(page_num: int, bbox: Tuple[float, float, float, float]) -> str:
                cell_lines = get_lines_in_rect(bbox, lines_by_page.get(page_num, []))
                return _text_from_table_cell_lines(cell_lines, page_num, bbox[0], bbox[2], logger)

            res = detect_tables_rapidtable_slanet(
                pdf_path,
                logger,
                page_info,
                pdf_lines,
                layout_regions_by_page,
                text_resolver=_resolve_cell_text,
            )
            logger.log(f"Hybrid: RapidTable SLANet+ thread finished. Found {len(res)} tables.")
            return _mark_table_source(res, "rapidtable_slanet")
        except Exception as e:
            logger.log(f"Hybrid: RapidTable SLANet+ failed: {e}")
            return []

    # Run in parallel
    tables_onnx = []
    tables_img2table = []
    tables_legacy = []
    tables_rapidtable = []
    
    workers = 2
    if RAPIDTABLE_AVAILABLE and detect_tables_rapidtable_slanet is not None:
        workers += 1
    if GMFT_AVAILABLE and detect_tables_gmft is not None:
        workers += 1
    with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
        future_onnx = executor.submit(_run_gmft_onnx)
        future_img = executor.submit(_run_img2table)
        future_rapid = executor.submit(_run_rapidtable)
        future_legacy = executor.submit(_run_legacy_gmft) if GMFT_AVAILABLE and detect_tables_gmft is not None else None
        
        tables_onnx = future_onnx.result()
        tables_img2table = future_img.result()
        tables_rapidtable = future_rapid.result()
        tables_legacy = future_legacy.result() if future_legacy is not None else []

    # Group by Page
    onnx_by_page = {}
    for t in tables_onnx:
        if t.page not in onnx_by_page: onnx_by_page[t.page] = []
        onnx_by_page[t.page].append(t)
        
    img2_by_page = {}
    for t in tables_img2table:
        if t.page not in img2_by_page: img2_by_page[t.page] = []
        img2_by_page[t.page].append(t)

    legacy_by_page = {}
    for t in tables_legacy:
        if t.page not in legacy_by_page: legacy_by_page[t.page] = []
        legacy_by_page[t.page].append(t)

    rapid_by_page = {}
    for t in tables_rapidtable:
        if t.page not in rapid_by_page: rapid_by_page[t.page] = []
        rapid_by_page[t.page].append(t)
        
    final_tables = []
    all_pages = set(onnx_by_page.keys()) | set(img2_by_page.keys()) | set(legacy_by_page.keys()) | set(rapid_by_page.keys())
    
    for page in sorted(all_pages):
        chosen = _select_table_candidates_for_page(
            page,
            {
                "gmft_onnx": onnx_by_page.get(page, []),
                "img2table": img2_by_page.get(page, []),
                "legacy_gmft": legacy_by_page.get(page, []),
                "rapidtable_slanet": rapid_by_page.get(page, []),
            },
            layout_regions_by_page,
            logger,
        )
        final_tables.extend(chosen)
            
    return final_tables


def _group_tables_by_page(tables: List[TableRegion]) -> Dict[int, List[TableRegion]]:
    by_page: Dict[int, List[TableRegion]] = {}
    for table in tables or []:
        by_page.setdefault(int(getattr(table, "page", 0) or 0), []).append(table)
    for page_tables in by_page.values():
        page_tables.sort(
            key=lambda t: (
                float(getattr(t, "y_top", 0.0) or 0.0),
                float(getattr(t, "x_left", 0.0) or 0.0),
            )
        )
    return by_page


def _layout_pages_with_tables(layout_regions_by_page: Optional[Dict[int, List[dict]]]) -> List[int]:
    pages = []
    for page, regions in (layout_regions_by_page or {}).items():
        if any(region.get("type") == "table" and region.get("bbox_pdf") for region in regions or []):
            pages.append(int(page))
    return sorted(set(pages))


def _same_table_shapes(left: List[TableRegion], right: List[TableRegion]) -> bool:
    if len(left) != len(right) or not left:
        return False
    left_shapes = [
        (int(getattr(t, "row_count", 0) or 0), int(getattr(t, "col_count", 0) or 0))
        for t in sorted(left, key=lambda t: (float(getattr(t, "y_top", 0.0) or 0.0), float(getattr(t, "x_left", 0.0) or 0.0)))
    ]
    right_shapes = [
        (int(getattr(t, "row_count", 0) or 0), int(getattr(t, "col_count", 0) or 0))
        for t in sorted(right, key=lambda t: (float(getattr(t, "y_top", 0.0) or 0.0), float(getattr(t, "x_left", 0.0) or 0.0)))
    ]
    return all(lr > 0 and lc > 0 and (lr, lc) == (rr, rc) for (lr, lc), (rr, rc) in zip(left_shapes, right_shapes))


def _candidate_set_empty_cell_ratio(tables: List[TableRegion]) -> float:
    total = 0
    empty = 0
    for table in tables or []:
        rows = int(getattr(table, "row_count", 0) or 0)
        cols = int(getattr(table, "col_count", 0) or 0)
        raw_cells = getattr(table, "cells", []) or []
        for r in range(rows):
            for c in range(cols):
                total += 1
                text = ""
                if r < len(raw_cells) and c < len(raw_cells[r]):
                    text = _clean_extracted_text(str(raw_cells[r][c]))
                if not text:
                    empty += 1
    return empty / total if total else 1.0


def _candidate_set_average_ocr_fit(tables: List[TableRegion], pdf_lines: List[TextLine]) -> float:
    if not tables:
        return -5.0
    try:
        from scanindex.core.tables.postprocess_v2 import table_ocr_fit_score
    except Exception:
        return 0.0
    return sum(table_ocr_fit_score(table, pdf_lines) for table in tables) / max(len(tables), 1)


def _choose_docling_first_candidates(
    page: int,
    candidate_sets: Dict[str, List[TableRegion]],
    layout_regions_by_page: Optional[Dict[int, List[dict]]],
    pdf_lines: List[TextLine],
    logger: Logger,
) -> List[TableRegion]:
    """
    Geometry/OCR consensus selector for GMFT and Docling TableFormer v1 ONNX.

    Both engines run on the same DocLayout table boxes. Selection is based on
    layout overlap, OCR-box fit, and grid richness. Raw empty-cell ratio is kept
    as a diagnostic only because some engines assign text later in the common
    geometry postprocess.
    """
    docling_tables = candidate_sets.get("docling_tableformer", []) or []
    gmft_tables = candidate_sets.get("gmft_onnx_layout", []) or []
    layout_bboxes = _layout_table_bboxes(layout_regions_by_page, page)

    if not docling_tables:
        return gmft_tables
    if not gmft_tables:
        return docling_tables

    docling_score = _candidate_set_score(docling_tables, layout_bboxes)
    gmft_score = _candidate_set_score(gmft_tables, layout_bboxes)
    docling_matches = _candidate_set_layout_matches(docling_tables, layout_bboxes)
    gmft_matches = _candidate_set_layout_matches(gmft_tables, layout_bboxes)
    docling_usable = _candidate_set_structurally_usable(docling_tables)
    gmft_usable = _candidate_set_structurally_usable(gmft_tables)
    docling_ocr = _candidate_set_average_ocr_fit(docling_tables, pdf_lines)
    gmft_ocr = _candidate_set_average_ocr_fit(gmft_tables, pdf_lines)
    docling_empty = _candidate_set_empty_cell_ratio(docling_tables)
    gmft_empty = _candidate_set_empty_cell_ratio(gmft_tables)
    docling_cells = _candidate_set_grid_cells(docling_tables)
    gmft_cells = _candidate_set_grid_cells(gmft_tables)
    docling_nonempty = _candidate_set_nonempty_cells(docling_tables)
    gmft_nonempty = _candidate_set_nonempty_cells(gmft_tables)

    use_gmft = False
    reason = ""
    if not docling_usable and gmft_usable:
        use_gmft = True
        reason = "docling unusable"
    elif gmft_matches > docling_matches:
        use_gmft = True
        reason = "gmft matches more layout boxes"
    elif gmft_usable and gmft_ocr >= docling_ocr + 0.70 and gmft_score >= docling_score - 1.50:
        use_gmft = True
        reason = "gmft has materially better OCR coverage"
    elif (
        gmft_usable
        and gmft_matches == docling_matches
        and gmft_ocr >= docling_ocr + 0.03
        and gmft_score >= docling_score - 0.50
        and gmft_cells >= max(1, int(docling_cells * 0.70))
    ):
        use_gmft = True
        reason = "gmft has cleaner OCR fit with comparable geometry"
    elif (
        gmft_usable
        and gmft_matches == docling_matches
        and gmft_cells >= docling_cells + max(4, int(docling_cells * 0.12))
        and gmft_ocr >= docling_ocr - 0.50
        and gmft_score >= docling_score - 1.00
    ):
        use_gmft = True
        reason = "gmft provides richer grid with comparable geometry/OCR fit"

    summary = (
        f"docling_v1_onnx:S{docling_score:.2f}/O{docling_ocr:.2f}/E{docling_empty:.2f}/"
        f"M{docling_matches}/C{docling_cells}/N{docling_nonempty}; "
        f"gmft:S{gmft_score:.2f}/O{gmft_ocr:.2f}/E{gmft_empty:.2f}/"
        f"M{gmft_matches}/C{gmft_cells}/N{gmft_nonempty}"
    )
    if use_gmft:
        logger.log(f"  Page {page}: Selected gmft_onnx_layout by GMFT+Docling selector ({reason}; {summary})")
        return gmft_tables

    logger.log(f"  Page {page}: Selected docling_tableformer by Docling-first selector ({summary})")
    return docling_tables


def detect_tables_doclayout_gmft_docling(
    pdf_path: str,
    logger: Logger,
    page_info: dict,
    pdf_lines: List[TextLine],
    layout_regions_by_page: Optional[Dict[int, List[dict]]] = None,
) -> List[TableRegion]:
    """
    Production table pipeline:
    DocLayout provides table bboxes, GMFT-ONNX and Docling TableFormer v1 ONNX
    recognize structure on those same regions, then a geometry/OCR-fit scorer
    chooses the best page candidate set. No text-specific hardcoding is used.
    """
    if not layout_regions_by_page:
        return []

    layout_pages = _layout_pages_with_tables(layout_regions_by_page)
    if not layout_pages:
        return []

    logger.log("Running DocLayout-anchored tables (DocLayout bbox + GMFT-ONNX + Docling TableFormer v1 ONNX)...")

    import concurrent.futures

    def _run_gmft_layout():
        if not GMFT_ONNX_AVAILABLE or detect_tables_gmft_onnx_on_layout_regions is None:
            logger.log("DocLayout table pipeline: GMFT-ONNX structure recognizer not available")
            return []
        try:
            logger.log("DocLayout table pipeline: Starting GMFT-ONNX structure recognizer...")
            res = detect_tables_gmft_onnx_on_layout_regions(
                pdf_path,
                logger,
                page_info,
                pdf_lines,
                layout_regions_by_page,
                "cpu",
            )
            logger.log(f"DocLayout table pipeline: GMFT-ONNX finished. Found {len(res)} tables.")
            return _mark_table_source(res, "gmft_onnx_layout")
        except Exception as exc:
            logger.log(f"DocLayout table pipeline: GMFT-ONNX failed: {exc}")
            return []

    def _run_docling():
        if not DOCLING_TABLEFORMER_AVAILABLE or detect_tables_docling_tableformer_v1_onnx is None:
            logger.log("DocLayout table pipeline: Docling TableFormer v1 ONNX not available")
            return []
        try:
            logger.log("DocLayout table pipeline: Starting Docling TableFormer v1 ONNX...")
            res = detect_tables_docling_tableformer_v1_onnx(
                pdf_path,
                logger,
                page_info,
                pdf_lines,
                layout_regions_by_page,
                dpi=144,
                pad_pt=0.0,
                num_threads=4,
            )
            logger.log(f"DocLayout table pipeline: Docling TableFormer v1 ONNX finished. Found {len(res)} tables.")
            return _mark_table_source(res, "docling_tableformer")
        except Exception as exc:
            logger.log(f"DocLayout table pipeline: Docling TableFormer v1 ONNX failed: {exc}")
            return []

    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
        future_gmft = executor.submit(_run_gmft_layout)
        future_docling = executor.submit(_run_docling)
        tables_gmft = future_gmft.result()
        tables_docling = future_docling.result()

    gmft_by_page = _group_tables_by_page(tables_gmft)
    docling_by_page = _group_tables_by_page(tables_docling)

    final_tables: List[TableRegion] = []
    for page in layout_pages:
        candidate_sets = {
            "gmft_onnx_layout": gmft_by_page.get(page, []),
            "docling_tableformer": docling_by_page.get(page, []),
        }
        chosen_tables = _choose_docling_first_candidates(
            page,
            candidate_sets,
            layout_regions_by_page,
            pdf_lines,
            logger,
        )
        final_tables.extend(chosen_tables)

    logger.log(f"DocLayout table pipeline: selected {len(final_tables)} tables")
    return final_tables


def detect_tables_doclayout_rapidtable(
    pdf_path: str,
    logger: Logger,
    page_info: dict,
    pdf_lines: List[TextLine],
    layout_regions_by_page: Optional[Dict[int, List[dict]]] = None,
) -> List[TableRegion]:
    """
    Portable-first table pipeline:
    DocLayout provides table bboxes and RapidTable SLANet+ recognizes structure.
    The common geometry/OCR postprocess still assigns and repairs cell text.
    """
    if not layout_regions_by_page:
        return []
    if not RAPIDTABLE_AVAILABLE or detect_tables_rapidtable_slanet is None:
        logger.log("RapidTable primary pipeline: RapidTable SLANet+ not available")
        return []

    layout_pages = _layout_pages_with_tables(layout_regions_by_page)
    if not layout_pages:
        return []

    logger.log("Running DocLayout-anchored tables (DocLayout bbox + RapidTable SLANet+ primary)...")
    try:
        lines_by_page: Dict[int, List[TextLine]] = {}
        for line in pdf_lines or []:
            lines_by_page.setdefault(line.page, []).append(line)

        def _resolve_cell_text(page_num: int, bbox: Tuple[float, float, float, float]) -> str:
            cell_lines = get_lines_in_rect(bbox, lines_by_page.get(page_num, []))
            return _text_from_table_cell_lines(cell_lines, page_num, bbox[0], bbox[2], logger)

        tables = detect_tables_rapidtable_slanet(
            pdf_path,
            logger,
            page_info,
            pdf_lines,
            layout_regions_by_page,
            text_resolver=_resolve_cell_text,
        )
        tables = _mark_table_source(tables, "rapidtable_slanet")
        logger.log(f"RapidTable primary pipeline: selected {len(tables)} tables")
        return tables
    except Exception as exc:
        logger.log(f"RapidTable primary pipeline failed: {exc}")
        return []


def _is_digital_text_pdf(pdf_path: str) -> bool:
    """True when the companion OCR JSON came from native digital PDF extraction."""
    try:
        from scanindex.core.pdf.text_extractor import is_digital_ocr_output

        return bool(is_digital_ocr_output(pdf_path))
    except Exception:
        return False


def _page_source_modes_from_companion(companion_data: Optional[dict]) -> Dict[int, str]:
    pages = (companion_data or {}).get("pages") or []
    pipeline_mode = (
        ((companion_data or {}).get("pipeline") or {}).get("ocr") or {}
    ).get("source_mode")
    modes: Dict[int, str] = {}
    for idx, page in enumerate(pages, 1):
        mode = str((page or {}).get("source_mode") or pipeline_mode or "").lower()
        if mode in {"scan", "digital", "mixed"}:
            modes[idx] = mode
    return modes


def _filter_lines_to_pages(lines: List[TextLine], pages: Set[int]) -> List[TextLine]:
    if not pages:
        return []
    return [line for line in lines or [] if int(line.page) in pages]


def _filter_layout_regions_to_pages(
    layout_regions_by_page: Optional[Dict[int, List[dict]]],
    pages: Set[int],
) -> Dict[int, List[dict]]:
    if not layout_regions_by_page or not pages:
        return {}
    return {
        int(page): regions
        for page, regions in layout_regions_by_page.items()
        if int(page) in pages
    }


def _table_overlap_ratio(a: TableRegion, b: TableRegion) -> float:
    if int(getattr(a, "page", 0) or 0) != int(getattr(b, "page", 0) or 0):
        return 0.0
    ax0, ay0, ax1, ay1 = _table_bbox(a)
    bx0, by0, bx1, by1 = _table_bbox(b)
    ix0 = max(ax0, bx0)
    iy0 = max(ay0, by0)
    ix1 = min(ax1, bx1)
    iy1 = min(ay1, by1)
    inter = max(0.0, ix1 - ix0) * max(0.0, iy1 - iy0)
    area_a = max(1.0, (ax1 - ax0) * (ay1 - ay0))
    area_b = max(1.0, (bx1 - bx0) * (by1 - by0))
    return max(inter / area_a, inter / area_b)


def _merge_native_and_model_tables(
    native_tables: List[TableRegion],
    model_tables: List[TableRegion],
    logger: Logger,
) -> List[TableRegion]:
    """Prefer native table text for overlapping digital regions; keep model tables elsewhere."""
    result = list(native_tables or [])
    skipped = 0
    for table in model_tables or []:
        if any(_table_overlap_ratio(table, native) >= 0.70 for native in native_tables or []):
            skipped += 1
            continue
        result.append(table)
    if native_tables and model_tables:
        logger.log(
            f"Combined table engines: kept {len(native_tables)} native table(s), "
            f"{len(model_tables) - skipped} model table(s), skipped {skipped} overlap(s)"
        )
    return sorted(result, key=lambda t: (t.page, t.y_top, t.y_bottom))


def _model_table_candidate_pages(
    page_modes: Dict[int, str],
    layout_regions_by_page: Dict[int, List[dict]],
) -> Set[int]:
    candidate_pages = {
        int(page)
        for page, regions in (layout_regions_by_page or {}).items()
        if regions
    }
    pages: Set[int] = set()
    for page, mode in page_modes.items():
        if mode in {"scan", "mixed"}:
            pages.add(int(page))
        elif mode == "digital" and int(page) in candidate_pages:
            pages.add(int(page))
    if not page_modes:
        pages.update(candidate_pages)
    return pages


def _assign_table_lines_preserving_native_text(
    native_tables: List[TableRegion],
    model_tables: List[TableRegion],
    pdf_lines: List[TextLine],
    logger: Logger,
) -> Set[int]:
    assigned: Set[int] = set()
    if native_tables:
        assigned.update(assign_ocr_lines_to_table_cells_by_geometry(
            native_tables,
            pdf_lines,
            logger,
            rebuild_cells=False,
        ))
    if model_tables:
        assigned.update(assign_ocr_lines_to_table_cells_by_geometry(
            model_tables,
            pdf_lines,
            logger,
        ))
    return assigned


def _native_bbox_tuple(bbox) -> Optional[Tuple[float, float, float, float]]:
    if bbox is None:
        return None
    try:
        values = tuple(float(v) for v in bbox[:4])
    except Exception:
        return None
    if len(values) != 4 or values[2] <= values[0] or values[3] <= values[1]:
        return None
    return values


def _cluster_axis_positions(values: List[float], tolerance: float = 2.0) -> List[float]:
    if not values:
        return []
    clusters: List[List[float]] = []
    for value in sorted(float(v) for v in values):
        if clusters and abs(value - (sum(clusters[-1]) / len(clusters[-1]))) <= tolerance:
            clusters[-1].append(value)
        else:
            clusters.append([value])
    return [sum(cluster) / len(cluster) for cluster in clusters]


def _even_intervals(start: float, end: float, count: int) -> List[Tuple[float, float]]:
    if count <= 0 or end <= start:
        return []
    width = (end - start) / count
    return [(start + width * idx, start + width * (idx + 1)) for idx in range(count)]


def _native_table_column_intervals(table_obj, table_bbox: Tuple[float, float, float, float], col_count: int) -> List[Tuple[float, float]]:
    rows = list(getattr(table_obj, "rows", []) or [])
    for row in rows:
        row_cells = list(getattr(row, "cells", []) or [])
        if len(row_cells) < col_count:
            continue
        intervals: List[Tuple[float, float]] = []
        ok = True
        for col_idx in range(col_count):
            bbox = _native_bbox_tuple(row_cells[col_idx])
            if bbox is None:
                ok = False
                break
            intervals.append((bbox[0], bbox[2]))
        if ok and all(intervals[idx][1] <= intervals[idx + 1][0] + 2.0 for idx in range(len(intervals) - 1)):
            return intervals

    x_values = [table_bbox[0], table_bbox[2]]
    for row in rows:
        for cell_bbox in list(getattr(row, "cells", []) or []):
            bbox = _native_bbox_tuple(cell_bbox)
            if bbox is not None:
                x_values.extend([bbox[0], bbox[2]])
    boundaries = [
        value for value in _cluster_axis_positions(x_values, 2.0)
        if table_bbox[0] - 2.0 <= value <= table_bbox[2] + 2.0
    ]
    if len(boundaries) == col_count + 1:
        return [(boundaries[idx], boundaries[idx + 1]) for idx in range(col_count)]
    return _even_intervals(table_bbox[0], table_bbox[2], col_count)


def _native_table_row_bands(table_obj, table_bbox: Tuple[float, float, float, float], row_count: int) -> List[Tuple[float, float]]:
    rows = list(getattr(table_obj, "rows", []) or [])
    bands: List[Tuple[float, float]] = []
    for row in rows[:row_count]:
        bbox = _native_bbox_tuple(getattr(row, "bbox", None))
        if bbox is None:
            row_cells = [_native_bbox_tuple(cell) for cell in list(getattr(row, "cells", []) or [])]
            row_cells = [cell for cell in row_cells if cell is not None]
            if row_cells:
                bbox = (
                    min(cell[0] for cell in row_cells),
                    min(cell[1] for cell in row_cells),
                    max(cell[2] for cell in row_cells),
                    max(cell[3] for cell in row_cells),
                )
        if bbox is not None:
            bands.append((bbox[1], bbox[3]))
    if len(bands) == row_count and all(bands[idx][1] <= bands[idx + 1][0] + 2.0 for idx in range(len(bands) - 1)):
        return bands

    y_values = [table_bbox[1], table_bbox[3]]
    for row in rows:
        for cell_bbox in list(getattr(row, "cells", []) or []):
            bbox = _native_bbox_tuple(cell_bbox)
            if bbox is not None:
                y_values.extend([bbox[1], bbox[3]])
    boundaries = [
        value for value in _cluster_axis_positions(y_values, 2.0)
        if table_bbox[1] - 2.0 <= value <= table_bbox[3] + 2.0
    ]
    if len(boundaries) == row_count + 1:
        return [(boundaries[idx], boundaries[idx + 1]) for idx in range(row_count)]
    return _even_intervals(table_bbox[1], table_bbox[3], row_count)


def _normalize_native_cells(raw_cells, row_count: int, col_count: int) -> List[List[str]]:
    rows: List[List[str]] = []
    for row_idx in range(row_count):
        raw_row = raw_cells[row_idx] if raw_cells and row_idx < len(raw_cells) else []
        row: List[str] = []
        for col_idx in range(col_count):
            value = raw_row[col_idx] if col_idx < len(raw_row) else ""
            row.append(_clean_extracted_text(str(value or "")))
        rows.append(row)
    return rows


def _native_span_columns(
    bbox: Tuple[float, float, float, float],
    col_intervals: List[Tuple[float, float]],
    fallback_col: int,
) -> Tuple[int, int]:
    overlapping: List[int] = []
    for idx, (col_x0, col_x1) in enumerate(col_intervals):
        col_width = max(col_x1 - col_x0, 1e-6)
        overlap = max(0.0, min(bbox[2], col_x1) - max(bbox[0], col_x0))
        center = (col_x0 + col_x1) / 2.0
        if overlap / col_width >= 0.45 or bbox[0] - 1.0 <= center <= bbox[2] + 1.0:
            overlapping.append(idx)
    if not overlapping:
        col = min(max(0, fallback_col), max(0, len(col_intervals) - 1))
        return (col, col)
    return (min(overlapping), max(overlapping))


def _native_table_spans(
    table_obj,
    cells: List[List[str]],
    col_intervals: List[Tuple[float, float]],
) -> Dict[int, List[Tuple[int, int, str]]]:
    span_rows: Dict[int, List[Tuple[int, int, str]]] = {}
    rows = list(getattr(table_obj, "rows", []) or [])
    row_count = len(cells)
    col_count = len(cells[0]) if cells else 0
    for row_idx, row_obj in enumerate(rows[:row_count]):
        row_cells = list(getattr(row_obj, "cells", []) or [])
        row_spans: List[Tuple[int, int, str]] = []
        for cell_idx, cell_bbox in enumerate(row_cells[:col_count]):
            bbox = _native_bbox_tuple(cell_bbox)
            if bbox is None:
                continue
            text = cells[row_idx][cell_idx] if cell_idx < len(cells[row_idx]) else ""
            if not text:
                continue
            start_col, end_col = _native_span_columns(bbox, col_intervals, cell_idx)
            if end_col <= start_col:
                continue
            row_spans.append((start_col, end_col, text))
            cells[row_idx][start_col] = text
            for col_idx in range(start_col + 1, end_col + 1):
                if col_idx < len(cells[row_idx]):
                    cells[row_idx][col_idx] = ""
        if row_spans:
            seen = set()
            unique_spans = []
            for span in sorted(row_spans, key=lambda item: (item[0], item[1])):
                key = (span[0], span[1], _norm_table_key(span[2]))
                if key in seen:
                    continue
                seen.add(key)
                unique_spans.append(span)
            span_rows[row_idx] = unique_spans
    return span_rows


def _native_table_region_from_pymupdf(
    page_num: int,
    table_obj,
    source: str,
    page_size: Optional[Tuple[float, float]] = None,
) -> Optional[TableRegion]:
    table_bbox = _native_bbox_tuple(getattr(table_obj, "bbox", None))
    if table_bbox is None:
        return None
    row_count = int(getattr(table_obj, "row_count", 0) or 0)
    col_count = int(getattr(table_obj, "col_count", 0) or 0)
    if row_count < 2 or col_count < 2 or col_count > 30:
        return None

    try:
        raw_cells = table_obj.extract()
    except Exception:
        raw_cells = []
    cells = _normalize_native_cells(raw_cells, row_count, col_count)
    nonempty = sum(1 for row in cells for cell in row if cell)
    if nonempty < 2 or nonempty / max(row_count * col_count, 1) < 0.08:
        return None
    if source == "pymupdf_native_text" and page_size:
        page_w, page_h = page_size
        width_cover = (table_bbox[2] - table_bbox[0]) / max(float(page_w), 1.0)
        height_cover = (table_bbox[3] - table_bbox[1]) / max(float(page_h), 1.0)
        nonempty_texts = [cell for row in cells for cell in row if cell]
        short_fragment_ratio = sum(
            1 for cell in nonempty_texts
            if len(cell) <= 2 and not re.fullmatch(r"\d+([.,]\d+)?", cell)
        ) / max(len(nonempty_texts), 1)
        if (
            (row_count > 25 and width_cover > 0.55 and height_cover > 0.35)
            or (row_count > 40 and height_cover > 0.50)
            or (col_count > 8 and row_count > 15)
            or (col_count >= 6 and short_fragment_ratio > 0.30)
        ):
            return None

    col_intervals = _native_table_column_intervals(table_obj, table_bbox, col_count)
    row_bands = _native_table_row_bands(table_obj, table_bbox, row_count)
    if len(col_intervals) != col_count or len(row_bands) != row_count:
        return None

    cell_bboxes = [
        [
            (col_intervals[col_idx][0], row_bands[row_idx][0], col_intervals[col_idx][1], row_bands[row_idx][1])
            for col_idx in range(col_count)
        ]
        for row_idx in range(row_count)
    ]
    horizontal_text_spans = _native_table_spans(table_obj, cells, col_intervals)

    region = TableRegion(
        page=page_num,
        y_top=table_bbox[1],
        y_bottom=table_bbox[3],
        cells=cells,
        row_count=row_count,
        col_count=col_count,
        cell_bboxes=cell_bboxes,
    )
    setattr(region, "x_left", table_bbox[0])
    setattr(region, "x_right", table_bbox[2])
    setattr(region, "source", source)
    setattr(region, "native_table", True)
    if horizontal_text_spans:
        setattr(region, "horizontal_text_spans", horizontal_text_spans)
    return region


def _native_table_seen_key(region: TableRegion) -> Tuple[int, int, int, int, int]:
    bbox = _table_bbox(region)
    return (
        int(region.page),
        int(round(bbox[0])),
        int(round(bbox[1])),
        int(round(bbox[2])),
        int(round(bbox[3])),
    )


def detect_tables_pymupdf_native(
    pdf_path: str,
    logger: Logger,
    pdf_lines: List[TextLine],
    include_text_strategy: bool = False,
) -> List[TableRegion]:
    """
    Extract tables from digital PDFs using PyMuPDF's vector/text table finder.

    This path trusts native PDF geometry and extracted cell text, then only uses
    OCR line geometry later to suppress duplicate paragraph output. It avoids
    OCR/model table repairs that can rewrite correct digital cell values.
    """
    strategies = [
        ("pymupdf_native_lines", {}),
    ]
    if include_text_strategy or os.environ.get("OCRTOOL_ENABLE_PYMUPDF_TEXT_TABLES") == "1":
        strategies.append((
            "pymupdf_native_text",
            {
                "vertical_strategy": "text",
                "horizontal_strategy": "text",
                "min_words_vertical": 2,
                "min_words_horizontal": 1,
            },
        ))

    tables: List[TableRegion] = []
    seen = set()
    try:
        with fitz.open(pdf_path) as doc:
            for page_idx, page in enumerate(doc):
                page_added = False
                for source, kwargs in strategies:
                    if source == "pymupdf_native_text" and page_added:
                        continue
                    try:
                        import contextlib
                        import io

                        with contextlib.redirect_stdout(io.StringIO()):
                            found = page.find_tables(**kwargs)
                    except Exception as exc:
                        if source == "pymupdf_native_lines":
                            logger.log(f"PyMuPDF native table finder failed on page {page_idx + 1}: {exc}")
                        continue
                    for table_obj in list(getattr(found, "tables", []) or []):
                        region = _native_table_region_from_pymupdf(
                            page_idx + 1,
                            table_obj,
                            source,
                            (float(page.rect.width), float(page.rect.height)),
                        )
                        if region is None:
                            continue
                        key = _native_table_seen_key(region)
                        if key in seen:
                            continue
                        seen.add(key)
                        tables.append(region)
                        page_added = True
    except Exception as exc:
        logger.log(f"PyMuPDF native table extraction failed: {exc}")
        return []

    if tables:
        logger.log(f"PyMuPDF native table extractor: selected {len(tables)} tables with native cell text")
    else:
        logger.log("PyMuPDF native table extractor: no tables found")
    return tables


def detect_tables(
    pdf_path: str, 
    logger: Logger, 
    page_info: dict, 
    pdf_lines: List[TextLine],
    layout_regions_by_page: Optional[Dict[int, List[dict]]] = None
) -> List[TableRegion]:
    """
    Unified table detection function.
    """
    doclayout_tables = detect_tables_doclayout_gmft_docling(
        pdf_path,
        logger,
        page_info,
        pdf_lines,
        layout_regions_by_page,
    )
    if not doclayout_tables:
        logger.log("DocLayout GMFT + Docling TableFormer v1 ONNX produced no tables; no alternate table fallback is enabled")
    return doclayout_tables


# ============================================================================
# CREATE DOCX WITH POSITIONS
# ============================================================================

def set_paragraph_font(para, font_name="Times New Roman", font_size=14, bold: Optional[bool] = None):
    """Set font for all runs in paragraph."""
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bool(bold)
        # Set for Asian text too
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


# def get_lines_in_rect moved up


def _same_visual_cell_region(
    bbox_a: Tuple[float, float, float, float],
    bbox_b: Tuple[float, float, float, float],
) -> bool:
    """True only when the structure engine duplicated the same visual cell bbox."""
    if not (any(bbox_a) and any(bbox_b)):
        return False
    area_a = max((bbox_a[2] - bbox_a[0]) * (bbox_a[3] - bbox_a[1]), 1e-6)
    area_b = max((bbox_b[2] - bbox_b[0]) * (bbox_b[3] - bbox_b[1]), 1e-6)
    ix0 = max(bbox_a[0], bbox_b[0])
    iy0 = max(bbox_a[1], bbox_b[1])
    ix1 = min(bbox_a[2], bbox_b[2])
    iy1 = min(bbox_a[3], bbox_b[3])
    inter = max(0.0, ix1 - ix0) * max(0.0, iy1 - iy0)
    return inter / area_a >= 0.88 and inter / area_b >= 0.88


def _word_center_in_bbox(word: dict, bbox: Tuple[float, float, float, float]) -> bool:
    try:
        x = float(word.get("x", 0.0))
        y = float(word.get("y", 0.0))
        w = float(word.get("w", 0.0))
        h = float(word.get("h", 0.0))
    except Exception:
        return False
    cx = x + w / 2.0
    cy = y + h / 2.0
    x0, y0, x1, y1 = bbox
    return x0 <= cx <= x1 and y0 <= cy <= y1


def _cell_words_for_bbox(
    bbox: Tuple[float, float, float, float],
    pdf_lines_page: List[TextLine],
) -> List[Tuple[dict, TextLine]]:
    words: List[Tuple[dict, TextLine]] = []
    for line in pdf_lines_page:
        if _bbox_overlap_ratio(_line_bbox(line), bbox) <= 0 and not (
            bbox[0] <= line.x + line.width / 2.0 <= bbox[2]
            and bbox[1] <= line.y + line.height / 2.0 <= bbox[3]
        ):
            continue
        for word in _word_items_for_line(line):
            if _word_center_in_bbox(word, bbox):
                words.append((word, line))
    return words


def _cell_text_should_be_bold(
    bbox: Tuple[float, float, float, float],
    pdf_lines_page: List[TextLine],
) -> bool:
    cell_words = _cell_words_for_bbox(bbox, pdf_lines_page)
    if not cell_words:
        return False
    known = [(word, line) for word, line in cell_words if _has_known_gray(word.get("fg_gray", 128))]
    if len(known) < max(1, int(len(cell_words) * 0.45)):
        return False
    bold_count = sum(1 for word, line in known if _word_is_visually_bold(word, line))
    return bold_count / max(len(known), 1) >= 0.45


def _row_has_continuation_body(row: List[str], cols: int) -> bool:
    if cols <= 2:
        return False
    body_end = max(2, cols - 1)
    return any(_clean_extracted_text(str(row[c])) for c in range(2, body_end))


def _blank_leading_continuation_row(row: List[str], cols: int) -> bool:
    if cols < 3:
        return False
    first = _clean_extracted_text(str(row[0])) if len(row) > 0 else ""
    second = _clean_extracted_text(str(row[1])) if len(row) > 1 else ""
    return not first and not second and _row_has_continuation_body(row, cols)


def _blank_continuation_vmerge_spans(table_region: TableRegion) -> List[Tuple[int, int, int]]:
    """
    Vertical-merge leading descriptor columns over blank continuation rows.

    This restores rowspans such as an item number/title spanning several
    scoring criteria, including when that span crosses a physical PDF page.
    The signal is structural: an indexed row anchors the span, and following
    rows must have blank leading cells while retaining body text.
    """
    cols = int(getattr(table_region, "col_count", 0) or 0)
    rows = [
        list(row[:cols]) + [""] * max(0, cols - len(row))
        for row in (getattr(table_region, "cells", []) or [])
    ]
    if cols < 4 or len(rows) < 2:
        return []

    spans: List[Tuple[int, int, int]] = []
    for c in (0, 1):
        r = 0
        while r < len(rows) - 1:
            anchor_first = _clean_extracted_text(str(rows[r][0]))
            anchor_text = _clean_extracted_text(str(rows[r][c]))
            if not anchor_text or not _looks_like_table_index_marker(anchor_first):
                r += 1
                continue

            end = r
            j = r + 1
            while j < len(rows) and _blank_leading_continuation_row(rows[j], cols):
                if _clean_extracted_text(str(rows[j][c])):
                    break
                end = j
                j += 1

            if end > r:
                spans.append((r, end, c))
                r = end + 1
            else:
                r += 1

    def _overlaps_existing(start: int, end: int, col: int) -> bool:
        return any(col == sc and start <= se and end >= ss for ss, se, sc in spans)

    # A continuation page can begin in the middle of a vertical span, so there
    # is no anchor text on that physical page. Merge the blank leading run
    # itself to avoid redrawing horizontal rules through the continued cell.
    for c in (0, 1):
        r = 0
        while r < len(rows):
            if not _blank_leading_continuation_row(rows[r], cols):
                r += 1
                continue
            start = r
            while r + 1 < len(rows) and _blank_leading_continuation_row(rows[r + 1], cols):
                r += 1
            end = r
            if end > start and not _overlaps_existing(start, end, c):
                spans.append((start, end, c))
            r += 1
    return spans


def _set_tc_border_nil(tc, edge: str) -> None:
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    edge_el = borders.find(qn(f'w:{edge}'))
    if edge_el is None:
        edge_el = OxmlElement(f'w:{edge}')
        borders.append(edge_el)
    # Some DOCX renderers keep table-style insideH lines even when a cell
    # border is set to nil. A white explicit edge reliably suppresses the
    # visual rule while preserving the table grid and cell addressing.
    edge_el.set(qn('w:val'), 'single')
    edge_el.set(qn('w:sz'), '4')
    edge_el.set(qn('w:space'), '0')
    edge_el.set(qn('w:color'), 'FFFFFF')


def _set_tc_border_black(tc, edge: str) -> None:
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    edge_el = borders.find(qn(f'w:{edge}'))
    if edge_el is None:
        edge_el = OxmlElement(f'w:{edge}')
        borders.append(edge_el)
    edge_el.set(qn('w:val'), 'single')
    edge_el.set(qn('w:sz'), '4')
    edge_el.set(qn('w:space'), '0')
    edge_el.set(qn('w:color'), '000000')


def _apply_explicit_cell_grid_borders(table) -> None:
    for row in table._tbl.tr_lst:
        for tc in row.tc_lst:
            for edge in ("top", "bottom", "left", "right"):
                _set_tc_border_black(tc, edge)


def _suppress_ooxml_internal_span_borders(table,
                                          spans: List[Tuple[int, int, int]],
                                          logger: Logger) -> int:
    applied = 0
    for start, end, col in spans:
        try:
            rows = table._tbl.tr_lst
            if start >= len(rows) or end >= len(rows):
                continue
            ok = True
            for row_idx in range(start, end):
                if col >= len(rows[row_idx].tc_lst) or col >= len(rows[row_idx + 1].tc_lst):
                    ok = False
                    break
                _set_tc_border_nil(rows[row_idx].tc_lst[col], 'bottom')
                _set_tc_border_nil(rows[row_idx + 1].tc_lst[col], 'top')
            if ok:
                applied += 1
        except Exception as exc:
            logger.log(f"  Leading border suppression failed at r{start}-{end}c{col}: {exc}")
    return applied


def _set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn('w:noWrap')) is None:
        tc_pr.append(OxmlElement('w:noWrap'))


def _compact_wrapped_marker_text(text: str) -> str:
    cleaned = _clean_extracted_text(text)
    if re.fullmatch(r"[IVXLCDM](?:\s+[IVXLCDM])+", cleaned, flags=re.IGNORECASE):
        return _norm_table_key(cleaned)
    if re.fullmatch(r"\d(?:\s+\d)+", cleaned):
        return _norm_table_key(cleaned)
    return text



def add_table_to_doc(doc: Document, table_region: TableRegion, pdf_lines_page: List[TextLine], logger: Logger, page_info: dict = None):
    """Add a table to the document with PDF-verified horizontal merging."""
    rows = table_region.row_count
    cols = table_region.col_count
    pg_info = (page_info or {}).get(table_region.page, {})
    table_font_size = int(round(float(pg_info.get("docx_table_font_pt", pg_info.get("docx_body_font_pt", 13)))))
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set proportional column widths from cell bboxes. Pick the row whose
    # neighboring boxes have the smallest horizontal gaps; header text boxes are
    # often much narrower than the ruled cells and make poor width references.
    col_widths_pt = [0.0] * cols
    best_width_score = float("inf")
    table_bbox = _table_bbox(table_region)
    for ref_row in range(rows):
        row_bboxes = []
        for c in range(cols):
            if ref_row < len(table_region.cell_bboxes) and c < len(table_region.cell_bboxes[ref_row]):
                bx = table_region.cell_bboxes[ref_row][c]
                row_bboxes.append(tuple(float(v) for v in bx[:4]))
            else:
                row_bboxes.append((0.0, 0.0, 0.0, 0.0))
        # Good reference row: all widths > 0 AND no adjacent cells share same bbox (not merged)
        all_positive = all(bx[2] > bx[0] for bx in row_bboxes)
        no_merges = True
        if all_positive:
            for c in range(len(row_bboxes) - 1):
                if row_bboxes[c] == row_bboxes[c + 1]:
                    no_merges = False
                    break
        if not (all_positive and no_merges):
            continue

        boundaries = [float(table_bbox[0])]
        gap_sum = 0.0
        overlap_sum = 0.0
        valid_boundaries = True
        for c in range(cols - 1):
            left = row_bboxes[c]
            right = row_bboxes[c + 1]
            gap_sum += max(0.0, right[0] - left[2])
            overlap_sum += max(0.0, left[2] - right[0])
            boundary = (left[2] + right[0]) / 2.0
            if boundary <= boundaries[-1] + 1.0:
                valid_boundaries = False
                break
            boundaries.append(boundary)
        boundaries.append(float(table_bbox[2]))
        if not valid_boundaries:
            continue

        widths = [boundaries[i + 1] - boundaries[i] for i in range(cols)]
        if not widths or any(w <= 1.0 for w in widths):
            continue
        width_score = gap_sum + overlap_sum * 2.0
        if width_score < best_width_score:
            best_width_score = width_score
            col_widths_pt = widths

    total_w = sum(col_widths_pt)
    if total_w > 0:
        # Set fixed layout to prevent Word from auto-resizing
        tbl_xml = table._tbl
        tblPr = tbl_xml.tblPr if tbl_xml.tblPr is not None else OxmlElement('w:tblPr')
        # Remove existing tblLayout if any
        existing_layout = tblPr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tblPr.remove(existing_layout)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        # Also disable autofit
        autofit = OxmlElement('w:tblW')
        autofit.set(qn('w:w'), '0')
        autofit.set(qn('w:type'), 'auto')
        existing_w = tblPr.find(qn('w:tblW'))
        if existing_w is not None:
            tblPr.remove(existing_w)
        tblPr.append(autofit)

        # Apply proportional widths using Twips for precision
        # Use actual DOCX content area width (depends on page orientation)
        pg_num = table_region.page
        _pi = page_info or {}
        pg_w = _pi.get(pg_num, {}).get("width", 595)
        pg_info = _pi.get(pg_num, {})
        left_margin_pt = float(pg_info.get("docx_left_margin_pt", pg_w * 0.10))
        right_margin_pt = float(pg_info.get("docx_right_margin_pt", pg_w * 0.07))
        content_pt = max(pg_w - left_margin_pt - right_margin_pt, pg_w * 0.45)
        content_cm = content_pt / 72.0 * 2.54
        content_twips = int(round(content_pt * 20.0))
        col_twips = [
            max(80, int(round(content_twips * (col_widths_pt[c] / total_w))))
            for c in range(cols)
        ]
        twip_delta = content_twips - sum(col_twips)
        if col_twips:
            col_twips[-1] = max(80, col_twips[-1] + twip_delta)

        autofit.set(qn('w:w'), str(content_twips))
        autofit.set(qn('w:type'), 'dxa')

        existing_grid = tbl_xml.find(qn('w:tblGrid'))
        if existing_grid is not None:
            tbl_xml.remove(existing_grid)
        tbl_grid = OxmlElement('w:tblGrid')
        for col_twip in col_twips:
            grid_col = OxmlElement('w:gridCol')
            grid_col.set(qn('w:w'), str(col_twip))
            tbl_grid.append(grid_col)
        tbl_xml.insert(1, tbl_grid)

        for c in range(cols):
            col_ratio = col_widths_pt[c] / total_w
            col_cm = col_ratio * content_cm
            col_width = Cm(col_cm)
            for row_obj in table.rows:
                row_obj.cells[c].width = col_width
                tc_pr = row_obj.cells[c]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn('w:tcW'))
                if tc_w is None:
                    tc_w = OxmlElement('w:tcW')
                    tc_pr.append(tc_w)
                tc_w.set(qn('w:w'), str(col_twips[c]))
                tc_w.set(qn('w:type'), 'dxa')

    # Pre-calculate row y-ranges and PDF text for verification
    row_pdf_texts = {}
    row_source_pages = list(getattr(table_region, "row_source_pages", []) or [])
    lines_by_page_for_table: Dict[int, List[TextLine]] = {}
    for line in pdf_lines_page:
        lines_by_page_for_table.setdefault(int(line.page), []).append(line)
    for r in range(rows):
        # Find min y and max y for this row based on cell bboxes
        min_y = 10000
        max_y = 0
        has_bbox = False
        for c in range(cols):
            bbox = table_region.cell_bboxes[r][c]
            if any(bbox):
                min_y = min(min_y, bbox[1])
                max_y = max(max_y, bbox[3])
                has_bbox = True
        
        if has_bbox:
            # Extract text lines from PDF that fall within this row's Y range
            # Add some tolerance
            tolerance = 2.0
            row_page = row_source_pages[r] if r < len(row_source_pages) else table_region.page
            candidate_lines = lines_by_page_for_table.get(int(row_page), pdf_lines_page)
            row_lines = [
                l.text for l in candidate_lines
                if min_y - tolerance <= l.y + l.height/2 <= max_y + tolerance
            ]
            row_pdf_texts[r] = " ".join(row_lines)
        else:
            row_pdf_texts[r] = ""

    visited = set()
    horizontal_text_spans = getattr(table_region, "horizontal_text_spans", {}) or {}
    native_table = bool(getattr(table_region, "native_table", False))

    def _row_horizontal_spans(row_idx: int) -> List[Tuple[int, int, str]]:
        raw = horizontal_text_spans.get(row_idx)
        if not raw:
            return []
        if isinstance(raw, tuple) and len(raw) >= 3:
            return [(int(raw[0]), int(raw[1]), str(raw[2]))]
        spans: List[Tuple[int, int, str]] = []
        if isinstance(raw, list):
            for item in raw:
                if isinstance(item, (tuple, list)) and len(item) >= 3:
                    spans.append((int(item[0]), int(item[1]), str(item[2])))
        return spans
    
    for row_idx in range(rows):
        # Calculate table occurrences for this row 
        # (How many times each string appears in the table row)
        # We need this to compare with PDF counts.
        table_row_counts = {}
        if row_idx < len(table_region.cells):
             for val in table_region.cells[row_idx]:
                 val_str = str(val)
                 table_row_counts[val_str] = table_row_counts.get(val_str, 0) + 1
        
        pdf_row_text = row_pdf_texts.get(row_idx, "")
        
        for col_idx in range(cols):
            if (row_idx, col_idx) in visited:
                continue

            row_span_specs = _row_horizontal_spans(row_idx)
            active_span = None
            covered_by_span = False
            for span_start, span_end, span_text in row_span_specs:
                if col_idx == span_start:
                    active_span = (span_start, span_end, span_text)
                    break
                if span_start < col_idx <= span_end:
                    covered_by_span = True
            if covered_by_span and active_span is None:
                visited.add((row_idx, col_idx))
                continue
                
            # Get current text
            current_text = ""
            if row_idx < len(table_region.cells) and col_idx < len(table_region.cells[row_idx]):
                 current_text = str(table_region.cells[row_idx][col_idx])
             
            colspan = 1
            rowspan = 1 # Always 1 as requested (No vertical merging)
            row_vals = table_region.cells[row_idx] if row_idx < len(table_region.cells) else []
            if active_span:
                span_start, span_end, span_text = active_span
                if col_idx == span_start:
                    current_text = str(span_text)
                    colspan = max(1, int(span_end) - int(span_start) + 1)

            # Header cells from TATR sometimes land in the visual center of a
            # multi-column span with empty neighbor cells. Start the merge from
            # the first empty neighbor so the DOCX has one logical header cell
            # instead of a blank + text + blank sequence.
            if not native_table and not current_text.strip() and row_idx == 0 and row_idx < len(table_region.cells):
                next_col = col_idx + 1
                while next_col < cols:
                    next_text = str(row_vals[next_col]) if next_col < len(row_vals) else ""
                    if next_text.strip():
                        if "CONGTAC" in _norm_table_key(next_text):
                            current_text = next_text
                            colspan = next_col - col_idx + 1
                            while col_idx + colspan < cols:
                                tail_text = str(row_vals[col_idx + colspan]) if col_idx + colspan < len(row_vals) else ""
                                if tail_text.strip():
                                    break
                                lower_text = ""
                                if row_idx + 1 < len(table_region.cells):
                                    lower_row = table_region.cells[row_idx + 1]
                                    lower_text = str(lower_row[col_idx + colspan]) if col_idx + colspan < len(lower_row) else ""
                                if colspan >= 2 and _clean_extracted_text(lower_text):
                                    break
                                colspan += 1
                        break
                    next_col += 1

            # Section rows often have a numeric marker in the first column and
            # one label spanning every remaining column. Some detectors assign
            # the label only to the columns its glyphs touch and leave trailing
            # cells empty, so merge the whole row tail when the tail contains
            # only one distinct non-empty text value.
            if not native_table and current_text.strip() and col_idx == 1 and row_vals:
                marker = _clean_extracted_text(str(row_vals[0])) if row_vals else ""
                tail = [
                    _clean_extracted_text(str(row_vals[c])) if c < len(row_vals) else ""
                    for c in range(1, cols)
                ]
                non_empty_tail = [value for value in tail if value]
                if (
                    re.fullmatch(r"\d{1,3}", marker or "")
                    and non_empty_tail
                    and len(set(non_empty_tail)) == 1
                    and _clean_extracted_text(current_text) == non_empty_tail[0]
                ):
                    colspan = cols - col_idx
             
            # Merge Logic: Horizontal Only + PDF Verification
            if current_text.strip() and not native_table:
                # Count in PDF
                # Use simple substring count. 
                # Normalize spaces for robust check?
                # For now simple count.
                pdf_count = pdf_row_text.count(current_text)
                
                # Count in Table Row
                table_count = table_row_counts.get(current_text, 1)
                
                # Check duplication
                if table_count > pdf_count and table_count > 1 and not is_numeric_cell(current_text):
                    # Likely an artifact -> Merge allowed
                    should_merge = True
                else:
                    # Real data duplicate -> Do not merge
                    should_merge = False
                
                if should_merge and colspan == 1:
                    # Calculate max mergeable span based on content match
                    while col_idx + colspan < cols:
                        next_text = ""
                        if row_idx < len(table_region.cells) and (col_idx + colspan) < len(table_region.cells[row_idx]):
                            next_text = str(table_region.cells[row_idx][col_idx + colspan])
                        
                        if next_text == current_text:
                            colspan += 1
                        else:
                            break
            
            # Mark visited
            for c in range(col_idx, col_idx + colspan):
                visited.add((row_idx, c))
            
            # Apply Merge
            cell = table.cell(row_idx, col_idx)
            
            if colspan > 1:
                try:
                    # Clear cells before merge
                    for mc in range(col_idx, col_idx + colspan):
                        table.cell(row_idx, mc).text = ""
                    # Merge Horizontal
                    right_cell = table.cell(row_idx, col_idx + colspan - 1)
                    cell.merge(right_cell)
                    # Remove extra paragraphs from merge
                    merged = table.cell(row_idx, col_idx)
                    while len(merged.paragraphs) > 1:
                        p_el = merged.paragraphs[-1]._element
                        p_el.getparent().remove(p_el)
                    logger.log(f"  Merged {colspan} cells at r{row_idx}c{col_idx} (Text: '{current_text}')")
                except Exception as e:
                    logger.log(f"  Merge failed at r{row_idx}c{col_idx}: {e}")
            
            # Set Text
            current_text = _compact_wrapped_marker_text(current_text)
            cell.text = current_text
            if _looks_like_table_index_marker(current_text) or is_numeric_cell(current_text):
                _set_cell_no_wrap(cell)
            cell_bbox = (0.0, 0.0, 0.0, 0.0)
            if row_idx < len(table_region.cell_bboxes) and col_idx < len(table_region.cell_bboxes[row_idx]):
                cell_bbox = tuple(table_region.cell_bboxes[row_idx][col_idx])
            row_page = row_source_pages[row_idx] if row_idx < len(row_source_pages) else table_region.page
            cell_page_lines = lines_by_page_for_table.get(int(row_page), pdf_lines_page)
            cell_bold = _cell_text_should_be_bold(cell_bbox, cell_page_lines) if any(cell_bbox) else False
            cell_font_size = table_font_size
            marker_for_fit = _clean_extracted_text(current_text)
            if (
                col_idx in (0, cols - 1)
                and len(marker_for_fit) > 1
                and (_looks_like_table_index_marker(marker_for_fit) or is_numeric_cell(marker_for_fit))
            ):
                cell_font_size = max(9, table_font_size - 3)
            
            # Formatting
            for para in cell.paragraphs:
                # Logic: Numeric -> Center, Text -> Left
                if is_numeric_cell(current_text):
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    
                if para.runs:
                    set_paragraph_font(para, font_size=cell_font_size, bold=cell_bold if cell_bold else None)
                else:
                    run = para.add_run()
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(cell_font_size)
                    if cell_bold:
                        run.font.bold = True

    if getattr(table_region, "enable_blank_continuation_vmerge", False):
        spans = _blank_continuation_vmerge_spans(table_region)
        _apply_explicit_cell_grid_borders(table)
        applied = _suppress_ooxml_internal_span_borders(table, spans, logger)
        if applied:
            logger.log(f"  Suppressed {applied} leading continuation border span(s)")
    
    # Vertical merges are opt-in. python-docx cell addressing can shift after
    # rowspan merges in dense scanned tables, which is worse than keeping
    # repeated/blank page-local cells readable.
    if (
        os.environ.get("OCRTOOL_ENABLE_VERTICAL_TABLE_MERGE") == "1"
        and not getattr(table_region, "disable_vertical_merge", False)
    ):
      for c in reversed(range(cols)):
        r = 0
        while r < rows - 1:
            text_r = table_region.cells[r][c].strip() if r < len(table_region.cells) and c < len(table_region.cells[r]) else ""
            text_r1 = table_region.cells[r + 1][c].strip() if r + 1 < len(table_region.cells) and c < len(table_region.cells[r + 1]) else ""
            if not text_r or not text_r1 or text_r != text_r1:
                r += 1
                continue
            # Check bbox alignment (same column boundaries)
            bbox_r = table_region.cell_bboxes[r][c] if r < len(table_region.cell_bboxes) and c < len(table_region.cell_bboxes[r]) else (0, 0, 0, 0)
            bbox_r1 = table_region.cell_bboxes[r + 1][c] if r + 1 < len(table_region.cell_bboxes) and c < len(table_region.cell_bboxes[r + 1]) else (0, 0, 0, 0)
            if not _same_visual_cell_region(tuple(bbox_r), tuple(bbox_r1)):
                r += 1
                continue
            # Count span
            span = 1
            while r + span < rows:
                next_text = table_region.cells[r + span][c].strip() if r + span < len(table_region.cells) and c < len(table_region.cells[r + span]) else ""
                if next_text != text_r:
                    break
                span += 1
            if span > 1:
                try:
                    # Clear ALL cells in merge range
                    for mr in range(r, r + span):
                        table.cell(mr, c).text = ""
                    top_cell = table.cell(r, c)
                    bottom_cell = table.cell(r + span - 1, c)
                    top_cell.merge(bottom_cell)
                    # Remove extra paragraphs from merge (keeps only first)
                    merged_cell = table.cell(r, c)
                    while len(merged_cell.paragraphs) > 1:
                        p_el = merged_cell.paragraphs[-1]._element
                        p_el.getparent().remove(p_el)
                    merged_cell.paragraphs[0].text = text_r
                    logger.log(f"  Vertical merge {span} cells at r{r}c{c} (Text: '{text_r[:30]}')")
                except Exception as e:
                    logger.log(f"  Vertical merge failed at r{r}c{c}: {e}")
            r += span

    logger.log(f"  Added table: {rows}x{cols}")


def refine_table_structure(tables: List[TableRegion], logger: Logger):
    """
    Refines table structure by enforcing geometric column consistency.
    Specifically targets rows where Column 0 and Column 1 are incorrectly merged
    (e.g., '2. Cong ty...' should be '2.' | 'Cong ty...').
    """
    for t_idx, table in enumerate(tables):
        # We need at least 2 columns to split Col 0 and Col 1
        if table.col_count < 2:
            continue
            
        # 1. Calculate Dominant Boundary for Column 0
        # Collect x2 of Col 0 for rows where Col 0 and Col 1 are DISTINCT
        col0_boundaries = []
        
        rows = table.row_count
        cols = table.col_count
        
        for r in range(rows):
            # Check if Col 0 and Col 1 are distinct
            # We use cell_bboxes to check.
            # cells_grid might have been populated or not.
            
            bbox0 = table.cell_bboxes[r][0]
            bbox1 = table.cell_bboxes[r][1]
            
            # If bbox0 != bbox1, they are distinct cells (visually)
            # And bbox0 must be valid
            if any(bbox0) and any(bbox1) and bbox0 != bbox1:
                # Store x2 of col 0
                col0_boundaries.append(bbox0[2])
                
        if not col0_boundaries:
            continue
            
        # Median boundary
        col0_boundaries.sort()
        dominant_x2 = col0_boundaries[len(col0_boundaries) // 2]
        
        # Tolerance ? say 5 points
        
        # 2. Iterate rows and checking for Merges crossing this boundary
        candidates = 0
        
        for r in range(rows):
            bbox0 = table.cell_bboxes[r][0]
            bbox1 = table.cell_bboxes[r][1]
            
            # Check if merged (same bbox)
            if any(bbox0) and bbox0 == bbox1:
                # This is a merge across Col 0 and 1.
                # Check if it SHOULD be split.
                
                # Condition A: Text implies a split? (e.g. "1. Text")
                # Condition B: The cell spans significantly across the dominant_x2
                # bbox0 is (x1, y1, x2, y2)
                
                cell_x1, cell_y1, cell_x2, cell_y2 = bbox0
                
                # Check coverage
                # If the cell starts before the boundary and ends well after it
                if cell_x1 < dominant_x2 - 10 and cell_x2 > dominant_x2 + 20: 
                    # Likely a missing separator
                    
                    # Split it!
                    candidates += 1
                    
                    # Update BBoxes
                    # New BBox 0: (x1, y1, dominant_x2, y2)
                    new_bbox0 = (cell_x1, cell_y1, dominant_x2, cell_y2)
                    
                    # New BBox 1: (dominant_x2, y1, x2, y2)
                    new_bbox1 = (dominant_x2, cell_y1, cell_x2, cell_y2)
                    
                    table.cell_bboxes[r][0] = new_bbox0
                    table.cell_bboxes[r][1] = new_bbox1
                    
                    # Also update subsequent columns if they were merged?
                    # If Col 2 was also merged (triple merge), we set Col 1 to be the rest?
                    # For now, just split Col 0 from the blob.
                    # If the blob was Col 0+1+2, now Col 0 is small, Col 1 is (Old - Small). Col 2 is (Old - Small) (still merged with 1)
                    
                    # Propagate changes to underlying logic (Text Extraction)
                    # We need to re-extract text for these two new cells.
                    # But we don't have the original PDF TextLines here easily accessible?
                    # Wait, we prefer to use the existing text if possible. 
                    
                    current_text = table.cells[r][0] # The merged text
                    
                    # Regex split heuristic
                    # Look for "Digits." or "Roman." at start
                    match = re.match(r'^(\d+\.|[IVX]+\.)\s+(.+)$', current_text, re.DOTALL)
                    if match:
                        txt0 = match.group(1)
                        txt1 = match.group(2)
                        table.cells[r][0] = txt0
                        table.cells[r][1] = txt1
                    else:
                        # Fallback: Just put it all in Col 1? Or keep in Col 0 (overflow)?
                        # If we split the cell, we should try to split text. 
                        # If regex fails, maybe it IS a real merge (header)?
                        # But geometry said it spans column 0.
                        # Leave text in Col 0, empty in Col 1? -> Result: Col 0 expands visuals? No, we fixed width.
                        # Let's start with Regex only.
                        pass
        
        if candidates > 0:
            logger.log(f"  Refined Table {t_idx+1}: Force-split {candidates} rows at x={dominant_x2:.1f}")


def create_docx_from_pdf(
    pdf_path: str,
    output_path: str,
    log_path: str = None,
    no_log_file: bool = False,
    metadata: dict = None
) -> Tuple[bool, str, str]:
    """
    Create a new DOCX from PDF text with positions preserved.
    Returns: (Success, Message, LogContent)
    """
    layout_profiles: Dict[str, dict] = {}
    estimated_first_line_indent_pt = 0.0

    def _default_layout_profile(orientation: str = "portrait") -> dict:
        if orientation == "landscape":
            page_w, page_h = 841.89, 595.28
        else:
            page_w, page_h = 595.28, 841.89
        return {
            "page_width_pt": page_w,
            "page_height_pt": page_h,
            "left_margin_pt": page_w * 0.12,
            "right_margin_pt": page_w * 0.08,
            "top_margin_pt": page_h * 0.07,
            "bottom_margin_pt": page_h * 0.07,
            "body_font_pt": 14.0,
            "table_font_pt": 13.0,
            "first_line_indent_pt": page_w * 0.045,
        }

    def _apply_standard_section_layout(section, orientation: str = "portrait") -> None:
        """Apply detected page size/margins; fall back to proportional A4."""
        profile = layout_profiles.get(orientation) or _default_layout_profile(orientation)
        section.page_width = Pt(profile["page_width_pt"])
        section.page_height = Pt(profile["page_height_pt"])
        section.top_margin = Pt(profile["top_margin_pt"])
        section.bottom_margin = Pt(profile["bottom_margin_pt"])
        section.left_margin = Pt(profile["left_margin_pt"])
        section.right_margin = Pt(profile["right_margin_pt"])

    def _estimate_docx_layout_profiles(
        page_info_local: dict,
        lines: List[TextLine],
        tables: List[TableRegion],
        paragraph_items: List[Tuple[str, TextLine, bool]],
    ) -> Dict[str, dict]:
        grouped: Dict[str, Dict[str, List[float]]] = {}
        table_bboxes_by_page: Dict[int, List[Tuple[float, float, float, float]]] = {}
        for table in tables or []:
            table_bboxes_by_page.setdefault(table.page, []).append(_table_bbox(table))

        for pg, info in (page_info_local or {}).items():
            page_w = float(info.get("width", 595.28) or 595.28)
            page_h = float(info.get("height", 841.89) or 841.89)
            orientation = "landscape" if page_w > page_h else "portrait"
            group = grouped.setdefault(
                orientation,
                {
                    "page_width": [],
                    "page_height": [],
                    "left": [],
                    "right_gap": [],
                    "top": [],
                    "bottom_gap": [],
                    "font": [],
                    "table_font": [],
                    "indent": [],
                },
            )
            group["page_width"].append(page_w)
            group["page_height"].append(page_h)

            page_lines = [
                line for line in lines
                if line.page == pg
                and not line.is_footnote
                and line.content_type != 4
                and not is_page_number_text(line.text)
                and line.width >= page_w * 0.08
            ]
            body_band = [
                line for line in page_lines
                if page_h * 0.04 <= line.y <= page_h * 0.96
            ]
            lefts = [line.x for line in body_band]
            rights = [line.x + line.width for line in body_band]
            tops = [line.y for line in body_band]
            bottoms = [line.y + line.height for line in body_band]
            fonts = [line.font_size for line in body_band if line.font_size > 0]
            for bbox in table_bboxes_by_page.get(pg, []):
                lefts.append(bbox[0])
                rights.append(bbox[2])
                tops.append(bbox[1])
                bottoms.append(bbox[3])

            if lefts and rights:
                left = _percentile(lefts, 0.10, page_w * 0.12)
                right = _percentile(rights, 0.90, page_w * 0.90)
                group["left"].append(_clamp(left, page_w * 0.025, page_w * 0.24))
                group["right_gap"].append(_clamp(page_w - right, page_w * 0.025, page_w * 0.24))
            if tops:
                group["top"].append(_clamp(_percentile(tops, 0.05, page_h * 0.07), page_h * 0.025, page_h * 0.18))
            if bottoms:
                bottom_gap = page_h - _percentile(bottoms, 0.95, page_h * 0.93)
                group["bottom_gap"].append(_clamp(bottom_gap, page_h * 0.025, page_h * 0.18))
            if fonts:
                group["font"].append(_median(fonts, 14.0))

        page_left_by_page: Dict[int, float] = {}
        for pg, info in (page_info_local or {}).items():
            page_w = float(info.get("width", 595.28) or 595.28)
            page_lines = [
                line for line in lines
                if line.page == pg
                and not line.is_footnote
                and line.content_type != 4
                and line.width >= page_w * 0.08
            ]
            if page_lines:
                page_left_by_page[pg] = _percentile([line.x for line in page_lines], 0.10, page_w * 0.12)

        for _, first_line, is_footnote in paragraph_items or []:
            if is_footnote:
                continue
            merged_lines = _get_merged_lines(first_line)
            continuation_lefts = [
                line.x for line in merged_lines[1:]
                if line.page == first_line.page and abs(line.x - first_line.x) > 1.0
            ]
            base_left = min(continuation_lefts) if continuation_lefts else page_left_by_page.get(first_line.page)
            if base_left is None:
                continue
            page_w = float((page_info_local or {}).get(first_line.page, {}).get("width", 595.28) or 595.28)
            indent = first_line.x - base_left
            if page_w * 0.012 <= indent <= page_w * 0.16:
                orientation = "landscape" if page_info_local.get(first_line.page, {}).get("width", 0) > page_info_local.get(first_line.page, {}).get("height", 1) else "portrait"
                if orientation not in grouped:
                    grouped[orientation] = {
                        "page_width": [],
                        "page_height": [],
                        "left": [],
                        "right_gap": [],
                        "top": [],
                        "bottom_gap": [],
                        "font": [],
                        "table_font": [],
                        "indent": [],
                    }
                grouped[orientation]["indent"].append(indent)

        profiles: Dict[str, dict] = {}
        for orientation in ("portrait", "landscape"):
            defaults = _default_layout_profile(orientation)
            group = grouped.get(orientation)
            if not group:
                profiles[orientation] = defaults
                continue
            page_w = _median(group.get("page_width", []), defaults["page_width_pt"])
            page_h = _median(group.get("page_height", []), defaults["page_height_pt"])
            body_font = _clamp(round(_median(group.get("font", []), 14.0)), 12.0, 14.0)
            table_font = _clamp(body_font - 1.0, 11.0, 13.0)
            profiles[orientation] = {
                "page_width_pt": page_w,
                "page_height_pt": page_h,
                "left_margin_pt": _median(group.get("left", []), defaults["left_margin_pt"]),
                "right_margin_pt": _median(group.get("right_gap", []), defaults["right_margin_pt"]),
                "top_margin_pt": defaults["top_margin_pt"],
                "bottom_margin_pt": defaults["bottom_margin_pt"],
                "body_font_pt": body_font,
                "table_font_pt": table_font,
                "first_line_indent_pt": _median(group.get("indent", []), defaults["first_line_indent_pt"]),
            }
        return profiles

    if not no_log_file and log_path is None:
        base = os.path.splitext(output_path)[0]
        log_path = base.replace("_final", "_merge") + ".log"
    
    if no_log_file:
        log_path = None

    logger = Logger(log_path)
    logger.log("=" * 60)
    logger.log("PDF to DOCX Position-Preserving Converter v4")
    logger.log("=" * 60)
    logger.log(f"PDF: {pdf_path}")
    logger.log(f"Output: {output_path}")
    
    try:
        # Step 1: Extract text lines with positions
        pdf_lines, page_info = extract_pdf_lines(pdf_path, logger)

        # Step 1b: Enrich lines with Screen AI metadata + layout regions from companion JSON
        layout_regions_by_page = {}  # page_num -> [regions]
        companion_data, companion_path = _load_canonical_companion_data(pdf_path, logger, "layout/enrichment")
        if companion_data:
            enrich_lines_from_json(pdf_lines, str(companion_path or pdf_path), logger)
            for _pi, _pg in enumerate(companion_data.get("pages", [])):
                _lr = _pg.get("layout_regions", [])
                if _lr:
                    layout_regions_by_page[_pi + 1] = _lr
            if layout_regions_by_page:
                logger.log(f"Loaded layout regions for {len(layout_regions_by_page)} pages")

        if not layout_regions_by_page:
            layout_regions_by_page = analyze_layout_regions_for_pdf(pdf_path, page_info, logger)
            if layout_regions_by_page:
                save_layout_regions_to_companion(
                    companion_path,
                    companion_data,
                    layout_regions_by_page,
                    logger,
                )

        if layout_regions_by_page:
            prepare_layout_region_reading_orders(layout_regions_by_page, page_info, logger)

        # Step 1c: Match lines to layout regions (assigns semantic_type)
        if layout_regions_by_page:
            try:
                from scanindex.core.tables.layout_analyzer import match_lines_to_regions
                for pg_num, regions in layout_regions_by_page.items():
                    pg_lines = [l for l in pdf_lines if l.page == pg_num]
                    pw = page_info.get(pg_num, {}).get("width", 595)
                    ph = page_info.get(pg_num, {}).get("height", 842)
                    # Prefer bbox_pdf written by OCR JSON decoration. This
                    # keeps semantic matching independent from render DPI.
                    if any(region.get("bbox_pdf") for region in regions):
                        match_regions = []
                        for region in regions:
                            bbox_pdf = region.get("bbox_pdf")
                            if not bbox_pdf:
                                continue
                            item = dict(region)
                            item["bbox"] = bbox_pdf
                            match_regions.append(item)
                        match_lines_to_regions(pg_lines, match_regions, 1.0, 1.0)
                    else:
                        # Legacy JSON fallback: bbox is in image pixels.
                        scale_x = pw / (pw / (72.0 / 200.0))  # = 72/200
                        scale_y = ph / (ph / (72.0 / 200.0))
                        match_lines_to_regions(pg_lines, regions, scale_x, scale_y)
                tagged = sum(1 for l in pdf_lines if l.semantic_type)
                logger.log(f"Tagged {tagged}/{len(pdf_lines)} lines with semantic types")
            except ImportError:
                logger.log("layout_analyzer not available, skipping semantic tagging")
            except Exception as e:
                logger.log(f"Semantic tagging failed: {e}")

        pdf_lines = filter_figure_ocr_noise(pdf_lines, layout_regions_by_page, logger)

        # Step 2: Detect tables. Digital PDFs get a native/vector-first path
        # so correct embedded text is not rewritten by OCR/model table repairs.
        table_cache_meta = _current_table_cache_meta(
            pdf_path,
            pdf_lines,
            layout_regions_by_page,
            companion_data,
        )
        table_regions = load_table_structures_from_companion(
            companion_data,
            table_cache_meta,
            logger,
        )
        if table_regions is not None:
            table_assigned_ids = assign_ocr_lines_to_table_cells_by_geometry(
                table_regions,
                pdf_lines,
                logger,
                rebuild_cells=False,
            )
        else:
            page_source_modes = _page_source_modes_from_companion(companion_data)
            digital_text_pdf = _is_digital_text_pdf(pdf_path)
            if digital_text_pdf and not page_source_modes:
                page_source_modes = {page: "digital" for page in page_info.keys()}

            native_candidate_pages = {
                page for page, mode in page_source_modes.items()
                if mode in {"digital", "mixed"}
            }
            native_tables: List[TableRegion] = []
            if native_candidate_pages:
                native_tables = [
                    table for table in detect_tables_pymupdf_native(pdf_path, logger, pdf_lines)
                    if int(getattr(table, "page", 0) or 0) in native_candidate_pages
                ]
                if native_tables:
                    native_tables = filter_false_positive_tables(native_tables, layout_regions_by_page, logger)
                else:
                    logger.log("Native table extraction found no candidate tables; model table pipeline may be used")

            model_candidate_pages = _model_table_candidate_pages(
                page_source_modes,
                layout_regions_by_page,
            )
            model_tables: List[TableRegion] = []
            if model_candidate_pages:
                model_tables = detect_tables(
                    pdf_path,
                    logger,
                    page_info,
                    _filter_lines_to_pages(pdf_lines, model_candidate_pages),
                    _filter_layout_regions_to_pages(layout_regions_by_page, model_candidate_pages),
                )

            if model_tables:
                model_layout_regions = _filter_layout_regions_to_pages(
                    layout_regions_by_page,
                    {int(getattr(table, "page", 0) or 0) for table in model_tables},
                )
                model_tables = repair_continued_tables(
                    model_tables, model_layout_regions, pdf_lines, page_info, logger
                )
                model_tables = split_stacked_tables(model_tables, logger)
                model_tables = postprocess_table_layout_grids(model_tables, model_layout_regions, logger)
                model_tables = trim_empty_trailing_table_columns(model_tables, logger)
                model_tables = stabilize_page_local_table_grids(
                    model_tables, model_layout_regions, page_info, logger
                )
                model_tables = stabilize_continuation_table_schemas_from_geometry(
                    model_tables, model_layout_regions, page_info, pdf_lines, logger
                )
                model_tables = filter_false_positive_tables(model_tables, model_layout_regions, logger)
                model_tables = repair_table_row_gaps_from_ocr(model_tables, pdf_lines, logger)
                table_assigned_ids = _assign_table_lines_preserving_native_text(
                    native_tables,
                    model_tables,
                    pdf_lines,
                    logger,
                )
                try:
                    from scanindex.core.tables.postprocess_v2 import postprocess_tables_v2

                    postprocess_tables_v2(model_tables, pdf_lines, logger)
                except Exception as exc:
                    logger.log(f"V2 table postprocess failed, keeping V1 cells: {exc}")
                model_tables = repair_shifted_leading_table_cells(model_tables, logger)
                model_tables = fuse_page_local_continuation_rows(model_tables, logger)
                model_tables = stabilize_continued_table_column_schemas(
                    model_tables, model_layout_regions, page_info, pdf_lines, logger
                )
                if _continued_table_flow_enabled():
                    model_tables = compose_continued_tables_for_word_flow(
                        model_tables, model_layout_regions, page_info, pdf_lines, logger
                    )
                model_tables = repair_section_heading_rows(model_tables, logger)
                model_tables = repair_split_table_header_fragments(model_tables, logger)
                model_tables = repair_wrapped_group_header_tokens(model_tables, logger)
                for table in model_tables:
                    if not getattr(table, "skip_render", False):
                        setattr(table, "enable_blank_continuation_vmerge", True)
                table_regions = _merge_native_and_model_tables(native_tables, model_tables, logger)
            else:
                table_regions = list(native_tables)
                table_assigned_ids = assign_ocr_lines_to_table_cells_by_geometry(
                    table_regions,
                    pdf_lines,
                    logger,
                    rebuild_cells=False,
                )

            save_table_structures_to_companion(
                companion_path,
                companion_data,
                table_regions,
                table_cache_meta,
                logger,
            )

        # --- NEW: Refine Structure ---
        # User requested to disable custom geometric refinement.
        # refine_table_structure(table_regions, logger)

        # Create table lookup by page and y-range
        table_map = {}  # page -> [(y_top, y_bottom, table)]
        for table in table_regions:
            if table.page not in table_map:
                table_map[table.page] = []
            table_map[table.page].append((table.y_top, table.y_bottom, table))

        # Step 3: Create new document
        doc = Document()

        # Set A4 page and administrative-document margins.
        for section in doc.sections:
            _apply_standard_section_layout(section, "portrait")

        # Set default font and paragraph spacing
        style = doc.styles['Normal']
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = Pt(20)

        def _apply_word_flow_paragraph_format(
            para,
            *,
            is_footnote: bool = False,
            is_centered: bool = False,
            is_caption: bool = False,
            use_first_line_indent: bool = False,
        ) -> None:
            pf = para.paragraph_format
            if is_footnote:
                pf.line_spacing = Pt(14)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.first_line_indent = None
                return

            if is_centered or is_caption:
                pf.line_spacing = Pt(18)
                pf.space_before = Pt(3)
                pf.space_after = Pt(3)
                pf.first_line_indent = None
                return

            pf.line_spacing = Pt(20)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.first_line_indent = (
                Pt(estimated_first_line_indent_pt)
                if use_first_line_indent and estimated_first_line_indent_pt > 0.5
                else None
            )

        # Step 4: Filter out table-assigned lines, then merge the rest into paragraphs
        filtered_lines = []
        unresolved_table_lines = []
        preserved_table_residuals = 0
        for line in pdf_lines:
            if id(line) in table_assigned_ids:
                residuals = getattr(line, "_table_residual_fragments", []) or []
                if residuals:
                    filtered_lines.extend(residuals)
                    preserved_table_residuals += len(residuals)
                continue
            if _looks_like_qr_access_artifact(line.text):
                logger.log(f"Skipped QR/link access artifact line on page {line.page}")
                continue

            inside_rendered_table = False
            inside_skipped_table = False
            if line.page in table_map:
                for y_top, y_bottom, table in table_map[line.page]:
                    if y_top <= line.y + line.height / 2.0 <= y_bottom and _line_overlaps_table_area(line, table):
                        if getattr(table, "skip_render", False):
                            inside_skipped_table = True
                        else:
                            inside_rendered_table = True
                        break
            if inside_skipped_table:
                continue
            if inside_rendered_table:
                unresolved_table_lines.append(line)
                continue
            filtered_lines.append(line)

        if unresolved_table_lines:
            logger.log(f"Skipped {len(unresolved_table_lines)} unresolved table-overlap lines after table mapping")
        if preserved_table_residuals:
            logger.log(f"Kept {preserved_table_residuals} non-table text fragment(s) from table-crossing OCR lines")

        filtered_lines = split_digital_dual_header_spans(filtered_lines, page_info, logger)
        filtered_lines = filter_publication_running_edge_lines(filtered_lines, page_info, logger)
        filtered_lines = split_lines_crossing_multi_column_gutters(filtered_lines, page_info, logger)
        filtered_lines = apply_multi_column_reading_order(filtered_lines, page_info, logger)
        if layout_regions_by_page:
            filtered_lines = apply_layout_region_reading_order(
                filtered_lines,
                layout_regions_by_page,
                page_info,
                logger,
            )

        # Detect dual-column headers before merging
        dual_headers = detect_dual_column_headers(filtered_lines, page_info, logger, layout_regions_by_page)
        dual_footers = detect_dual_column_footers(filtered_lines, page_info, logger)
        dual_header_line_ids = set()
        for pg, (left, right) in dual_headers.items():
            for l in left + right:
                dual_header_line_ids.add(id(l))
        dual_footer_line_ids = set()
        for pg, (left, right) in dual_footers.items():
            for l in left + right:
                dual_footer_line_ids.add(id(l))
        # Remove dual-column header/footer lines from paragraph merging
        dual_column_line_ids = dual_header_line_ids | dual_footer_line_ids
        if dual_column_line_ids:
            filtered_lines = [l for l in filtered_lines if id(l) not in dual_column_line_ids]

        # Merge lines into paragraphs
        paragraphs = merge_lines_to_paragraphs(filtered_lines, page_info, logger)

        # Step 4b: Compute statistics from actual data for adaptive formatting
        body_heights = [l.height for l in filtered_lines if not l.is_footnote and l.height > 0]
        median_h = sorted(body_heights)[len(body_heights) // 2] if body_heights else 20

        layout_profiles = _estimate_docx_layout_profiles(page_info, filtered_lines, table_regions, paragraphs)
        portrait_profile = layout_profiles.get("portrait") or _default_layout_profile("portrait")
        estimated_first_line_indent_pt = float(portrait_profile.get("first_line_indent_pt", 0.0))
        for section in doc.sections:
            _apply_standard_section_layout(section, "portrait")
        style.font.size = Pt(float(portrait_profile.get("body_font_pt", 14.0)))
        for pg, info in page_info.items():
            orientation = "landscape" if info.get("width", 0) > info.get("height", 1) else "portrait"
            profile = layout_profiles.get(orientation) or _default_layout_profile(orientation)
            info["docx_left_margin_pt"] = profile["left_margin_pt"]
            info["docx_right_margin_pt"] = profile["right_margin_pt"]
            info["docx_body_font_pt"] = profile["body_font_pt"]
            info["docx_table_font_pt"] = profile["table_font_pt"]
        logger.log(
            "Estimated DOCX layout: "
            + "; ".join(
                f"{name} page={profile['page_width_pt']:.1f}x{profile['page_height_pt']:.1f} "
                f"margins L/R/T/B={profile['left_margin_pt']:.1f}/"
                f"{profile['right_margin_pt']:.1f}/{profile['top_margin_pt']:.1f}/"
                f"{profile['bottom_margin_pt']:.1f}, indent={profile['first_line_indent_pt']:.1f}"
                for name, profile in sorted(layout_profiles.items())
            )
        )

        # OCR foreground gray is normalized at document level for word-level
        # emphasis. Page stats are kept as a secondary guard for uneven scans.
        gray_values_by_page: Dict[int, List[float]] = {}
        doc_gray_values: List[float] = []
        for line in pdf_lines:
            for word in _word_items_for_line(line):
                gray = word.get("fg_gray", 128)
                if _has_known_gray(gray):
                    gray_f = float(gray)
                    gray_values_by_page.setdefault(line.page, []).append(gray_f)
                    doc_gray_values.append(gray_f)
        style_stats_by_page: Dict[int, dict] = {}
        doc_dark_cutoff = _percentile(doc_gray_values, 0.12, 128.0)
        doc_q25 = _percentile(doc_gray_values, 0.25, 128.0)
        doc_q75 = _percentile(doc_gray_values, 0.75, 128.0)
        doc_median = _median(doc_gray_values, 128.0)
        doc_iqr = max(doc_q75 - doc_q25, 1.0)
        doc_style_stats = {
            "doc_known_gray_count": len(doc_gray_values),
            "doc_gray_median": doc_median,
            "doc_gray_iqr": doc_iqr,
            "doc_bold_gray_cutoff": doc_dark_cutoff if len(doc_gray_values) >= 80 else None,
            "doc_bold_z_cutoff": _gray_z(doc_dark_cutoff, doc_median, doc_iqr) if len(doc_gray_values) >= 80 else 0.0,
        }
        for pg, values in gray_values_by_page.items():
            page_dark_cutoff = _percentile(values, 0.12, 128.0)
            q25 = _percentile(values, 0.25, 128.0)
            q75 = _percentile(values, 0.75, 128.0)
            style_stats_by_page[pg] = {
                **doc_style_stats,
                "known_gray_count": len(values),
                "gray_median": _median(values, 128.0),
                "page_bold_gray_cutoff": page_dark_cutoff if len(values) >= 40 else None,
                "gray_q25": q25,
                "gray_iqr": max(q75 - q25, 1.0),
            }
        for line in pdf_lines:
            setattr(line, "_page_style_stats", style_stats_by_page.get(line.page, doc_style_stats))

        # Margins: reuse margin_map computed by merge_lines_to_paragraphs
        # (same logic already used for paragraph merging)
        computed_margins = {}  # page -> (left, right)
        page_lines_map = {}
        for line in filtered_lines:
            if line.page not in page_lines_map:
                page_lines_map[line.page] = []
            page_lines_map[line.page].append(line)
        for pg, p_lines in page_lines_map.items():
            from collections import Counter
            x_rounded = [round(l.x / 2.0) * 2.0 for l in p_lines]
            common = Counter(x_rounded).most_common(1)
            left = common[0][0] if common else 0
            all_rights = sorted([l.x + l.width for l in p_lines])
            right = all_rights[int(len(all_rights) * 0.98)] if all_rights else page_info.get(pg, {}).get("width", 595)
            computed_margins[pg] = (left, right)

        # Step 5: Add paragraphs and tables to document using a unified sorted list
        # This ensures strict order based on Y position

        doc_elements = []
        use_reading_order = any(getattr(l, "order", 0) > 0 for l in pdf_lines)

        # Add paragraphs
        for para_text, first_line, is_footnote in paragraphs:
            doc_elements.append({
                "type": "para",
                "page": first_line.page,
                "y": first_line.y,
                "order": first_line.order,
                "data": (para_text, first_line, is_footnote)
            })

        # Add tables
        for table in table_regions:
            if getattr(table, "skip_render", False):
                continue
            table_lines = [
                l for l in pdf_lines
                if l.page == table.page and table.y_top <= l.y_center <= table.y_bottom
            ]
            table_order = min((l.order for l in table_lines), default=1000000)
            table_order = _layout_order_for_bbox(
                table.page,
                _matching_layout_bbox_for_table(table, layout_regions_by_page) or _table_bbox(table),
                layout_regions_by_page,
                table_order,
            )
            doc_elements.append({
                "type": "table",
                "page": table.page,
                "y": table.y_top,
                "order": table_order,
                "data": table
            })

        # Add figure regions from layout analysis
        for pg_num, regions in layout_regions_by_page.items():
            for r_idx, region in enumerate(regions):
                if region["type"] == "figure":
                    bbox_pdf = region.get("bbox_pdf", region["bbox"])
                    fig_bbox = tuple(float(v) for v in bbox_pdf[:4])
                    nearby_qr_text = False
                    for line in pdf_lines:
                        if line.page != pg_num or not _looks_like_qr_access_artifact(line.text):
                            continue
                        lx0, ly0, lx1, ly1 = _line_bbox(line)
                        x_overlap = max(0.0, min(lx1, fig_bbox[2]) - max(lx0, fig_bbox[0]))
                        near_y = ly0 <= fig_bbox[3] + 45.0 and ly1 >= fig_bbox[1] - 20.0
                        if _bbox_overlap_ratio((lx0, ly0, lx1, ly1), fig_bbox) > 0.01 or (near_y and x_overlap > 0):
                            nearby_qr_text = True
                            break
                    if nearby_qr_text:
                        logger.log(f"Skipped QR/link access figure on page {pg_num}")
                        continue
                    pg_width = page_info.get(pg_num, {}).get("width", 595)
                    fig_center_x = (fig_bbox[0] + fig_bbox[2]) / 2
                    if pg_num in dual_footers:
                        footer_left, footer_right = dual_footers[pg_num]
                        footer_lines = footer_left + footer_right
                        footer_y0 = min((line.y for line in footer_lines), default=fig_bbox[1]) - 30
                        footer_y1 = max((line.y + line.height for line in footer_lines), default=fig_bbox[3]) + 50
                        overlaps_footer_band = fig_bbox[1] <= footer_y1 and fig_bbox[3] >= footer_y0
                        if (
                            any(_bbox_overlap_ratio(_line_bbox(line), fig_bbox) > 0.05 for line in footer_right)
                            or (fig_center_x >= pg_width * 0.45 and overlaps_footer_band)
                        ):
                            logger.log(f"Skipped signature/stamp figure on page {pg_num}")
                            continue
                    else:
                        anchors = [
                            line for line in pdf_lines
                            if line.page == pg_num and _unaccent_upper(line.text).startswith("NOI NHAN")
                        ]
                        if anchors:
                            anchor = sorted(anchors, key=lambda line: line.y)[-1]
                            if fig_center_x >= pg_width * 0.45 and fig_bbox[1] >= anchor.y - 30:
                                logger.log(f"Skipped signature/stamp figure on page {pg_num}")
                                continue
                    doc_elements.append({
                        "type": "figure",
                        "page": pg_num,
                        "y": bbox_pdf[1],
                        "order": _layout_order_for_bbox(
                            pg_num,
                            fig_bbox,
                            layout_regions_by_page,
                            1000000 + int(bbox_pdf[1]),
                        ),
                        "data": {"bbox": region["bbox"], "page": pg_num, "idx": r_idx, "bbox_pdf": fig_bbox}
                    })

        # Add dual-column headers
        for pg, (left, right) in dual_headers.items():
            top_y = min(l.y for l in left + right)
            doc_elements.append({
                "type": "dual_header",
                "page": pg,
                "y": top_y,
                "order": min(l.order for l in left + right),
                "data": (left, right)
            })

        # Add dual-column recipient/signature footers
        for pg, (left, right) in dual_footers.items():
            top_y = min(l.y for l in left + right)
            doc_elements.append({
                "type": "dual_footer",
                "page": pg,
                "y": top_y,
                "order": min(l.order for l in left + right),
                "data": (left, right)
            })

        # Sort by OCR reading order when canonical JSON is available. Pure Y sorting
        # interleaves two-column signature/recipient blocks and hurts editable text
        # accuracy; fallback PDFs without JSON keep the original coordinate sort.
        if use_reading_order:
            doc_elements.sort(key=lambda x: (x["page"], x.get("order", 1000000), x["y"]))
        else:
            doc_elements.sort(key=lambda x: (x["page"], x["y"]))

        # Open source PDF for figure extraction (if layout regions have figures)
        doc_pdf = None
        has_figures = any(e["type"] == "figure" for e in doc_elements)
        if has_figures:
            try:
                # Prefer original input (higher quality than OCR overlay)
                original_pdf = pdf_path.replace("_ocr.pdf", ".pdf")
                doc_pdf = fitz.open(original_pdf if os.path.exists(original_pdf) else pdf_path)
            except Exception:
                doc_pdf = None

        # Detect page orientations from PDF page dimensions
        page_orientations = {}  # page_num -> 'portrait' or 'landscape'
        for pg, info in page_info.items():
            page_orientations[pg] = 'landscape' if info["width"] > info["height"] else 'portrait'

        tables_added_count = 0
        first_elem_page = doc_elements[0]["page"] if doc_elements else 1
        current_orientation = page_orientations.get(first_elem_page, 'portrait')
        for section in doc.sections:
            _apply_standard_section_layout(section, current_orientation)
        last_rendered_page = None
        table_render_pages = {
            int(table.page)
            for table in table_regions
            if not getattr(table, "skip_render", False)
        }
        table_bottom_pages = {
            int(table.page)
            for table in table_regions
            if (
                not getattr(table, "skip_render", False)
                and float(table.y_bottom) >= float(page_info.get(table.page, {}).get("height", 842)) * 0.86
            )
        }

        for elem_index, elem in enumerate(doc_elements):
            elem_page = elem["page"]
            needed_orient = page_orientations.get(elem_page, 'portrait')

            page_changed = last_rendered_page is not None and elem_page != last_rendered_page
            orientation_changed = needed_orient != current_orientation
            table_boundary = (
                page_changed
                and (
                    int(elem_page) in table_render_pages
                    or int(last_rendered_page) in table_render_pages
                )
            )
            page_break_likely_implicit = page_changed and int(last_rendered_page) in table_bottom_pages
            if (
                page_changed
                and not orientation_changed
                and (
                    os.environ.get("OCRTOOL_PRESERVE_SOURCE_PAGE_BREAKS") == "1"
                    or (table_boundary and not page_break_likely_implicit)
                )
            ):
                doc.add_page_break()
                logger.log(f"Inserted page break before source page {elem_page}")
            elif orientation_changed:
                from docx.enum.section import WD_ORIENT, WD_SECTION
                new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                if needed_orient == 'landscape':
                    new_section.orientation = WD_ORIENT.LANDSCAPE
                    _apply_standard_section_layout(new_section, "landscape")
                else:
                    new_section.orientation = WD_ORIENT.PORTRAIT
                    _apply_standard_section_layout(new_section, "portrait")
                current_orientation = needed_orient
                logger.log(f"Page {elem_page}: switched to {needed_orient}")
            last_rendered_page = elem_page

            if elem["type"] == "para":
                para_text, first_line, is_footnote = elem["data"]
                para_text = _repair_drop_cap_join(para_text, first_line)

                # Skip separators (horizontal rules)
                if first_line.content_type == 4:
                    continue

                # Note: "abandon" from layout model is unreliable (low conf, false positives)
                # e.g. "ĐẢNG CỘNG SẢN VIỆT NAM" falsely tagged as abandon
                # → don't skip any text based on layout model alone

                # Use layout/KIE signals for emphasis, but keep alignment
                # governed by the original line geometry.
                sem = (first_line.semantic_type or "").strip().lower()
                numbered_heading = _looks_like_numbered_heading(para_text)
                is_list_item = _looks_like_list_item(para_text)
                is_doc_subject = "DOC_SUBJECT" in getattr(first_line, "kie_labels", set())

                # Detect centered from geometry only. Layout "title" is a style
                # hint; section headings can be tagged title while still being
                # visually left-aligned in the source.
                left_m, right_m = computed_margins.get(first_line.page, (0, page_info.get(first_line.page, {}).get("width", 595)))
                content_width = right_m - left_m
                if content_width > 0:
                    line_center = first_line.x + first_line.width / 2
                    page_center = left_m + content_width / 2
                    is_short = first_line.width < content_width * 0.75
                    left_gap = max(0.0, first_line.x - left_m)
                    right_gap = max(0.0, right_m - (first_line.x + first_line.width))
                    is_centered = (
                        is_short
                        and abs(line_center - page_center) < content_width * 0.08
                        and abs(left_gap - right_gap) < content_width * 0.14
                    )
                else:
                    is_short = False
                    is_centered = False

                # Relative font size from line height
                # Only heading-like lines get larger font
                is_bold = _looks_like_visual_bold_heading(
                    para_text,
                    sem,
                    numbered_heading,
                    is_doc_subject,
                    is_centered,
                )
                if (
                    not is_bold
                    and is_centered
                    and not is_list_item
                    and len(_clean_extracted_text(para_text)) <= 180
                    and _line_visual_bold_ratio(first_line) >= 0.65
                ):
                    is_bold = True
                ratio = first_line.height / median_h if median_h > 0 else 1.0
                base_font_size = float(page_info.get(first_line.page, {}).get("docx_body_font_pt", 14.0))
                is_heading = (
                    numbered_heading
                    or is_doc_subject
                    or (is_bold and (sem == "title" or _mostly_uppercase_text(para_text)))
                )
                if is_heading and ratio > 1.4:
                    para_font_size = base_font_size + 4
                elif is_heading and ratio > 1.15:
                    para_font_size = base_font_size + 2
                else:
                    para_font_size = base_font_size

                # Determine alignment
                para = doc.add_paragraph()
                if (sem == "figure_caption" or sem == "table_caption"):
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    _apply_word_flow_paragraph_format(
                        para,
                        is_footnote=is_footnote,
                        is_centered=True,
                        is_caption=True,
                    )
                elif is_centered and not numbered_heading:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    _apply_word_flow_paragraph_format(
                        para,
                        is_footnote=is_footnote,
                        is_centered=True,
                    )
                elif is_heading:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    _apply_word_flow_paragraph_format(
                        para,
                        is_footnote=is_footnote,
                        use_first_line_indent=True,
                    )
                elif is_list_item:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    _apply_word_flow_paragraph_format(
                        para,
                        is_footnote=is_footnote,
                        use_first_line_indent=True,
                    )
                else:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    _apply_word_flow_paragraph_format(
                        para,
                        is_footnote=is_footnote,
                        use_first_line_indent=True,
                    )

                # Italic detection:
                # - Handwritten/signature content type
                # - Parenthesized text like "(Kèm theo Công văn...)" = subtitle/caption
                is_handwritten = first_line.content_type in (1, 8)
                # Italic: only from content_type (handwritten/signature)
                is_italic = is_handwritten

                # Format text with superscripts
                add_text_with_superscripts(para, para_text, first_line, is_footnote,
                                          bold=is_bold, font_size_pt=para_font_size,
                                          italic=is_italic)

            elif elem["type"] == "dual_header":
                left_lines, right_lines = elem["data"]
                section = doc.sections[-1]
                content_width_pt = max(
                    float(section.page_width.pt - section.left_margin.pt - section.right_margin.pt),
                    120.0,
                )
                add_dual_header_paragraphs(doc, left_lines, right_lines, content_width_pt, logger)

            elif elem["type"] == "dual_footer":
                left_lines, right_lines = elem["data"]
                section = doc.sections[-1]
                content_width_pt = max(
                    float(section.page_width.pt - section.left_margin.pt - section.right_margin.pt),
                    120.0,
                )
                add_dual_header_table(doc, left_lines, right_lines, content_width_pt, logger)

            elif elem["type"] == "figure":
                # Extract figure image from PDF page
                try:
                    fig_data = elem["data"]
                    fig_page = doc_pdf[fig_data["page"] - 1]  # 0-based
                    # Render page at 200 DPI and crop
                    fig_pix = fig_page.get_pixmap(dpi=200)
                    from PIL import Image as PILImage
                    fig_img = PILImage.frombytes("RGB", [fig_pix.width, fig_pix.height], fig_pix.samples)
                    if fig_data.get("bbox_pdf"):
                        fx0, fy0, fx1, fy1 = (float(v) for v in fig_data["bbox_pdf"][:4])
                        sx = fig_pix.width / max(float(fig_page.rect.width), 1.0)
                        sy = fig_pix.height / max(float(fig_page.rect.height), 1.0)
                        bx = (fx0 * sx, fy0 * sy, fx1 * sx, fy1 * sy)
                    else:
                        bx = fig_data["bbox"]  # image pixel coords
                    # Crop with padding
                    pad = 5
                    crop_box = (
                        max(0, int(bx[0]) - pad), max(0, int(bx[1]) - pad),
                        min(fig_img.width, int(bx[2]) + pad), min(fig_img.height, int(bx[3]) + pad)
                    )
                    cropped = fig_img.crop(crop_box)
                    # Save temp and insert
                    import tempfile
                    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                    cropped.save(tmp.name)
                    tmp.close()
                    # Size: fit within page width
                    fig_w_inches = cropped.width / 200  # 200 DPI
                    max_w = 5.5  # inches, safe for portrait
                    if fig_w_inches > max_w:
                        fig_w_inches = max_w
                    doc.add_picture(tmp.name, width=Inches(fig_w_inches))
                    os.unlink(tmp.name)
                    logger.log(f"  Added figure from page {fig_data['page']}")
                except Exception as e:
                    logger.log(f"  Figure extraction failed: {e}")

            elif elem["type"] == "table":
                table = elem["data"]
                source_pages = set(getattr(table, "source_pages", []) or [table.page])
                page_lines = [l for l in pdf_lines if l.page in source_pages]
                add_table_to_doc(doc, table, page_lines, logger, page_info)
                tables_added_count += 1
                if (
                    elem_index < len(doc_elements) - 1
                    and doc_elements[elem_index + 1]["page"] == elem_page
                ):
                    doc.add_paragraph()  # Space after table when more content follows
        
        # Cleanup
        if doc_pdf:
            doc_pdf.close()

        # Step 5: Embed metadata into DOCX properties
        if metadata:
            try:
                props = doc.core_properties
                if metadata.get("co_quan_ban_hanh"):
                    props.author = metadata["co_quan_ban_hanh"]
                if metadata.get("trich_yeu"):
                    props.title = metadata["trich_yeu"]
                if metadata.get("loai_van_ban"):
                    props.subject = metadata["loai_van_ban"]
                if metadata.get("so_ky_hieu"):
                    props.keywords = metadata["so_ky_hieu"]
            except Exception as e:
                logger.log(f"Warning: Failed to set DOCX properties: {e}")

        # Step 6: Save
        doc.save(output_path)
        
        # Summary
        logger.log("\n" + "=" * 60)
        logger.log("SUMMARY")
        logger.log("=" * 60)
        logger.log(f"Text lines: {len(pdf_lines)}")
        logger.log(f"Tables: {tables_added_count}")
        logger.log(f"Output: {output_path}")
        logger.save()
        
        return True, f"Success: {output_path}", logger.get_log_text()
        
    except Exception as e:
        import traceback
        logger.log(f"ERROR: {traceback.format_exc()}")
        logger.save()
        return False, str(e), logger.get_log_text()


def create_final_docx_v2(base_path: str) -> Tuple[bool, str, str]:
    """Convenience function."""
    for suffix in ["_ocr.pdf", "_final.docx"]:
        if base_path.endswith(suffix):
            base_path = base_path[:-len(suffix)]
            break
    
    return create_docx_from_pdf(
        base_path + "_ocr.pdf",
        base_path + "_final.docx"
    )


if __name__ == "__main__":
    import sys
    if len(sys.argv) >= 2:
        success, msg, logs = create_final_docx_v2(sys.argv[1])
        print(f"\n{'SUCCESS' if success else 'FAILED'}: {msg}")
    else:
        print("Usage: python table_anchored_merger.py <base_path>")
