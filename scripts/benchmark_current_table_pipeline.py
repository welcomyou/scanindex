from __future__ import annotations

import argparse
import copy
import csv
import json
import re
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from rapidfuzz.distance import Levenshtein

from table_eval_metrics import compare_table_grid_lists
from table_anchored_merger import (
    TextLine,
    assign_ocr_lines_to_table_cells_by_geometry,
    create_docx_from_pdf,
    detect_tables,
    enrich_lines_from_json,
    extract_pdf_lines,
    filter_false_positive_tables,
    filter_figure_ocr_noise,
    postprocess_table_layout_grids,
    repair_continued_tables,
    split_stacked_tables,
)
from table_postprocess_v2 import postprocess_tables_v2


class QuietLogger:
    def __init__(self):
        self.lines: list[str] = []

    def log(self, msg: str):
        self.lines.append(str(msg))


@dataclass
class ScanTiming:
    preprocess_sec: float
    raw_detect_sec: float
    postprocess_sec: float
    table_total_sec: float
    full_docx_sec: float


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "")
    return text.strip()


def docx_tables(path: Path) -> list[tuple[int, int, str]]:
    doc = Document(str(path))
    out = []
    for table in doc.tables:
        rows = len(table.rows)
        cols = len(table.columns)
        text = clean_text(" ".join(cell.text for row in table.rows for cell in row.cells))
        out.append((rows, cols, text))
    return out


def docx_table_grids(path: Path) -> list[list[list[str]]]:
    doc = Document(str(path))
    out: list[list[list[str]]] = []
    for table in doc.tables:
        out.append([[clean_text(cell.text) for cell in row.cells] for row in table.rows])
    return out


def data_tables(tables: list[tuple[int, int, str]]) -> list[tuple[int, int, str]]:
    return [t for t in tables if t[:2] != (1, 2)]


def data_table_grids(tables: list[list[list[str]]]) -> list[list[list[str]]]:
    return [t for t in tables if not (len(t) == 1 and max((len(row) for row in t), default=0) == 2)]


def table_region_data(tables) -> list[tuple[int, int, str]]:
    out = []
    for table in tables:
        if getattr(table, "skip_render", False):
            continue
        rows = int(getattr(table, "row_count", 0) or 0)
        cols = int(getattr(table, "col_count", 0) or 0)
        text = clean_text(" ".join(str(cell) for row in getattr(table, "cells", []) for cell in row))
        out.append((rows, cols, text))
    return out


def table_region_grids(tables) -> list[list[list[str]]]:
    out: list[list[list[str]]] = []
    for table in tables:
        if getattr(table, "skip_render", False):
            continue
        rows = int(getattr(table, "row_count", 0) or 0)
        cols = int(getattr(table, "col_count", 0) or 0)
        raw_cells = getattr(table, "cells", []) or []
        grid = [
            [
                clean_text(raw_cells[r][c]) if r < len(raw_cells) and c < len(raw_cells[r]) else ""
                for c in range(cols)
            ]
            for r in range(rows)
        ]
        out.append(grid)
    return out


def shape_score(gt: tuple[int, int], out: tuple[int, int]) -> float:
    gr, gc = gt
    or_, oc = out
    if gr <= 0 or gc <= 0 or or_ <= 0 or oc <= 0:
        return 0.0
    return (min(gr, or_) / max(gr, or_)) * (min(gc, oc) / max(gc, oc))


def text_acc(gt: str, out: str) -> float:
    gt = clean_text(gt)
    out = clean_text(out)
    if not gt and not out:
        return 1.0
    if not gt:
        return 0.0
    return max(0.0, 1.0 - Levenshtein.distance(gt, out) / len(gt))


def compare_tables(gt_tables: list[tuple[int, int, str]], out_tables: list[tuple[int, int, str]]) -> dict[str, Any]:
    n = max(len(gt_tables), len(out_tables))
    details = []
    shape_scores = []
    text_scores = []
    for idx in range(n):
        gt = gt_tables[idx] if idx < len(gt_tables) else (0, 0, "")
        out = out_tables[idx] if idx < len(out_tables) else (0, 0, "")
        ss = shape_score(gt[:2], out[:2])
        ta = text_acc(gt[2], out[2])
        shape_scores.append(ss)
        text_scores.append(ta)
        details.append({
            "index": idx,
            "gt_shape": list(gt[:2]),
            "out_shape": list(out[:2]),
            "shape_score": ss,
            "text_acc": ta,
        })
    return {
        "count": len(out_tables),
        "gt_count": len(gt_tables),
        "shape_acc": sum(shape_scores) / n if n else 1.0,
        "text_acc": sum(text_scores) / n if n else 1.0,
        "details": details,
    }


def compare_table_cells(gt_tables: list[list[list[str]]], out_tables: list[list[list[str]]]) -> dict[str, Any]:
    return compare_table_grid_lists(gt_tables, out_tables)


def load_layout_regions_by_page(json_path: Path) -> dict[int, list[dict[str, Any]]]:
    data = json.loads(json_path.read_text(encoding="utf-8"))
    return {
        page_idx: page.get("layout_regions", [])
        for page_idx, page in enumerate(data.get("pages", []), 1)
    }


def prepare_pipeline_inputs(ocr_pdf: Path, logger: QuietLogger) -> tuple[list[TextLine], dict[int, list[dict[str, Any]]], dict]:
    pdf_lines, page_info = extract_pdf_lines(str(ocr_pdf), logger)
    json_path = Path(str(ocr_pdf) + ".json")
    layout_regions_by_page: dict[int, list[dict[str, Any]]] = {}
    if json_path.exists():
        enrich_lines_from_json(pdf_lines, str(json_path), logger)
        layout_regions_by_page = load_layout_regions_by_page(json_path)
        if layout_regions_by_page:
            try:
                from layout_analyzer import match_lines_to_regions

                for pg_num, regions in layout_regions_by_page.items():
                    pg_lines = [line for line in pdf_lines if line.page == pg_num]
                    pw = page_info.get(pg_num, {}).get("width", 595)
                    ph = page_info.get(pg_num, {}).get("height", 842)
                    scale_x = pw / (pw / (72.0 / 200.0))
                    scale_y = ph / (ph / (72.0 / 200.0))
                    match_lines_to_regions(pg_lines, regions, scale_x, scale_y)
            except Exception as exc:
                logger.log(f"Semantic tagging failed: {exc}")
    pdf_lines = filter_figure_ocr_noise(pdf_lines, layout_regions_by_page, logger)
    return pdf_lines, layout_regions_by_page, page_info


def postprocess_current_tables(tables, layout_regions_by_page, pdf_lines, page_info, logger: QuietLogger):
    processed = repair_continued_tables(tables, layout_regions_by_page, pdf_lines, page_info, logger)
    processed = split_stacked_tables(processed, logger)
    processed = postprocess_table_layout_grids(processed, layout_regions_by_page, logger)
    processed = filter_false_positive_tables(processed, layout_regions_by_page, logger)
    assign_ocr_lines_to_table_cells_by_geometry(processed, pdf_lines, logger)
    processed = postprocess_tables_v2(processed, pdf_lines, logger)
    return [table for table in processed if not getattr(table, "skip_render", False)]


def groundtruth_docx_for_scan(scan: str, groundtruth_dir: Path, gt03_docx: Path) -> Path:
    if scan == "03":
        return gt03_docx
    return groundtruth_dir / f"groundtruth{scan}.docx"


def run_scan(scan: str, args) -> dict[str, Any]:
    ocr_pdf = args.ocr_dir / f"scan{scan}_ocr.pdf"
    logger = QuietLogger()

    prep_t0 = time.perf_counter()
    pdf_lines, layout_regions_by_page, page_info = prepare_pipeline_inputs(ocr_pdf, logger)
    preprocess_sec = time.perf_counter() - prep_t0

    raw_t0 = time.perf_counter()
    raw_tables = detect_tables(str(ocr_pdf), logger, page_info, pdf_lines, layout_regions_by_page)
    raw_detect_sec = time.perf_counter() - raw_t0
    raw_data = table_region_data(raw_tables)
    raw_sources = [getattr(table, "source", "") for table in raw_tables]

    post_t0 = time.perf_counter()
    post_tables = postprocess_current_tables(copy.deepcopy(raw_tables), layout_regions_by_page, pdf_lines, page_info, logger)
    postprocess_sec = time.perf_counter() - post_t0

    full_out_dir = args.out_dir / "current_full_docx"
    full_out_dir.mkdir(parents=True, exist_ok=True)
    full_docx = full_out_dir / f"scan{scan}_final.docx"
    full_t0 = time.perf_counter()
    full_ok, full_msg, full_logs = create_docx_from_pdf(str(ocr_pdf), str(full_docx), no_log_file=True)
    full_docx_sec = time.perf_counter() - full_t0

    gt_data = data_tables(docx_tables(groundtruth_docx_for_scan(scan, args.groundtruth_dir, args.gt03_docx)))
    post_data = table_region_data(post_tables)
    full_data = data_tables(docx_tables(full_docx)) if full_docx.exists() else []

    timings = ScanTiming(
        preprocess_sec=preprocess_sec,
        raw_detect_sec=raw_detect_sec,
        postprocess_sec=postprocess_sec,
        table_total_sec=raw_detect_sec + postprocess_sec,
        full_docx_sec=full_docx_sec,
    )

    return {
        "scan": scan,
        "raw": compare_tables(gt_data, raw_data),
        "postprocess": compare_tables(gt_data, post_data),
        "final_docx": compare_tables(gt_data, full_data),
        "gt_shapes": [list(t[:2]) for t in gt_data],
        "raw_shapes": [list(t[:2]) for t in raw_data],
        "postprocess_shapes": [list(t[:2]) for t in post_data],
        "final_docx_shapes": [list(t[:2]) for t in full_data],
        "raw_sources": raw_sources,
        "timing": timings.__dict__,
        "full_docx_ok": bool(full_ok),
        "full_docx_msg": full_msg,
        "log_tail": logger.lines[-80:],
        "full_log_tail": full_logs.splitlines()[-80:] if isinstance(full_logs, str) else [],
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--ocr-dir", type=Path, default=Path(r"D:\App\ocrtool\temp\groundtruth5_pipeline"))
    parser.add_argument("--groundtruth-dir", type=Path, default=Path(r"C:\Users\nhquan\Downloads\groundtruth"))
    parser.add_argument("--gt03-docx", type=Path, default=Path(r"D:\App\ocrtool\temp\groundtruth4_scan_word\groundtruth03_converted.docx"))
    parser.add_argument("--out-dir", type=Path, default=Path(r"D:\App\ocrtool\temp\cpu_table_bench_current"))
    parser.add_argument("--scans", nargs="+", default=["01", "02", "03", "04", "05"])
    args = parser.parse_args()

    args.out_dir.mkdir(parents=True, exist_ok=True)
    scans = [str(scan).zfill(2) for scan in args.scans]
    comparison = [run_scan(scan, args) for scan in scans]

    result = {
        "model": "current doclayout + gmft_onnx + img2table pipeline",
        "comparison": comparison,
    }
    (args.out_dir / "comparison.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")

    with (args.out_dir / "summary.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            "scan",
            "gt_data_count",
            "raw_count",
            "post_count",
            "final_count",
            "raw_shape_acc",
            "post_shape_acc",
            "final_shape_acc",
            "raw_text_acc",
            "post_text_acc",
            "final_text_acc",
            "preprocess_sec",
            "raw_detect_sec",
            "postprocess_sec",
            "table_total_sec",
            "full_docx_sec",
            "raw_sources",
        ])
        for item in comparison:
            timing = item["timing"]
            writer.writerow([
                item["scan"],
                item["raw"]["gt_count"],
                item["raw"]["count"],
                item["postprocess"]["count"],
                item["final_docx"]["count"],
                item["raw"]["shape_acc"],
                item["postprocess"]["shape_acc"],
                item["final_docx"]["shape_acc"],
                item["raw"]["text_acc"],
                item["postprocess"]["text_acc"],
                item["final_docx"]["text_acc"],
                timing["preprocess_sec"],
                timing["raw_detect_sec"],
                timing["postprocess_sec"],
                timing["table_total_sec"],
                timing["full_docx_sec"],
                "|".join(item["raw_sources"]),
            ])

    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
