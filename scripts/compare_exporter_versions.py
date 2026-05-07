"""
Compare OCR PDF -> DOCX exporters across two code worktrees.

The script can run both exporters against the same OCR PDFs, then compares the
DOCX outputs with ground-truth DOCX files using full-text and table-focused
metrics. It is intentionally data-driven: no scan-specific table rules are used.
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import re
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document

try:
    from rapidfuzz.distance import Levenshtein
except Exception:  # pragma: no cover - slow fallback for portable use
    Levenshtein = None


SCANS = ("01", "02", "03", "04")


def _norm_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _words(text: str) -> list[str]:
    return re.findall(r"\S+", text or "")


def _levenshtein(a: Any, b: Any) -> int:
    if Levenshtein is not None:
        return int(Levenshtein.distance(a, b))

    if len(a) < len(b):
        a, b = b, a
    prev = list(range(len(b) + 1))
    for i, ca in enumerate(a, 1):
        cur = [i]
        for j, cb in enumerate(b, 1):
            cur.append(min(prev[j] + 1, cur[-1] + 1, prev[j - 1] + (ca != cb)))
        prev = cur
    return prev[-1]


def _cell_text(cell) -> str:
    return _norm_text(cell.text)


def _table_matrix(table) -> list[list[str]]:
    return [[_cell_text(cell) for cell in row.cells] for row in table.rows]


def _table_text(matrix: list[list[str]]) -> str:
    return "\n".join("\t".join(row) for row in matrix)


def _shape(matrix: list[list[str]]) -> list[int]:
    return [len(matrix), max((len(row) for row in matrix), default=0)]


def _shape_score(gt_shape: list[int], out_shape: list[int]) -> float:
    if gt_shape == out_shape:
        return 1.0
    gr, gc = gt_shape
    orow, oc = out_shape
    row_score = min(gr, orow) / max(gr, orow, 1)
    col_score = min(gc, oc) / max(gc, oc, 1)
    return row_score * col_score


def read_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraphs = [_norm_text(p.text) for p in doc.paragraphs if _norm_text(p.text)]
    tables = [_table_matrix(t) for t in doc.tables]
    table_texts = [_table_text(t) for t in tables]
    full_text = "\n".join(paragraphs + table_texts)
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "table_shapes": [_shape(t) for t in tables],
        "table_texts": table_texts,
        "full_text": _norm_text(full_text),
    }


def text_accuracy(gt: str, out: str) -> dict[str, float]:
    gt = _norm_text(gt)
    out = _norm_text(out)
    char_den = max(1, len(gt))
    gt_words = _words(gt)
    out_words = _words(out)
    word_den = max(1, len(gt_words))
    cer = _levenshtein(gt, out) / char_den
    wer = _levenshtein(gt_words, out_words) / word_den
    return {
        "char_acc": max(0.0, 1.0 - cer),
        "word_acc": max(0.0, 1.0 - wer),
        "cer": cer,
        "wer": wer,
    }


def compare_tables(gt_doc: dict[str, Any], out_doc: dict[str, Any]) -> dict[str, Any]:
    gt_tables = gt_doc["tables"]
    out_tables = out_doc["tables"]
    n = max(len(gt_tables), len(out_tables))
    rows = []
    for idx in range(n):
        if idx >= len(gt_tables):
            rows.append({
                "index": idx,
                "gt_shape": [0, 0],
                "out_shape": _shape(out_tables[idx]),
                "shape_score": 0.0,
                "text_acc": 0.0,
            })
            continue
        if idx >= len(out_tables):
            rows.append({
                "index": idx,
                "gt_shape": _shape(gt_tables[idx]),
                "out_shape": [0, 0],
                "shape_score": 0.0,
                "text_acc": 0.0,
            })
            continue

        gt_text = _norm_text(_table_text(gt_tables[idx]))
        out_text = _norm_text(_table_text(out_tables[idx]))
        text_acc = max(0.0, 1.0 - (_levenshtein(gt_text, out_text) / max(1, len(gt_text))))
        rows.append({
            "index": idx,
            "gt_shape": _shape(gt_tables[idx]),
            "out_shape": _shape(out_tables[idx]),
            "shape_score": _shape_score(_shape(gt_tables[idx]), _shape(out_tables[idx])),
            "text_acc": text_acc,
        })

    if not rows:
        return {"count": 0, "shape_acc": 1.0, "text_acc": 1.0, "details": []}

    return {
        "count": len(out_tables),
        "gt_count": len(gt_tables),
        "shape_acc": sum(r["shape_score"] for r in rows) / len(rows),
        "text_acc": sum(r["text_acc"] for r in rows) / len(rows),
        "details": rows,
    }


def run_exporter(worktree: Path, input_pdf: Path, output_docx: Path) -> None:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    code = (
        "from table_anchored_merger import create_docx_from_pdf; "
        "ok,msg,log=create_docx_from_pdf(r'%s', r'%s', no_log_file=False); "
        "print(ok); print(msg); "
        "raise SystemExit(0 if ok else 1)"
    ) % (str(input_pdf), str(output_docx))
    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"
    env.setdefault("GMFT_ONNX_THREADS", "4")
    subprocess.run([sys.executable, "-c", code], cwd=str(worktree), env=env, check=True)


def gt_paths(groundtruth_dir: Path, gt03_docx: Path | None) -> dict[str, Path]:
    paths = {
        "01": groundtruth_dir / "groundtruth01.docx",
        "02": groundtruth_dir / "groundtruth02.docx",
        "04": groundtruth_dir / "groundtruth04.docx",
    }
    if gt03_docx:
        paths["03"] = gt03_docx
    else:
        paths["03"] = groundtruth_dir / "groundtruth03.docx"
    return paths


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--old-worktree", type=Path, required=True)
    parser.add_argument("--new-worktree", type=Path, required=True)
    parser.add_argument("--ocr-dir", type=Path, required=True)
    parser.add_argument("--groundtruth-dir", type=Path, required=True)
    parser.add_argument("--gt03-docx", type=Path)
    parser.add_argument("--out-dir", type=Path, required=True)
    parser.add_argument("--skip-run", action="store_true")
    args = parser.parse_args()

    old_dir = args.out_dir / "old"
    new_dir = args.out_dir / "new"
    gt = gt_paths(args.groundtruth_dir, args.gt03_docx)

    if not args.skip_run:
        for scan in SCANS:
            input_pdf = args.ocr_dir / f"scan{scan}_ocr.pdf"
            run_exporter(args.old_worktree, input_pdf, old_dir / f"scan{scan}.docx")
            run_exporter(args.new_worktree, input_pdf, new_dir / f"scan{scan}.docx")

    results = []
    for scan in SCANS:
        gt_doc = read_docx(gt[scan])
        old_doc = read_docx(old_dir / f"scan{scan}.docx")
        new_doc = read_docx(new_dir / f"scan{scan}.docx")
        old_text = text_accuracy(gt_doc["full_text"], old_doc["full_text"])
        new_text = text_accuracy(gt_doc["full_text"], new_doc["full_text"])
        old_tables = compare_tables(gt_doc, old_doc)
        new_tables = compare_tables(gt_doc, new_doc)
        results.append({
            "scan": scan,
            "old": {
                **old_text,
                "tables": old_tables,
                "table_shapes": old_doc["table_shapes"],
            },
            "new": {
                **new_text,
                "tables": new_tables,
                "table_shapes": new_doc["table_shapes"],
            },
            "groundtruth_table_shapes": gt_doc["table_shapes"],
        })

    args.out_dir.mkdir(parents=True, exist_ok=True)
    json_path = args.out_dir / "comparison.json"
    csv_path = args.out_dir / "summary.csv"
    json_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=[
            "scan",
            "old_char_acc",
            "new_char_acc",
            "old_word_acc",
            "new_word_acc",
            "old_table_count",
            "new_table_count",
            "gt_table_count",
            "old_table_shape_acc",
            "new_table_shape_acc",
            "old_table_text_acc",
            "new_table_text_acc",
        ])
        writer.writeheader()
        for item in results:
            writer.writerow({
                "scan": item["scan"],
                "old_char_acc": item["old"]["char_acc"],
                "new_char_acc": item["new"]["char_acc"],
                "old_word_acc": item["old"]["word_acc"],
                "new_word_acc": item["new"]["word_acc"],
                "old_table_count": item["old"]["tables"]["count"],
                "new_table_count": item["new"]["tables"]["count"],
                "gt_table_count": item["old"]["tables"]["gt_count"],
                "old_table_shape_acc": item["old"]["tables"]["shape_acc"],
                "new_table_shape_acc": item["new"]["tables"]["shape_acc"],
                "old_table_text_acc": item["old"]["tables"]["text_acc"],
                "new_table_text_acc": item["new"]["tables"]["text_acc"],
            })

    print(json.dumps(results, ensure_ascii=False, indent=2))
    print(f"Wrote {json_path}")
    print(f"Wrote {csv_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
